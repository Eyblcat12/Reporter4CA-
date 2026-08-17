"""Run isolated, reproducible Reporter Pro DOCX benchmarks.

Every trial runs in a fresh Python process. This gives process-cold measurements
and a trustworthy per-trial peak working set while explicitly leaving the Windows
filesystem cache uncontrolled.
"""

from __future__ import annotations

import argparse
import ctypes
import hashlib
import importlib.metadata
import json
import os
import platform
import subprocess
import sys
import time
import uuid
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "apps" / "backend"
DEFAULT_MANIFEST = ROOT / "tests" / "fixtures" / "performance" / "manifest.json"
DEFAULT_ARTIFACT_ROOT = ROOT / "artifacts" / "benchmarks"
sys.path.insert(0, str(BACKEND))

from core.performance_metrics import (  # noqa: E402
    PerformanceMetrics,
    current_rss_mib,
    peak_rss_mib,
)

REPORT_TYPES = (
    "full",
    "server_only",
    "client_only",
    "summary",
    "technical",
    "incident_response",
)
FEATURE_FLAG_NAMES = (
    "AUTO_REPORT_PERF_METRICS",
    "AUTO_REPORT_COMPACT_PROTOTYPE",
    "AUTO_REPORT_PREPARED_TEMPLATE",
    "AUTO_REPORT_FAST_CELL",
    "AUTO_REPORT_UNIFIED_SCHEDULER",
    "AUTO_REPORT_PREVIEW_JOBS",
    "AUTO_REPORT_PREVIEW_CACHE",
)
DEPENDENCIES = ("python-docx", "pandas", "openpyxl", "docxtpl")
DEFAULT_TEMPLATE_BY_REPORT_TYPE = {
    "full": BACKEND / "templates" / "report_template.docx",
    "server_only": BACKEND / "templates" / "server_only" / "report_server_only_default.docx",
    "client_only": BACKEND / "templates" / "client_only" / "report_client_only_default.docx",
    "summary": BACKEND / "templates" / "summary" / "report_summary_default.docx",
    "technical": BACKEND / "templates" / "technical" / "report_technical_default.docx",
    "incident_response": BACKEND / "templates" / "report_template.docx",
}


class BenchmarkConfigurationError(ValueError):
    """Raised before an expensive trial when the benchmark input is not controlled."""


class BenchmarkCorrectnessError(RuntimeError):
    """Raised when a generated report differs from the fixture expectation."""


def atomic_write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    temporary.replace(path)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def percentile(values: Iterable[float], percent: float) -> float | None:
    ordered = sorted(float(value) for value in values)
    if not ordered:
        return None
    position = (len(ordered) - 1) * percent
    lower = int(position)
    upper = min(lower + 1, len(ordered) - 1)
    fraction = position - lower
    return ordered[lower] + (ordered[upper] - ordered[lower]) * fraction


def summarize_samples(values: Iterable[float]) -> dict[str, Any]:
    samples = [float(value) for value in values]
    if not samples:
        return {
            "sampleCount": 0,
            "min": None,
            "p50": None,
            "p95": None,
            "p95Published": False,
            "max": None,
        }
    return {
        "sampleCount": len(samples),
        "min": round(min(samples), 3),
        "p50": round(percentile(samples, 0.50) or 0.0, 3),
        "p95": round(percentile(samples, 0.95) or 0.0, 3) if len(samples) >= 10 else None,
        "p95Published": len(samples) >= 10,
        "max": round(max(samples), 3),
    }


def load_manifest(path: Path) -> dict[str, Any]:
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise BenchmarkConfigurationError(f"Fixture manifest does not exist: {path}") from exc
    except json.JSONDecodeError as exc:
        raise BenchmarkConfigurationError(f"Fixture manifest is not valid JSON: {path}") from exc
    if payload.get("schemaVersion") != 1 or not isinstance(payload.get("fixtures"), list):
        raise BenchmarkConfigurationError("Unsupported performance fixture manifest")
    return payload


def select_fixture(manifest: dict[str, Any], selector: str) -> dict[str, Any]:
    for fixture in manifest["fixtures"]:
        if fixture.get("id") == selector or fixture.get("profile") == selector:
            return fixture
    raise BenchmarkConfigurationError(f"Fixture profile/id is not in manifest: {selector}")


def resolve_fixture_path(manifest_path: Path, fixture: dict[str, Any]) -> Path:
    fixture_root = manifest_path.resolve().parent
    fixture_path = (fixture_root / str(fixture.get("file", ""))).resolve()
    try:
        fixture_path.relative_to(fixture_root)
    except ValueError as exc:
        raise BenchmarkConfigurationError("Fixture path escapes the manifest directory") from exc
    if not fixture_path.is_file():
        raise BenchmarkConfigurationError(f"Fixture file does not exist: {fixture_path}")
    return fixture_path


def validate_fixture(
    manifest_path: Path,
    fixture: dict[str, Any],
    *,
    report_type: str,
) -> tuple[Path, dict[str, Any]]:
    if report_type not in fixture.get("validReportTypes", []):
        raise BenchmarkConfigurationError(
            f"Fixture {fixture.get('id')} does not support report type {report_type}"
        )
    expected_by_type = fixture.get("expectedByReportType")
    if not isinstance(expected_by_type, dict) or report_type not in expected_by_type:
        raise BenchmarkConfigurationError("Fixture does not declare report-type expectations")
    fixture_path = resolve_fixture_path(manifest_path, fixture)
    actual_size = fixture_path.stat().st_size
    actual_hash = sha256_file(fixture_path)
    if actual_size != int(fixture.get("inputBytes", -1)):
        raise BenchmarkConfigurationError(
            f"Fixture size mismatch for {fixture.get('id')}: {actual_size}"
        )
    if actual_hash != fixture.get("sha256"):
        raise BenchmarkConfigurationError(f"Fixture SHA-256 mismatch for {fixture.get('id')}")
    return fixture_path, expected_by_type[report_type]


def resolve_template(report_type: str, template: Path | None) -> Path:
    path = (template or DEFAULT_TEMPLATE_BY_REPORT_TYPE[report_type]).resolve()
    if not path.is_file():
        raise BenchmarkConfigurationError(f"Template does not exist: {path}")
    if not zipfile.is_zipfile(path):
        raise BenchmarkConfigurationError(f"Template is not a valid DOCX package: {path}")
    return path


def _manifest_projection(manifest: dict[str, Any]) -> dict[str, Any]:
    return {
        "assetCount": int(manifest.get("assetCount", 0)),
        "serverCount": sum(
            item.get("assetType") == "server" for item in manifest.get("assets", [])
        ),
        "clientCount": sum(
            item.get("assetType") == "client" for item in manifest.get("assets", [])
        ),
        "findingCount": int(manifest.get("findingCount", 0)),
        "evidenceCount": int(manifest.get("evidenceCount", 0)),
        "assessmentCounts": dict(sorted(manifest.get("assessmentCounts", {}).items())),
    }


def assert_expected_manifest(actual: dict[str, Any], expected: dict[str, Any]) -> None:
    projection = _manifest_projection(actual)
    normalized_expected = {
        "assetCount": int(expected["assetCount"]),
        "serverCount": int(expected["serverCount"]),
        "clientCount": int(expected["clientCount"]),
        "findingCount": int(expected["findingCount"]),
        "evidenceCount": int(expected["evidenceCount"]),
        "assessmentCounts": dict(sorted(expected["assessmentCounts"].items())),
    }
    if projection != normalized_expected:
        differing_keys = [
            key for key in normalized_expected if projection.get(key) != normalized_expected[key]
        ]
        raise BenchmarkCorrectnessError(
            "Generated report manifest differs from fixture expectation: "
            + ", ".join(differing_keys)
        )


def run_trial(config: dict[str, Any]) -> dict[str, Any]:
    """Execute one benchmark trial in the current process."""

    # A direct worker invocation must execute the same feature flags recorded in
    # its config instead of merely echoing them into benchmark metadata.
    for name, value in dict(config.get("featureFlags", {})).items():
        if name in FEATURE_FLAG_NAMES:
            os.environ[name] = str(value)

    from core.docx_field_updater import refresh_docx_fields
    from core.input_parser import parse_input
    from core.report_generator import ReportType, generate_report
    from core.report_integrity import verify_report_document
    from docx import Document

    fixture_path = Path(config["fixturePath"]).resolve()
    template_path = Path(config["templatePath"]).resolve()
    output_path = Path(config["outputDocx"]).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    fixture_hash = sha256_file(fixture_path)
    template_hash = sha256_file(template_path)
    expected = config["expected"]
    metrics = PerformanceMetrics(
        run_id=str(config["runId"]),
        metadata={
            "trial": int(config["trial"]),
            "fixtureId": str(config["fixtureId"]),
            "fixtureProfile": str(config["fixtureProfile"]),
            "fixtureHash": fixture_hash,
            "templateHash": template_hash,
            "reportType": str(config["reportType"]),
            "assetCount": int(expected["assetCount"]),
            "serverCount": int(expected["serverCount"]),
            "clientCount": int(expected["clientCount"]),
            "cacheState": "process-cold/cache-miss",
            "auditMode": "docx-reopen",
            "wordFieldUpdater": "requested" if config["updateFields"] else "disabled",
            "featureFlags": config.get("featureFlags", {}),
        },
    )
    baseline_rss = current_rss_mib()
    product_started_ns = time.perf_counter_ns()
    cpu_started = time.process_time()

    with metrics.phase("inputParse"):
        payload = parse_input(fixture_path, default_section="clients")

    parsed_counts = {
        "serverCount": len(payload.get("servers", [])),
        "clientCount": len(payload.get("clients", [])),
    }
    if parsed_counts != {
        "serverCount": int(config["fixtureServerCount"]),
        "clientCount": int(config["fixtureClientCount"]),
    }:
        raise BenchmarkCorrectnessError("Parsed server/client counts differ from fixture manifest")

    with metrics.phase("documentBuild"):
        document = generate_report(
            payload,
            title=f"Reporter benchmark {config['fixtureId']}",
            organization="Reporter Pro benchmark",
            assessment_date="2026-07-30",
            template_path=template_path,
            report_type=ReportType(str(config["reportType"])),
            metrics=metrics,
        )
    generated_manifest = getattr(document, "_reporter_manifest", {})
    assert_expected_manifest(generated_manifest, expected)

    save_attributes: dict[str, Any] = {}
    with metrics.phase("saveZip", attributes=save_attributes):
        document.save(output_path)
        save_attributes["bytes"] = output_path.stat().st_size

    field_update = {"updated": False, "engine": "disabled"}
    if config["updateFields"]:
        with metrics.phase("wordFieldUpdate"):
            update = refresh_docx_fields(
                output_path,
                timeout_seconds=int(config["fieldUpdateTimeoutSeconds"]),
            )
        field_update = {
            "updated": bool(update.updated),
            "engine": str(update.engine),
        }

    product_finished_ns = time.perf_counter_ns()
    product_final_rss = current_rss_mib()
    product_peak_rss = peak_rss_mib()
    audit_started_ns = time.perf_counter_ns()
    with metrics.phase("packageValidation", latency_class="audit"):
        if not zipfile.is_zipfile(output_path):
            raise BenchmarkCorrectnessError("Generated output is not a valid DOCX ZIP package")
    reopen_attributes: dict[str, Any] = {}
    with metrics.phase("docxReopen", latency_class="audit", attributes=reopen_attributes):
        reopened = Document(output_path)
        reopen_attributes["count"] = len(reopened.paragraphs) + len(reopened.tables)
    with metrics.phase("semanticIntegrity", latency_class="audit"):
        reopened_integrity = verify_report_document(reopened, generated_manifest)
    with metrics.phase("artifactHash", latency_class="audit"):
        output_hash = sha256_file(output_path)
    audit_finished_ns = time.perf_counter_ns()

    integrity = getattr(document, "_reporter_integrity", {})
    if not integrity.get("valid"):
        raise BenchmarkCorrectnessError("In-memory report integrity result is not valid")
    finding_verification_applicable = str(config["reportType"]) in {
        "technical",
        "incident_response",
    }
    expected_findings = int(reopened_integrity.get("expectedFindings", 0))
    missing_findings = len(reopened_integrity.get("missingFindings", []))
    return {
        "schemaVersion": 1,
        "status": "passed",
        "trial": int(config["trial"]),
        "runId": str(config["runId"]),
        "metrics": metrics.public(
            product_latency_ms=(product_finished_ns - product_started_ns) / 1_000_000,
            audit_latency_ms=(audit_finished_ns - audit_started_ns) / 1_000_000,
        ),
        "resources": {
            "cpuTimeMs": round((time.process_time() - cpu_started) * 1000, 3),
            "rssBaselineMiB": round(baseline_rss, 3),
            "productFinalRssMiB": round(product_final_rss, 3),
            "productPeakRssMiB": round(product_peak_rss, 3),
            "rssFinalMiB": round(current_rss_mib(), 3),
            "peakRssMiB": round(peak_rss_mib(), 3),
        },
        "artifact": {
            "bytes": output_path.stat().st_size,
            "sha256": output_hash,
            "retained": bool(config["retainDocx"]),
        },
        "reportManifest": _manifest_projection(generated_manifest),
        "reopenedIntegrity": {
            "valid": bool(reopened_integrity.get("valid")),
            "expectedAssets": int(reopened_integrity.get("expectedAssets", 0)),
            "verifiedAssets": int(reopened_integrity.get("verifiedAssets", 0)),
            "expectedFindings": expected_findings,
            "findingVerificationApplicable": finding_verification_applicable,
            "verifiedFindings": (
                expected_findings - missing_findings if finding_verification_applicable else None
            ),
            "missingFindingCount": missing_findings,
        },
        "fieldUpdate": field_update,
    }


def _run_worker(config_path: Path, output_path: Path) -> int:
    try:
        config = json.loads(config_path.read_text(encoding="utf-8"))
        result = run_trial(config)
    except Exception as exc:
        result = {
            "schemaVersion": 1,
            "status": "failed",
            "error": {
                "code": "benchmark_trial_failed",
                "type": type(exc).__name__,
            },
        }
        atomic_write_json(output_path, result)
        print(f"Benchmark worker failed: {type(exc).__name__}", file=sys.stderr)
        return 1
    atomic_write_json(output_path, result)
    return 0


def _git_metadata() -> dict[str, Any]:
    result = {"commit": "unknown", "dirty": None}
    try:
        commit = subprocess.run(
            ["git", "rev-parse", "HEAD"],
            cwd=ROOT,
            capture_output=True,
            text=True,
            timeout=10,
            check=False,
        )
        if commit.returncode == 0:
            result["commit"] = commit.stdout.strip()
        dirty = subprocess.run(
            ["git", "status", "--porcelain", "--untracked-files=no"],
            cwd=ROOT,
            capture_output=True,
            text=True,
            timeout=10,
            check=False,
        )
        if dirty.returncode == 0:
            result["dirty"] = bool(dirty.stdout.strip())
    except (OSError, subprocess.SubprocessError):
        pass
    return result


def _total_memory_mib() -> int | None:
    if os.name == "nt":

        class MemoryStatus(ctypes.Structure):
            _fields_ = [
                ("dwLength", ctypes.c_ulong),
                ("dwMemoryLoad", ctypes.c_ulong),
                ("ullTotalPhys", ctypes.c_ulonglong),
                ("ullAvailPhys", ctypes.c_ulonglong),
                ("ullTotalPageFile", ctypes.c_ulonglong),
                ("ullAvailPageFile", ctypes.c_ulonglong),
                ("ullTotalVirtual", ctypes.c_ulonglong),
                ("ullAvailVirtual", ctypes.c_ulonglong),
                ("ullAvailExtendedVirtual", ctypes.c_ulonglong),
            ]

        status = MemoryStatus()
        status.dwLength = ctypes.sizeof(status)
        if ctypes.windll.kernel32.GlobalMemoryStatusEx(ctypes.byref(status)):
            return round(status.ullTotalPhys / 1024 / 1024)
        return None
    try:
        pages = os.sysconf("SC_PHYS_PAGES")
        page_size = os.sysconf("SC_PAGE_SIZE")
        return round(pages * page_size / 1024 / 1024)
    except (AttributeError, OSError, ValueError):
        return None


def collect_environment() -> dict[str, Any]:
    dependency_versions: dict[str, str] = {}
    for dependency in DEPENDENCIES:
        try:
            dependency_versions[dependency] = importlib.metadata.version(dependency)
        except importlib.metadata.PackageNotFoundError:
            dependency_versions[dependency] = "not-installed"
    return {
        "os": {
            "system": platform.system(),
            "release": platform.release(),
            "version": platform.version(),
            "machine": platform.machine(),
        },
        "python": platform.python_version(),
        "cpu": {
            "model": platform.processor() or "unknown",
            "logicalCount": os.cpu_count(),
        },
        "totalMemoryMiB": _total_memory_mib(),
        "dependencies": dependency_versions,
        "git": _git_metadata(),
    }


def aggregate_trials(trials: list[dict[str, Any]]) -> dict[str, Any]:
    passed = [trial for trial in trials if trial.get("status") == "passed"]
    phase_samples: dict[str, list[float]] = {}
    aggregate_samples: dict[tuple[str, str], dict[str, list[float]]] = {}
    for trial in passed:
        for phase in trial["metrics"]["phases"]:
            phase_samples.setdefault(phase["name"], []).append(float(phase["durationMs"]))
        for aggregate in trial["metrics"].get("aggregates", []):
            key = (str(aggregate["name"]), str(aggregate["category"]))
            samples = aggregate_samples.setdefault(
                key,
                {"count": [], "totalDurationMs": [], "maxDurationMs": []},
            )
            samples["count"].append(float(aggregate["count"]))
            samples["totalDurationMs"].append(float(aggregate["totalDurationMs"]))
            samples["maxDurationMs"].append(float(aggregate["maxDurationMs"]))
    return {
        "trialCount": len(trials),
        "passedTrials": len(passed),
        "failedTrials": len(trials) - len(passed),
        "productLatencyMs": summarize_samples(
            trial["metrics"]["productLatencyMs"] for trial in passed
        ),
        "auditLatencyMs": summarize_samples(trial["metrics"]["auditLatencyMs"] for trial in passed),
        "productPeakRssMiB": summarize_samples(
            trial["resources"].get("productPeakRssMiB", trial["resources"]["peakRssMiB"])
            for trial in passed
        ),
        "peakRssMiB": summarize_samples(trial["resources"]["peakRssMiB"] for trial in passed),
        "outputBytes": summarize_samples(trial["artifact"]["bytes"] for trial in passed),
        "phases": {
            name: summarize_samples(samples) for name, samples in sorted(phase_samples.items())
        },
        "aggregates": {
            f"{name}:{category}": {
                field: summarize_samples(values) for field, values in samples.items()
            }
            for (name, category), samples in sorted(aggregate_samples.items())
        },
    }


def run_benchmark(args: argparse.Namespace) -> tuple[dict[str, Any], Path]:
    if args.trials < 1 or args.trials > 100:
        raise BenchmarkConfigurationError("trials must be between 1 and 100")
    manifest_path = args.manifest.resolve()
    manifest = load_manifest(manifest_path)
    fixture = select_fixture(manifest, args.profile)
    fixture_path, expected = validate_fixture(
        manifest_path,
        fixture,
        report_type=args.report_type,
    )
    template_path = resolve_template(args.report_type, args.template)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
    output_path = (
        args.output.resolve()
        if args.output
        else DEFAULT_ARTIFACT_ROOT / timestamp / "benchmark.json"
    )
    run_dir = output_path.parent
    run_dir.mkdir(parents=True, exist_ok=True)
    feature_flags = {
        name: os.getenv(
            name,
            "1"
            if name
            in {
                "AUTO_REPORT_PREPARED_TEMPLATE",
                "AUTO_REPORT_FAST_CELL",
                "AUTO_REPORT_UNIFIED_SCHEDULER",
            }
            else "0",
        )
        for name in FEATURE_FLAG_NAMES
    }

    fixture_snapshot = {
        "schemaVersion": manifest["schemaVersion"],
        "fixtureSetId": manifest.get("fixtureSetId"),
        "generator": manifest.get("generator"),
        "fixture": fixture,
    }
    atomic_write_json(run_dir / "fixture-manifest.json", fixture_snapshot)

    trials: list[dict[str, Any]] = []
    for trial_number in range(1, args.trials + 1):
        run_id = uuid.uuid4().hex
        trial_output = run_dir / f"trial-{trial_number:02d}.json"
        retained_docx = (
            run_dir / f"trial-{trial_number:02d}.docx"
            if args.keep_docx
            else run_dir / f".trial-{trial_number:02d}.docx"
        )
        worker_config = {
            "runId": run_id,
            "trial": trial_number,
            "fixtureId": fixture["id"],
            "fixtureProfile": fixture["profile"],
            "fixturePath": str(fixture_path),
            "fixtureServerCount": fixture["serverCount"],
            "fixtureClientCount": fixture["clientCount"],
            "templatePath": str(template_path),
            "reportType": args.report_type,
            "expected": expected,
            "outputDocx": str(retained_docx),
            "retainDocx": bool(args.keep_docx),
            "updateFields": bool(args.update_fields),
            "fieldUpdateTimeoutSeconds": args.field_update_timeout,
            "featureFlags": feature_flags,
        }
        config_path = run_dir / f".trial-{trial_number:02d}.config.json"
        atomic_write_json(config_path, worker_config)
        environment = dict(os.environ)
        environment["PYTHONHASHSEED"] = "0"
        completed = subprocess.run(
            [
                sys.executable,
                str(Path(__file__).resolve()),
                "--worker-config",
                str(config_path),
                "--worker-output",
                str(trial_output),
            ],
            cwd=ROOT,
            env=environment,
            capture_output=True,
            text=True,
            check=False,
        )
        config_path.unlink(missing_ok=True)
        if not args.keep_docx:
            retained_docx.unlink(missing_ok=True)
        if trial_output.is_file():
            result = json.loads(trial_output.read_text(encoding="utf-8"))
        else:
            result = {
                "schemaVersion": 1,
                "status": "failed",
                "error": {
                    "code": "worker_result_missing",
                    "type": "WorkerProcessError",
                },
            }
            atomic_write_json(trial_output, result)
        result["workerExitCode"] = completed.returncode
        atomic_write_json(trial_output, result)
        trials.append(result)
        print(
            f"Trial {trial_number}/{args.trials}: {result['status']}",
            flush=True,
        )

    observed_cache_states = sorted(
        {
            str(trial.get("metrics", {}).get("metadata", {}).get("cacheState"))
            for trial in trials
            if trial.get("metrics", {}).get("metadata", {}).get("cacheState")
        }
    )
    report = {
        "schemaVersion": 1,
        "createdAt": datetime.now(timezone.utc).isoformat(),
        "benchmark": {
            "fixtureId": fixture["id"],
            "fixtureProfile": fixture["profile"],
            "fixtureHash": fixture["sha256"],
            "reportType": args.report_type,
            "templateHash": sha256_file(template_path),
            "trialIsolation": "fresh-process",
            "cacheState": (
                observed_cache_states[0] if len(observed_cache_states) == 1 else "mixed"
            ),
            "cacheStates": observed_cache_states,
            "osFilesystemCacheControlled": False,
            "wordFieldUpdater": "enabled" if args.update_fields else "disabled",
            "featureFlags": feature_flags,
        },
        "environment": collect_environment(),
        "summary": aggregate_trials(trials),
        "trials": trials,
        "passed": all(trial.get("status") == "passed" for trial in trials),
    }
    atomic_write_json(output_path, report)
    return report, output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Benchmark Reporter Pro DOCX generation")
    parser.add_argument("--manifest", type=Path, default=DEFAULT_MANIFEST)
    parser.add_argument("--profile", default="mixed")
    parser.add_argument("--report-type", choices=REPORT_TYPES, default="full")
    parser.add_argument("--template", type=Path)
    parser.add_argument("--trials", type=int, default=3)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--keep-docx", action="store_true")
    parser.add_argument("--update-fields", action="store_true")
    parser.add_argument("--field-update-timeout", type=int, default=180)
    parser.add_argument("--worker-config", type=Path, help=argparse.SUPPRESS)
    parser.add_argument("--worker-output", type=Path, help=argparse.SUPPRESS)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if args.worker_config or args.worker_output:
        if not args.worker_config or not args.worker_output:
            print("Both worker arguments are required", file=sys.stderr)
            return 2
        return _run_worker(args.worker_config.resolve(), args.worker_output.resolve())
    try:
        report, output_path = run_benchmark(args)
    except BenchmarkConfigurationError as exc:
        print(f"[ERROR] {exc}", file=sys.stderr)
        return 2
    print(f"Benchmark report: {output_path}")
    return 0 if report["passed"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
