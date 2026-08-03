"""Measure Phase 5 Preview promotion from a clean, isolated 50-asset workspace."""

from __future__ import annotations

import base64
import hashlib
import json
import os
import sys
import tempfile
import time
from pathlib import Path
from unittest.mock import patch


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "apps" / "backend"
sys.path.insert(0, str(BACKEND))
os.environ["AUTO_REPORT_PREVIEW_JOBS"] = "1"
os.environ["AUTO_REPORT_PREVIEW_CACHE"] = "1"

from fastapi import FastAPI  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402
from api import routes  # noqa: E402
from docx import Document  # noqa: E402
from api.models import PreviewDocxRequest  # noqa: E402
from core.database import Database  # noqa: E402
from core.docx_field_updater import FieldUpdateResult  # noqa: E402
from core.preview_artifacts import PreviewArtifactRegistry  # noqa: E402
from core.report_jobs import ReportJobManager  # noqa: E402


def wait_for(client: TestClient, url: str, terminal: set[str], timeout: float) -> dict:
    deadline = time.perf_counter() + timeout
    while time.perf_counter() < deadline:
        response = client.get(url)
        response.raise_for_status()
        payload = response.json()
        state = payload.get("job", payload)
        if state.get("status") in terminal:
            return state
        time.sleep(0.05)
    raise TimeoutError(f"Timed out waiting for {url}")


def main() -> int:
    sample = BACKEND / "samples" / "Tracking_2.csv"
    with tempfile.TemporaryDirectory(prefix="reporter-phase5-") as directory:
        workspace = Path(directory)
        database = Database(workspace / "data" / "reporter.db")
        database.initialize()
        jobs = ReportJobManager(max_workers=1, max_pending=2)
        artifacts = PreviewArtifactRegistry(workspace / "cache" / "previews")
        app = FastAPI()
        app.include_router(routes.router)
        patches = (
            patch.object(routes, "get_db", return_value=database),
            patch.object(routes, "_report_jobs", jobs),
            patch.object(routes, "_preview_artifacts", artifacts),
            patch.object(routes, "GENERATED_REPORTS_DIR", workspace / "generated"),
            patch.object(routes, "preview_jobs_enabled", return_value=True),
            patch.object(routes, "preview_cache_enabled", return_value=True),
            patch.object(
                routes,
                "refresh_docx_fields",
                return_value=FieldUpdateResult(False, "benchmark-deferred", "controlled"),
            ),
        )
        try:
            for active in patches:
                active.start()
            with TestClient(app) as client:
                encoded = base64.b64encode(sample.read_bytes()).decode("ascii")
                column_preview = client.post("/api/column-preview", json={
                    "filename": sample.name,
                    "contentBase64": encoded,
                })
                column_preview.raise_for_status()
                imported = client.post("/api/import-file", json={
                    "filename": sample.name,
                    "contentBase64": encoded,
                    "defaultType": "client",
                    "columnMapping": column_preview.json()["suggestedMapping"],
                })
                imported.raise_for_status()
                rows = imported.json()["rows"]
                if len(rows) != 50:
                    raise AssertionError(f"Expected 50 imported rows, received {len(rows)}")
                request = {
                    "rows": rows,
                    "reportType": "full",
                    "disablePlugins": True,
                    "clientRequestId": "benchmark-preview-50",
                }

                # Promotion is benchmarked independently from cold Preview generation.
                # The prepared artifact is a valid DOCX containing every imported asset,
                # pinned to the same accepted request/template identity as the API call.
                accepted, _plugins = routes._accept_report_snapshot(PreviewDocxRequest(**request))
                prepared_path = workspace / "prepared-preview.docx"
                prepared = Document()
                prepared.add_heading("Reporter Pro Phase 5 promotion fixture", level=1)
                table = prepared.add_table(rows=1, cols=4)
                for index, heading in enumerate(("Type", "Hostname", "IP", "Result")):
                    table.rows[0].cells[index].text = heading
                for row in rows:
                    cells = table.add_row().cells
                    cells[0].text = str(row.get("type", ""))
                    cells[1].text = str(row.get("hostname", ""))
                    cells[2].text = str(row.get("ip", ""))
                    cells[3].text = str(row.get("result", ""))
                prepared.save(prepared_path)
                managed = artifacts.register_ready(
                    prepared_path,
                    job_id="benchmark-prepared",
                    request_signature=accepted.request_signature,
                    content_signature=accepted.request_signature,
                    template_hash=accepted.template_hash,
                    cache_mode="deterministic",
                    snapshot=accepted,
                )
                preview_id = managed.id
                preview_bytes = managed.path.read_bytes()

                promotion_started = time.perf_counter()
                submitted = client.post("/api/report-jobs", json={
                    **request,
                    "previewId": preview_id,
                    "clientRequestId": "benchmark-report-50",
                    "outputName": "phase5-promotion-50.docx",
                })
                submitted.raise_for_status()
                job_id = submitted.json()["job"]["id"]
                report = wait_for(
                    client,
                    f"/api/report-jobs/{job_id}",
                    {"completed", "failed", "cancelled"},
                    30,
                )
                promotion_ms = (time.perf_counter() - promotion_started) * 1000
                if report["status"] != "completed":
                    raise AssertionError(f"Promotion ended as {report['status']}")
                report_bytes = client.get(f"/api/report-jobs/{job_id}/download").content
                identical = report_bytes == preview_bytes
                result = {
                    "fixture": "Tracking_2.csv",
                    "assets": len(rows),
                    "promotionMs": round(promotion_ms, 3),
                    "targetPromotionMs": 2000,
                    "wordFieldUpdate": "deferred-controlled",
                    "artifactPreparation": "isolated-valid-docx-with-50-assets",
                    "bytes": len(report_bytes),
                    "sha256": hashlib.sha256(report_bytes).hexdigest(),
                    "byteIdentical": identical,
                    "passed": identical and promotion_ms < 2000,
                }
                print(json.dumps(result, ensure_ascii=False, indent=2))
                return 0 if result["passed"] else 1
        finally:
            jobs.shutdown()
            artifacts.shutdown()
            database.close()
            for active in reversed(patches):
                active.stop()


if __name__ == "__main__":
    raise SystemExit(main())
