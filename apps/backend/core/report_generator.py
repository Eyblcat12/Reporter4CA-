from __future__ import annotations

from contextlib import nullcontext
from contextvars import ContextVar
from copy import deepcopy
from datetime import date, datetime
from enum import Enum
from functools import lru_cache
import gc
from io import BytesIO
import logging
from pathlib import Path
import time
from typing import Any

from core.config import (
    compact_prototype_enabled,
    fast_cell_enabled,
    prepared_template_cache_bytes,
    prepared_template_cache_entries,
    prepared_template_enabled,
)
from core.prepared_template import (
    PreparedTemplateCache,
    PreparedTemplateError,
)
from core.template_blueprint import (
    BLUEPRINT_SCHEMA_VERSION,
    TableBlueprint,
    UnsafeTableBlueprintError,
    compile_table_blueprint,
)
from core.threat_intelligence import normalize_iocs, normalize_mitre
from core.rule_engine import assessment_text, evaluate_payload
from core.report_integrity import build_report_manifest, verify_report_document
from core.performance_metrics import PerformanceMetrics


DEFAULT_TEXT_VALUE = "N/A"
DEFAULT_TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "templates" / "report_template.docx"
DEFAULT_PREPARED_TEMPLATE_CACHE_ROOT = (
    Path(__file__).resolve().parent.parent / "data" / "cache" / "prepared-templates"
)
REPORTER_HEADING_NUMBERING_MARKER = "52505452"
REPORTER_HEADING_TEXT_LEFT_TWIPS = "720"
DEFAULT_RESULT_TEXT = "Không ghi nhận dấu hiệu bất thường"
ANOMALY_RESULT_TEXT = "Phát hiện dấu hiệu bất thường"
_CURRENT_PERFORMANCE_METRICS: ContextVar[PerformanceMetrics | None] = ContextVar(
    "reporter_performance_metrics",
    default=None,
)
_CURRENT_BUILD_TRACKER: ContextVar[Any | None] = ContextVar(
    "reporter_build_tracker",
    default=None,
)
_LOGGER = logging.getLogger(__name__)


class _BuildWorkTracker:
    """Bounded cooperative checkpoints for long, in-memory DOCX builds."""

    def __init__(
        self,
        *,
        estimated_rows: int,
        check_cancelled: Any | None,
        on_progress: Any | None,
        metrics: PerformanceMetrics | None,
        batch_size: int = 32,
    ) -> None:
        self.estimated_rows = max(1, int(estimated_rows))
        self.check_cancelled = check_cancelled
        self.on_progress = on_progress
        self.metrics = metrics
        self.batch_size = min(max(int(batch_size), 25), 50)
        self.completed_rows = 0
        self._last_checkpoint_ns = time.perf_counter_ns()

    def start(self) -> None:
        self._checkpoint(force=True)

    def advance(self, rows: int = 1) -> None:
        self.completed_rows += max(0, int(rows))
        if self.completed_rows % self.batch_size == 0:
            self._checkpoint()

    def finish(self) -> None:
        self._checkpoint(force=True, completed=True)

    def _checkpoint(self, *, force: bool = False, completed: bool = False) -> None:
        if not force and self.completed_rows <= 0:
            return
        started_ns = self._last_checkpoint_ns
        if self.check_cancelled is not None:
            self.check_cancelled()
        if self.on_progress is not None:
            ratio = min(1.0, self.completed_rows / self.estimated_rows)
            value = 66 if completed else min(65, 45 + int(ratio * 20))
            self.on_progress(value, self.completed_rows)
        finished_ns = time.perf_counter_ns()
        if self.metrics is not None and self.completed_rows:
            self.metrics.record_aggregate(
                "buildCheckpoint",
                "rows",
                max(0, finished_ns - started_ns) / 1_000_000,
            )
        self._last_checkpoint_ns = finished_ns

# Từ khóa nhận diện máy có bất thường (cụm từ, không phải từ đơn lẻ)
_ANOMALY_KEYWORDS = {
    "phát hiện mã độc", "mã độc", "bất thường", "malware", "trojan",
    "virus", "backdoor", "webshell", "c2", "beacon", "cobalt",
    "emotet", "ransomware", "exploit", "suspicious", "infected",
    "compromised", "lateral", "persistence", "dropper",
    "keylogger", "reverse shell", "plugx", "shadowpad", "mimikatz",
}

# Từ khóa phủ định — nếu chứa keyword này → KHÔNG phải bất thường
_CLEAN_KEYWORDS = {
    "không phát hiện", "không ghi nhận", "clean", "not found",
    "no malware", "no threat", "bình thường",
}

METHOD_CHECKLIST_ITEMS = [
    ("Tiến hành thu thập và phân tích các dữ liệu, các log liên quan đến quá trình điều tra. Trường hợp log source đã được thu thập trên SIEM, thực hiện đánh giá trên SIEM.", 0),
    ("Kiểm tra rootkit.", 0),
    ("Xác định các Autorun Entry bất thường.", 0),
    ("Xác định các Service, process bất thường.", 0),
    ("Xác định các lập lịch bất thường.", 0),
    ("Xác định, phân tích các tệp tin bất thường.", 0),
    ("Kiểm tra các kết nối mạng bất thường.", 0),
    ("Xác định các tài khoản, nhóm quyền bất thường.", 0),
    ("Kiểm tra các file/thư mục chia sẻ bất thường.", 0),
    ("Kiểm tra cấu hình liên quan đến tunnel.", 0),
    ("Kiểm tra named pipe.", 0),
    ("Với webserver:", 0),
    ("Kiểm tra webshell.", 1),
    ("Kiểm tra các access log, web application log nhằm xác định các request bất thường.", 1),
    ("Kiểm tra phiên bản phần mềm và các gói cài đặt trên server.", 0),
    ("Trường hợp cần thiết, thực hiện trích xuất memory và điều tra, phân tích trên memory này.", 0),
]

TOOLS_USED = [
    "LiveResponseCollection-Cedarpelta.",
    "Windows SysInternal Suite.",
    "Kaspersky TDSSKiller.",
    "Kaspersky Virus Removal Tool.",
    "Process Hacker.",
]

ASSESSMENT_APPROACHES = [
    "Thực hiện Remote Access tới các thiết bị bằng tài khoản đặc quyền cao nhất (Administrator/root). Nhân sự thực hiện sử dụng các công cụ được liệt kê ở trên để phân tích, đánh giá trực tiếp trên thiết bị.",
    "Thực hiện phân tích các log source (log hệ thống, log kết nối, log ứng dụng, v.v.) thông qua hệ thống SIEM.",
    "Trường hợp sử dụng EDR và nhân sự đánh giá được phép truy cập hệ thống quản lý EDR, thực hiện phân tích các alert và phân tích realtime.",
]

DETAIL_CHECKLIST_ITEMS = [
    "Sử dụng công cụ thực hiện rà soát và đánh giá",
    "Tiến hành thu thập và phân tích các dữ liệu, các log liên quan đến quá trình điều tra",
    "Kiểm tra rootkit",
    "Xác định các Autorun Entry bất thường",
    "Xác định các Service, process, loaded DLL bất thường",
    "Xác định các schedule task bất thường",
    "Xác định, phân tích các tệp tin bất thường",
    "Kiểm tra các kết nối mạng bất thường",
    "Xác định các tài khoản, nhóm quyền bất thường",
    "Kiểm tra các file/thư mục chia sẻ bất thường",
    "Kiểm tra cấu hình liên quan đến tunnel",
    "Kiểm tra named pipe",
    "Kiểm tra Prefetch files",
    "Kiểm tra webshell",
    "Kiểm tra các access log, web application log nhằm xác định các request bất thường",
]

DEFAULT_RECOMMENDATIONS = [
    "Đã hoàn thành thực hiện gỡ bỏ mã độc trên các máy chủ.",
    "Thực hiện thêm IoCs trên nền tảng giám sát để rà quét diện rộng.",
    "Kiểm tra các khóa truy cập, tài khoản đặc quyền và các artifact nhạy cảm đang được sử dụng hợp lệ hay không. Thực hiện backup hoặc mã hóa tệp khi cần thiết.",
    "Thường xuyên rà soát hệ thống định kỳ 6 tháng/lần để đảm bảo an toàn.",
    "Không sử dụng phần mềm trái phép, crack hay remote quản trị từ xa nếu không được cho phép.",
    "Thu hồi quyền quản trị cao nhất trên từng máy chủ của các cá nhân người dùng trong hệ thống.",
]


# ---------------------------------------------------------------------------
# Report type enum
# ---------------------------------------------------------------------------

class ReportType(str, Enum):
    FULL = "full"
    SERVER_ONLY = "server_only"
    CLIENT_ONLY = "client_only"
    SUMMARY = "summary"
    TECHNICAL = "technical"
    INCIDENT_RESPONSE = "incident_response"


# ---------------------------------------------------------------------------
# ReportBuilder class — điểm trung tâm tạo báo cáo
# ---------------------------------------------------------------------------

class ReportBuilder:
    """Wraps report generation with configurable report_type support."""

    def __init__(
        self,
        data: dict[str, Any],
        *,
        title: str,
        organization: str,
        assessment_date: str | None = None,
        template_path: str | Path | None = None,
        report_type: ReportType = ReportType.FULL,
        metrics: PerformanceMetrics | None = None,
        check_cancelled: Any | None = None,
        on_build_progress: Any | None = None,
    ) -> None:
        self.data = data
        self.title = title
        self.organization = organization
        self.assessment_date = assessment_date
        self.template_path = template_path
        self.report_type = report_type
        self.metrics = metrics
        self.check_cancelled = check_cancelled
        self.on_build_progress = on_build_progress

    def build(self) -> Any:
        """Tạo document python-docx theo report_type."""
        document, using_template, _template_mode, _resolved_template = _create_base_document(
            self.template_path,
            report_type=self.report_type,
            metrics=self.metrics,
        )
        with _performance_phase(self.metrics, "documentConfigure"):
            _configure_document(document, using_template=using_template)

        with _performance_phase(self.metrics, "tokenOrCover"):
            if using_template:
                _replace_template_tokens(
                    document,
                    title=self.title,
                    organization=self.organization,
                    assessment_date=self.assessment_date,
                )
            else:
                _add_fallback_cover_page(
                    document,
                    title=self.title,
                    organization=self.organization,
                    assessment_date=self.assessment_date,
                )

        # Lọc dữ liệu theo report_type
        filtered_data = self._filter_data()
        with _performance_phase(self.metrics, "ruleEvaluation"):
            data = evaluate_payload(filtered_data)

        total_assets = len(data.get("servers", [])) + len(data.get("clients", []))
        tracker = _BuildWorkTracker(
            estimated_rows=max(64, total_assets * (len(DETAIL_CHECKLIST_ITEMS) + 4)),
            check_cancelled=self.check_cancelled,
            on_progress=self.on_build_progress,
            metrics=self.metrics,
        )
        metrics_token = _CURRENT_PERFORMANCE_METRICS.set(self.metrics)
        tracker_token = _CURRENT_BUILD_TRACKER.set(tracker)
        try:
            tracker.start()
            with _performance_phase(self.metrics, "reportBodyBuild"):
                if self.report_type == ReportType.INCIDENT_RESPONSE:
                    self._build_incident_response_report(document, data)
                elif self.report_type == ReportType.SUMMARY:
                    self._build_summary_report(document, data)
                elif self.report_type == ReportType.TECHNICAL:
                    self._build_technical_report(document, data)
                else:
                    self._build_standard_report(document, data)
            tracker.finish()
        finally:
            _CURRENT_BUILD_TRACKER.reset(tracker_token)
            _CURRENT_PERFORMANCE_METRICS.reset(metrics_token)

        with _performance_phase(self.metrics, "manifestBuild"):
            manifest = build_report_manifest(data, self.report_type.value)
        with _performance_phase(self.metrics, "integrityVerify"):
            verification = verify_report_document(document, manifest)
        document._reporter_manifest = manifest
        document._reporter_integrity = verification
        return document

    def _filter_data(self) -> dict[str, Any]:
        """Lọc dữ liệu theo report_type."""
        data = dict(self.data)
        if self.report_type == ReportType.SERVER_ONLY:
            data["clients"] = []
        elif self.report_type == ReportType.CLIENT_ONLY:
            data["servers"] = []
        return data

    def _build_standard_report(self, document: Any, data: dict[str, Any]) -> None:
        """Báo cáo đầy đủ hoặc chỉ server/client."""
        include_servers, include_clients = self._scope()
        _add_overview_section(
            document,
            data,
            organization=self.organization,
            assessment_date=self.assessment_date,
            include_servers=include_servers,
            include_clients=include_clients,
        )
        _add_results_section(
            document, data,
            include_servers=include_servers,
            include_clients=include_clients,
        )
        _add_investigation_section(
            document,
            include_servers=include_servers,
            include_clients=include_clients,
        )
        _add_remediation_section(
            document, data,
            include_servers=include_servers,
            include_clients=include_clients,
        )
        _add_ioc_section(document)
        _add_recommendations_section(
            document,
            include_servers=include_servers,
            include_clients=include_clients,
        )

    def _build_summary_report(self, document: Any, data: dict[str, Any]) -> None:
        """Build the concise three-part management summary."""
        _add_summary_overview_section(
            document,
            data,
            organization=self.organization,
            assessment_date=self.assessment_date,
        )
        _add_summary_analysis_section(document, data)
        _add_summary_conclusion_section(document, data)

    def _build_technical_report(self, document: Any, data: dict[str, Any]) -> None:
        """Build the evidence-oriented three-part technical report."""
        _add_technical_overview_section(
            document,
            data,
            organization=self.organization,
            assessment_date=self.assessment_date,
        )
        _add_technical_analysis_section(document, data)
        _add_technical_conclusion_section(document, data)

    def _scope(self) -> tuple[bool, bool]:
        if self.report_type == ReportType.SERVER_ONLY:
            return True, False
        if self.report_type == ReportType.CLIENT_ONLY:
            return False, True
        return True, True

    def _build_incident_response_report(self, document: Any, data: dict[str, Any]) -> None:
        """Build a dedicated incident-response report from imported assets and metadata."""
        metadata = data.get("metadata", {}) if isinstance(data.get("metadata"), dict) else {}
        assets = [*data.get("servers", []), *data.get("clients", [])]

        _add_heading(document, "Thông tin sự cố", level=1)
        incident_rows = [
            ["Mã sự cố", str(metadata.get("incident_id", metadata.get("incidentId", "N/A")))],
            ["Mức độ", str(metadata.get("severity", "Chưa xác định"))],
            ["Trạng thái", str(metadata.get("status", "Đang xử lý"))],
            ["Thời điểm phát hiện", str(metadata.get("detected_at", metadata.get("detectedAt", "N/A")))],
        ]
        table = _create_table(document, ["Thuộc tính", "Giá trị"], incident_rows, column_widths_mm=[52, 108])
        _style_table(table, left_align_columns={0, 1})

        _add_heading(document, "Tóm tắt điều hành", level=1)
        supplied_summary = metadata.get("executive_summary", metadata.get("executiveSummary", ""))
        verified_findings = sum(len(asset.get("findings", [])) for asset in assets)
        _add_body_paragraph(document, str(supplied_summary or (
            f"Phạm vi ghi nhận {len(assets)} tài sản và {verified_findings} finding có evidence từ dữ liệu đầu vào. "
            "Các kết luận chưa có bằng chứng liên kết không được tự động đưa vào tóm tắt."
        )))

        _add_heading(document, "Tài sản bị ảnh hưởng", level=1)
        asset_rows = [
            [str(index), str(asset.get("hostname", "N/A")), str(asset.get("ip", "N/A")), str(asset.get("result", "N/A"))]
            for index, asset in enumerate(assets, start=1)
        ] or [["", "", "", ""]]
        table = _create_table(
            document,
            ["STT", "Tài sản", "Địa chỉ IP", "Kết quả"],
            asset_rows,
            column_widths_mm=[12, 48, 38, 62],
        )
        _style_table(table, left_align_columns={1, 2, 3})

        _add_heading(document, "Dòng thời gian", level=1)
        timeline = metadata.get("timeline", [])
        timeline_rows = []
        if isinstance(timeline, list):
            for item in timeline:
                if isinstance(item, dict):
                    timeline_rows.append([
                        str(item.get("time", item.get("timestamp", "N/A"))),
                        str(item.get("event", item.get("description", ""))),
                        str(item.get("evidence", item.get("evidenceId", ""))),
                        str(item.get("relatedIocs", item.get("iocs", ""))),
                    ])
        table = _create_table(
            document, ["Thời gian", "Sự kiện", "Bằng chứng", "IoC liên quan"],
            timeline_rows or [["", "", "", ""]], column_widths_mm=[30, 60, 35, 35],
        )
        _style_table(table, left_align_columns={0, 1, 2, 3})

        _add_heading(document, "Phát hiện và bằng chứng", level=1)
        finding_rows = []
        for asset in assets:
            for finding in asset.get("findings", []):
                evidence = "; ".join(
                    f"{item.get('field')}: {item.get('value')}" for item in finding.get("evidence", [])
                )
                finding_rows.append([
                    str(asset.get("hostname", "N/A")),
                    str(finding.get("ruleId", "N/A")),
                    str(finding.get("severity", "N/A")),
                    evidence,
                ])
        table = _create_table(
            document,
            ["Tài sản", "Rule", "Mức độ", "Bằng chứng"],
            finding_rows or [["", "Không ghi nhận phát hiện", "", ""]],
            column_widths_mm=[35, 42, 25, 58],
        )
        _style_table(table, left_align_columns={0, 1, 2, 3})

        _add_heading(document, "Indicators of compromise (IoCs)", level=1)
        iocs = normalize_iocs(metadata.get("iocs", []), default_source="incident metadata")
        ioc_rows = [
            [str(index), item["type"], item["value"], "Valid" if item["valid"] else "Invalid", ", ".join(item["sources"])]
            for index, item in enumerate(iocs, start=1)
        ]
        table = _create_table(
            document, ["STT", "Loại", "Giá trị", "Kiểm tra", "Nguồn"],
            ioc_rows or [["", "", "", "", ""]],
        )
        _style_table(table, left_align_columns={1, 2, 3, 4})

        _add_heading(document, "MITRE ATT&CK", level=1)
        mappings = normalize_mitre(metadata.get("mitre", metadata.get("mitreMappings", [])))
        mitre_rows = [
            [item["technique"], item["tactic"], item["name"], item["evidence"], "Valid" if item["valid"] else "Needs evidence"]
            for item in mappings
        ]
        table = _create_table(
            document, ["Technique", "Tactic", "Tên", "Bằng chứng", "Trạng thái"],
            mitre_rows or [["", "", "", "", "Chưa có mapping"]],
        )
        _style_table(table, left_align_columns={0, 1, 2, 3, 4})

        _add_heading(document, "Ứng phó sự cố", level=1)
        for title, key, fallback in (
            ("Khoanh vùng", "containment_actions", "Chưa ghi nhận hành động khoanh vùng."),
            ("Loại bỏ", "eradication_actions", "Chưa ghi nhận hành động loại bỏ."),
            ("Khôi phục", "recovery_actions", "Chưa ghi nhận hành động khôi phục."),
        ):
            _add_heading(document, title, level=2)
            value = metadata.get(key, metadata.get("".join([key.split("_")[0], "Actions"]), []))
            action_rows = []
            values = value if isinstance(value, list) else [value] if value else []
            for action in values:
                if isinstance(action, dict):
                    action_rows.append([
                        str(action.get("action", action.get("description", ""))),
                        str(action.get("status", "Planned")),
                        str(action.get("owner", "Unassigned")),
                        str(action.get("evidence", "")),
                    ])
                elif str(action).strip():
                    action_rows.append([str(action), "Planned", "Unassigned", ""])
            table = _create_table(
                document, ["Hành động", "Trạng thái", "Phụ trách", "Bằng chứng"],
                action_rows or [[fallback, "Not started", "Unassigned", ""]],
            )
            _style_table(table, left_align_columns={0, 1, 2, 3})

        _add_heading(document, "Bài học kinh nghiệm", level=1)
        _add_body_paragraph(
            document,
            _metadata_text(metadata.get("lessons_learned", metadata.get("lessonsLearned")), "Chưa có nội dung."),
        )
        _add_recommendations_section(document)


# ---------------------------------------------------------------------------
# Backward-compatible top-level function
# ---------------------------------------------------------------------------

def generate_report(
    data: dict[str, Any],
    *,
    title: str,
    organization: str,
    assessment_date: str | None = None,
    template_path: str | Path | None = None,
    report_type: str | ReportType = ReportType.FULL,
    metrics: PerformanceMetrics | None = None,
    check_cancelled: Any | None = None,
    on_build_progress: Any | None = None,
) -> Any:
    if isinstance(report_type, str):
        try:
            report_type = ReportType(report_type)
        except ValueError:
            report_type = ReportType.FULL

    builder = ReportBuilder(
        data,
        title=title,
        organization=organization,
        assessment_date=assessment_date,
        template_path=template_path,
        report_type=report_type,
        metrics=metrics,
        check_cancelled=check_cancelled,
        on_build_progress=on_build_progress,
    )
    return builder.build()


def save_report(document: Any, output_path: str | Path) -> Path:
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_file)
    return output_file.resolve()


def render_preview_text(data: dict[str, Any]) -> str:
    servers = data.get("servers", [])
    clients = data.get("clients", [])
    total_assets = len(servers) + len(clients)

    lines = [
        "Preview du lieu bao cao",
        "======================",
        f"Servers : {len(servers)}",
        f"Clients : {len(clients)}",
        f"Tong so : {total_assets}",
        f"Bat thuong : {sum(1 for asset in [*servers, *clients] if asset.get('assessment', {}).get('classification') == 'anomaly')}",
        f"Can xac minh : {sum(1 for asset in [*servers, *clients] if asset.get('assessment', {}).get('classification') == 'needs_review')}",
        f"Chua du du lieu : {sum(1 for asset in [*servers, *clients] if asset.get('assessment', {}).get('classification') == 'insufficient_data')}",
        "",
        "May chu dau tien:",
    ]

    if servers:
        first_server = servers[0]
        lines.append(
            f"- {first_server['hostname']} | IP: {first_server.get('ip', 'N/A')} | OS: {first_server.get('os', 'N/A')}"
        )
    else:
        lines.append("- Khong co du lieu server")

    lines.append("")
    lines.append("May tram dau tien:")

    if clients:
        first_client = clients[0]
        lines.append(
            f"- {first_client['hostname']} | IP: {first_client.get('ip', 'N/A')} | OS: {first_client.get('os', 'N/A')}"
        )
    else:
        lines.append("- Khong co du lieu client")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Document creation helpers
# ---------------------------------------------------------------------------

def _create_base_document(
    template_path: str | Path | None,
    *,
    report_type: str | ReportType = ReportType.FULL,
    metrics: PerformanceMetrics | None = None,
) -> tuple[Any, bool, str, Path | None]:
    modules = _docx_modules()
    Document = modules["Document"]

    with _performance_phase(metrics, "templateResolve"):
        resolved_template = _resolve_template_path(template_path)
    if resolved_template is not None:
        normalized_report_type = (
            report_type.value if isinstance(report_type, ReportType) else str(report_type)
        )
        document = None
        template_mode = "none"
        if prepared_template_enabled():
            prepared_attributes: dict[str, Any] = {
                "cacheHit": False,
                "outcome": "failed",
            }
            try:
                with _performance_phase(
                    metrics,
                    "preparedTemplate",
                    attributes=prepared_attributes,
                ):
                    prepared = _get_prepared_template_cache().get_or_compile(
                        resolved_template,
                        normalized_report_type,
                        _compile_prepared_template,
                        blueprint_schema_version=BLUEPRINT_SCHEMA_VERSION,
                    )
                    prepared_attributes["cacheHit"] = prepared.cache_hit
                    prepared_attributes["outcome"] = "passed"
                metadata = prepared.manifest.get("metadata", {})
                if isinstance(metadata, dict):
                    template_mode = str(metadata.get("templateMode", "none"))
                captured_prototypes: dict[str, Any] = {}
                if template_mode == "full":
                    with _performance_phase(metrics, "prototypeCapture"):
                        prototype_source = Document(BytesIO(prepared.source_bytes))
                        _capture_template_prototypes(prototype_source)
                        captured_prototypes = _extract_template_prototypes(
                            prototype_source
                        )
                        del prototype_source
                        # Avoid overlapping the large source package with the
                        # prepared document in peak working-set measurements.
                        gc.collect()
                with _performance_phase(metrics, "templateLoad"):
                    document = Document(str(prepared.path))
                _apply_template_prototypes(captured_prototypes, document)
                if metrics is not None:
                    metrics.update_metadata({
                        "cacheState": (
                            "cache-warm/prepared-hit"
                            if prepared.cache_hit
                            else "process-cold/cache-miss"
                        )
                    })
            except Exception as exc:
                # Cache failures are never allowed to break a report that the
                # legacy engine can still generate. Do not log paths or content.
                _LOGGER.warning(
                    "Prepared template unavailable (%s); using legacy path.",
                    type(exc).__name__,
                )
                if metrics is not None:
                    metrics.update_metadata({
                        "cacheState": "prepared-fallback",
                        "preparedFallback": _prepared_fallback_reason(exc),
                    })
                document = None

        if document is None:
            with _performance_phase(metrics, "templateLoad"):
                document = Document(str(resolved_template))
            with _performance_phase(metrics, "templateDetect"):
                template_mode = _detect_template_mode(document)
            if template_mode == "full":
                with _performance_phase(metrics, "prototypeCapture"):
                    _capture_template_prototypes(document)
            with _performance_phase(metrics, "tocCleanup"):
                _clear_cached_toc_result(document)
            with _performance_phase(metrics, "templateTrim"):
                _trim_template_body(document)
        using_template = True
    else:
        with _performance_phase(metrics, "templateLoad"):
            document = Document()
        _apply_sample_page_layout(document)
        using_template = False
        template_mode = "none"

    _enable_field_updates(document)
    return document, using_template, template_mode, resolved_template


@lru_cache(maxsize=1)
def _get_prepared_template_cache() -> PreparedTemplateCache:
    return PreparedTemplateCache(
        DEFAULT_PREPARED_TEMPLATE_CACHE_ROOT,
        max_entries=prepared_template_cache_entries(),
        max_bytes=prepared_template_cache_bytes(),
    )


def _compile_prepared_template(source_bytes: bytes) -> tuple[bytes, dict[str, object]]:
    """Compile the same legacy cleanup/trim sequence into an immutable DOCX."""

    Document = _docx_modules()["Document"]
    document = Document(BytesIO(source_bytes))
    template_mode = _detect_template_mode(document)
    _clear_cached_toc_result(document)
    _trim_template_body(document)
    output = BytesIO()
    document.save(output)
    return output.getvalue(), {"templateMode": template_mode}


def _prepared_fallback_reason(exc: Exception) -> str:
    if isinstance(exc, PreparedTemplateError):
        return "cache_or_compile_error"
    if isinstance(exc, (OSError, PermissionError)):
        return "filesystem_error"
    return "unexpected_error"


def _extract_template_prototypes(source: Any) -> dict[str, Any]:
    """Detach process-local prototype/blueprint snapshots from a source DOCX."""

    captured: dict[str, Any] = {}
    for attribute in (
        "_codex_table_prototypes",
        "_reporter_compact_prototype_enabled",
        "_reporter_table_blueprints",
    ):
        if hasattr(source, attribute):
            captured[attribute] = getattr(source, attribute)
    return captured


def _apply_template_prototypes(captured: dict[str, Any], target: Any) -> None:
    for attribute, value in captured.items():
        setattr(target, attribute, value)


def warm_prepared_template(
    template_path: str | Path,
    report_type: str | ReportType,
) -> bool:
    """Prepare one compatible template without changing upload success semantics."""

    if not prepared_template_enabled():
        return False
    resolved = _resolve_template_path(template_path)
    if resolved is None:
        return False
    normalized = report_type.value if isinstance(report_type, ReportType) else str(report_type)
    _get_prepared_template_cache().get_or_compile(
        resolved,
        normalized,
        _compile_prepared_template,
        blueprint_schema_version=BLUEPRINT_SCHEMA_VERSION,
    )
    return True


def warm_prepared_template_safely(
    template_path: str | Path,
    report_type: str | ReportType,
) -> None:
    try:
        warm_prepared_template(template_path, report_type)
    except (OSError, PreparedTemplateError, RuntimeError, ValueError) as exc:
        _LOGGER.warning(
            "Prepared template warm-up deferred (%s).",
            type(exc).__name__,
        )


def bundled_template_sources(
    templates_root: str | Path,
) -> list[tuple[ReportType, Path]]:
    """Resolve the same bundled defaults used by the API without touching the DB."""

    root = Path(templates_root).expanduser().resolve()
    fallback = root / "report_template.docx"
    sources: list[tuple[ReportType, Path]] = []
    for report_type in ReportType:
        category = root / report_type.value
        candidates = sorted(
            path
            for path in category.glob("*.docx")
            if path.is_file() and not path.name.startswith("~")
        ) if category.is_dir() else []
        source = candidates[0] if candidates else fallback
        if source.is_file():
            sources.append((report_type, source))
    return sources


def warm_bundled_templates(
    templates_root: str | Path,
) -> list[dict[str, Any]]:
    """Prepare bundled defaults and return aggregate-only setup diagnostics."""

    results: list[dict[str, Any]] = []
    for report_type, source in bundled_template_sources(templates_root):
        started_ns = time.perf_counter_ns()
        outcome = "prepared"
        error_code = ""
        try:
            prepared = warm_prepared_template(source, report_type)
            if not prepared:
                outcome = "disabled"
        except (OSError, PreparedTemplateError, RuntimeError, ValueError) as exc:
            outcome = "deferred"
            error_code = type(exc).__name__
        results.append({
            "reportType": report_type.value,
            "outcome": outcome,
            "errorCode": error_code,
            "durationMs": round(
                (time.perf_counter_ns() - started_ns) / 1_000_000,
                3,
            ),
        })
    return results


def _performance_phase(
    metrics: PerformanceMetrics | None,
    name: str,
    *,
    attributes: dict[str, Any] | None = None,
) -> Any:
    """Return a no-op context unless performance collection was requested."""

    return (
        metrics.phase(name, attributes=attributes)
        if metrics is not None
        else nullcontext()
    )


@lru_cache(maxsize=1)
def _docx_modules() -> dict[str, Any]:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION_START
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt, RGBColor
        from docx.table import Table
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Can cai dat python-docx de tao bao cao. Hay chay: python -m pip install -r requirements.txt"
        ) from exc

    return {
        "Document": Document,
        "Mm": Mm,
        "OxmlElement": OxmlElement,
        "Pt": Pt,
        "RGBColor": RGBColor,
        "Table": Table,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "WD_CELL_VERTICAL_ALIGNMENT": WD_CELL_VERTICAL_ALIGNMENT,
        "WD_SECTION_START": WD_SECTION_START,
        "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
        "qn": qn,
    }


def _resolve_template_path(template_path: str | Path | None) -> Path | None:
    if template_path:
        resolved = Path(template_path).expanduser()
        if not resolved.is_absolute():
            resolved = (Path.cwd() / resolved).resolve()
        if not resolved.exists():
            raise FileNotFoundError(f"Khong tim thay template DOCX: {resolved}")
        return resolved

    if DEFAULT_TEMPLATE_PATH.exists():
        return DEFAULT_TEMPLATE_PATH
    return None


def _trim_template_body(document: Any) -> None:
    modules = _docx_modules()
    qn = modules["qn"]
    body = document._element.body
    children = list(body)

    first_heading_index: int | None = None
    last_section_index = len(children)
    if children and children[-1].tag == qn("w:sectPr"):
        last_section_index = len(children) - 1

    for index, child in enumerate(children[:last_section_index]):
        if child.tag != qn("w:p"):
            continue
        p_pr = child.find(qn("w:pPr"))
        if p_pr is None:
            continue
        p_style = p_pr.find(qn("w:pStyle"))
        style_id = p_style.get(qn("w:val")) if p_style is not None else ""
        text = "".join(child.itertext()).strip()
        if style_id == "Heading1" and text:
            first_heading_index = index
            break

    if first_heading_index is None:
        return

    for child in children[first_heading_index:last_section_index]:
        body.remove(child)


def _clear_cached_toc_result(document: Any) -> bool:
    """Remove a template's stale TOC display while preserving the TOC field.

    Word stores every visible TOC entry (including nested PAGEREF fields) in
    document.xml. Some retained templates contain hundreds of those cached
    entries. Keeping them makes browser previews stale and forces Word to
    process obsolete fields before it can rebuild the table.
    """

    modules = _docx_modules()
    OxmlElement = modules["OxmlElement"]
    qn = modules["qn"]
    body = document._element.body
    start_paragraph = None
    for child in body.iter(qn("w:p")):
        instructions = " ".join(
            (node.text or "") for node in child.iter(qn("w:instrText"))
        )
        if "TOC " in instructions.upper():
            start_paragraph = child
            break

    if start_paragraph is None:
        return False

    container = start_paragraph.getparent()
    container_children = list(container)
    start_index = container_children.index(start_paragraph)

    depth = 0
    outer_started = False
    separator_run = None
    end_index: int | None = None

    for index in range(start_index, len(container_children)):
        child = container_children[index]
        for field_char in child.iter(qn("w:fldChar")):
            field_type = field_char.get(qn("w:fldCharType"))
            if field_type == "begin":
                depth += 1
                outer_started = True
            elif field_type == "separate" and outer_started and depth == 1 and separator_run is None:
                separator_run = field_char.getparent()
            elif field_type == "end" and outer_started:
                if depth == 1 and separator_run is not None:
                    end_index = index
                    break
                depth = max(0, depth - 1)
        if end_index is not None:
            break

    if separator_run is None or end_index is None:
        return False

    paragraph_children = list(start_paragraph)
    try:
        separator_index = paragraph_children.index(separator_run)
    except ValueError:
        return False

    for child in paragraph_children[separator_index + 1:]:
        start_paragraph.remove(child)

    end_run = OxmlElement("w:r")
    end_field = OxmlElement("w:fldChar")
    end_field.set(qn("w:fldCharType"), "end")
    end_run.append(end_field)
    start_paragraph.append(end_run)

    for index in range(end_index, start_index, -1):
        container.remove(container_children[index])
    return True


def _detect_template_mode(document: Any) -> str:
    modules = _docx_modules()
    qn = modules["qn"]
    body = document._element.body
    children = list(body)

    first_heading_index: int | None = None
    last_section_index = len(children)
    if children and children[-1].tag == qn("w:sectPr"):
        last_section_index = len(children) - 1

    table_count = 0
    heading_count = 0
    caption_count = 0
    non_empty_paragraphs = 0

    for index, child in enumerate(children[:last_section_index]):
        if child.tag != qn("w:p"):
            continue

        p_pr = child.find(qn("w:pPr"))
        p_style = p_pr.find(qn("w:pStyle")) if p_pr is not None else None
        style_id = p_style.get(qn("w:val")) if p_style is not None else ""
        text = "".join(child.itertext()).strip()

        if first_heading_index is None and style_id == "Heading1" and text:
            first_heading_index = index
            break

    if first_heading_index is None:
        return "cover"

    for child in children[first_heading_index:last_section_index]:
        if child.tag == qn("w:tbl"):
            table_count += 1
            continue

        if child.tag != qn("w:p"):
            continue

        p_pr = child.find(qn("w:pPr"))
        p_style = p_pr.find(qn("w:pStyle")) if p_pr is not None else None
        style_id = p_style.get(qn("w:val")) if p_style is not None else ""
        text = "".join(child.itertext()).strip()

        if text:
            non_empty_paragraphs += 1
        if style_id in {"Heading1", "Heading2", "Heading3"} and text:
            heading_count += 1
        if style_id == "Caption" and text:
            caption_count += 1

    if table_count > 0 or heading_count > 1 or caption_count > 0 or non_empty_paragraphs > 3:
        return "full"
    return "cover"


def _capture_template_prototypes(document: Any) -> None:
    prototype_specs = {
        "inventory_server": ("STT", "Máy chủ", "Địa chỉ truy cập", "Phiên bản hệ điều hành"),
        "inventory_client": ("STT", "Máy trạm", "Địa chỉ truy cập", "Phiên bản hệ điều hành"),
        "summary_server": ("STT", "Máy chủ", "Kết quả rà soát đánh giá"),
        "summary_client": ("STT", "Máy trạm", "Kết quả rà soát đánh giá"),
        "detail": ("STT", "Hạng mục rà soát", "Kết quả rà soát"),
        "remediation_client": ("STT", "Máy trạm", "Địa chỉ IP", "Trạng thái"),
        "ioc": ("STT", "Thông tin", "Chi tiết"),
    }

    compact_enabled = compact_prototype_enabled()
    prototypes: dict[str, Any] = {}
    blueprints: dict[str, TableBlueprint] = {}
    for key, headers in prototype_specs.items():
        table = _find_table_by_headers(document, headers)
        if table is not None:
            prototypes[key] = deepcopy(table._tbl)
            if not compact_enabled:
                continue
            try:
                blueprints[key] = compile_table_blueprint(table)
            except UnsafeTableBlueprintError:
                # Complex/custom tables deliberately remain on the exact legacy
                # XML path. This is a normal classifier outcome, not an error.
                continue
            except Exception as exc:
                # The feature is experimental and must never make an otherwise
                # valid template fail. Do not include exception text or table
                # content in logs because templates may contain sensitive data.
                _LOGGER.warning(
                    "Compact prototype compilation failed (%s); using legacy path.",
                    type(exc).__name__,
                )

    if prototypes:
        setattr(document, "_codex_table_prototypes", prototypes)
    setattr(document, "_reporter_compact_prototype_enabled", compact_enabled)
    setattr(document, "_reporter_table_blueprints", blueprints)


def _find_table_by_headers(document: Any, headers: tuple[str, ...]) -> Any | None:
    for table in document.tables:
        if not table.rows:
            continue
        current_headers = tuple(cell.text.strip() for cell in table.rows[0].cells)
        if current_headers == headers:
            return table
    return None


def _replace_template_tokens(
    document: Any,
    *,
    title: str,
    organization: str,
    assessment_date: str | None,
) -> bool:
    token_map = {
        "{{TITLE}}": title,
        "{{ORGANIZATION}}": organization,
        "{{ASSESSMENT_DATE}}": assessment_date or str(date.today()),
        "{{ASSESSMENT_PERIOD}}": _format_assessment_period(assessment_date),
    }

    replaced_any = False
    for paragraph in _iter_document_paragraphs(document):
        if _replace_tokens_in_paragraph(paragraph, token_map):
            replaced_any = True
    return replaced_any


def _iter_document_paragraphs(document: Any) -> list[Any]:
    paragraphs: list[Any] = []
    paragraphs.extend(document.paragraphs)

    for table in document.tables:
        paragraphs.extend(_iter_table_paragraphs(table))

    for section in document.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
        for table in section.header.tables:
            paragraphs.extend(_iter_table_paragraphs(table))
        for table in section.footer.tables:
            paragraphs.extend(_iter_table_paragraphs(table))

    return paragraphs


def _iter_table_paragraphs(table: Any) -> list[Any]:
    paragraphs: list[Any] = []
    for row in table.rows:
        for cell in row.cells:
            paragraphs.extend(cell.paragraphs)
            for nested_table in cell.tables:
                paragraphs.extend(_iter_table_paragraphs(nested_table))
    return paragraphs


def _replace_tokens_in_paragraph(paragraph: Any, token_map: dict[str, str]) -> bool:
    if not paragraph.runs:
        return False

    combined_text = "".join(run.text for run in paragraph.runs)
    updated_text = combined_text
    for token, value in token_map.items():
        updated_text = updated_text.replace(token, value)

    if updated_text == combined_text:
        return False

    paragraph.runs[0].text = updated_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def _add_fallback_cover_page(
    document: Any,
    *,
    title: str,
    organization: str,
    assessment_date: str | None,
) -> None:
    modules = _docx_modules()
    WD_ALIGN_PARAGRAPH = modules["WD_ALIGN_PARAGRAPH"]

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.font.name = "Times New Roman"
    title_run.font.size = modules["Pt"](18)
    title_run.font.bold = True

    org_paragraph = document.add_paragraph()
    org_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_paragraph.add_run(f"Don vi: {organization}")
    _format_paragraph_runs(org_paragraph, preserve_color=False)

    date_paragraph = document.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.add_run(f"Thoi gian danh gia: {_format_assessment_period(assessment_date)}")
    _format_paragraph_runs(date_paragraph, preserve_color=False)

    document.add_page_break()


def _enable_field_updates(document: Any) -> None:
    modules = _docx_modules()
    OxmlElement = modules["OxmlElement"]
    qn = modules["qn"]

    settings = document.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    else:
        update_fields = existing
    update_fields.set(qn("w:val"), "true")

    # python-docx preserves the cached result stored in a template TOC. Mark
    # every main-document field dirty so Word knows the cached text is stale.
    # This is also the safe fallback on machines where no field engine exists.
    for field_char in document._element.iter(qn("w:fldChar")):
        if field_char.get(qn("w:fldCharType")) == "begin":
            field_char.set(qn("w:dirty"), "true")


def _apply_sample_page_layout(document: Any) -> None:
    modules = _docx_modules()
    Mm = modules["Mm"]

    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(25)
        section.right_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.header_distance = Mm(12)
        section.footer_distance = Mm(12)


def _configure_document(document: Any, *, using_template: bool = False) -> None:
    modules = _docx_modules()
    Pt = modules["Pt"]
    RGBColor = modules["RGBColor"]
    WD_ALIGN_PARAGRAPH = modules["WD_ALIGN_PARAGRAPH"]

    # A supplied template owns its typography and paragraph geometry. Any
    # direct style override here would also leak into headings by inheritance.
    if using_template:
        return

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(13)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.space_before = Pt(3)
    normal_style.paragraph_format.space_after = Pt(3)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        style.font.name = "Times New Roman"
        style.font.size = Pt(13)
        style.font.bold = True

    try:
        list_style = document.styles["List Paragraph"]
        list_style.font.name = "Times New Roman"
        list_style.font.size = Pt(13)
    except KeyError:
        pass


# ---------------------------------------------------------------------------
# Report sections
# ---------------------------------------------------------------------------

def _report_assets(data: dict[str, Any]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    servers = data.get("servers", [])
    clients = data.get("clients", [])
    return (
        servers if isinstance(servers, list) else [],
        clients if isinstance(clients, list) else [],
    )


def _asset_statistics(data: dict[str, Any]) -> tuple[int, int, int]:
    servers, clients = _report_assets(data)
    assets = [*servers, *clients]
    anomalous = sum(1 for asset in assets if _is_anomalous_asset(asset))
    return len(assets), anomalous, len(assets) - anomalous


def _add_summary_overview_section(
    document: Any,
    data: dict[str, Any],
    *,
    organization: str,
    assessment_date: str | None,
) -> None:
    servers, clients = _report_assets(data)
    org_label = _normalize_organization_name(organization)

    _add_heading(document, "Tổng quan", level=1)
    _add_heading(document, "Mục tiêu và phạm vi đánh giá", level=2)
    _add_body_paragraph(
        document,
        f"Báo cáo tổng hợp kết quả rà soát an toàn thông tin đối với các tài sản do {org_label} cung cấp. "
        f"Thời gian thực hiện: {_format_assessment_period(assessment_date)}.",
    )

    _add_heading(document, "Tài sản trong phạm vi", level=2)
    _add_body_paragraph(
        document,
        f"Phạm vi gồm {len(servers)} máy chủ và {len(clients)} máy trạm.",
    )
    _add_inventory_table(document, servers, "Máy chủ")
    _add_blank_paragraph(document)
    _add_inventory_table(document, clients, "Máy trạm")

    _add_heading(document, "Phương pháp thực hiện", level=2)
    for approach in ASSESSMENT_APPROACHES:
        _add_list_paragraph(document, approach, num_id=27)


def _add_summary_analysis_section(document: Any, data: dict[str, Any]) -> None:
    servers, clients = _report_assets(data)
    all_assets = [*servers, *clients]
    notable_assets = [asset for asset in all_assets if _is_anomalous_asset(asset)]

    _add_heading(document, "Kết quả và phân tích tổng hợp", level=1)
    _add_heading(document, "Kết quả tổng hợp máy chủ", level=2)
    _add_summary_table(document, servers, "Máy chủ")
    _add_heading(document, "Kết quả tổng hợp máy trạm", level=2)
    _add_summary_table(document, clients, "Máy trạm")

    _add_heading(document, "Các phát hiện đáng chú ý", level=2)
    if notable_assets:
        rows = [
            [
                str(index),
                str(asset.get("hostname", DEFAULT_TEXT_VALUE)),
                str(asset.get("ip", DEFAULT_TEXT_VALUE)),
                str(asset.get("notes") or asset.get("result") or ANOMALY_RESULT_TEXT),
            ]
            for index, asset in enumerate(notable_assets, start=1)
        ]
        table = _create_table(
            document,
            ["STT", "Tài sản", "Địa chỉ IP", "Phát hiện/ghi chú"],
            rows,
            column_widths_mm=[12, 42, 36, 70],
        )
        _style_table(table, left_align_columns={1, 2, 3})
    else:
        _add_body_paragraph(document, "Chưa ghi nhận phát hiện bất thường nổi bật trong dữ liệu được cung cấp.")

    total, anomalous, clean = _asset_statistics(data)
    _add_heading(document, "Đánh giá ảnh hưởng chung", level=2)
    _add_body_paragraph(
        document,
        f"Tổng cộng {total} tài sản được tổng hợp; {anomalous} tài sản có dấu hiệu cần xem xét thêm và "
        f"{clean} tài sản chưa ghi nhận dấu hiệu bất thường theo dữ liệu đầu vào.",
    )


def _add_summary_conclusion_section(document: Any, data: dict[str, Any]) -> None:
    total, anomalous, clean = _asset_statistics(data)
    _add_heading(document, "Kết luận và khuyến nghị", level=1)
    _add_heading(document, "Kết luận chung", level=2)
    _add_body_paragraph(
        document,
        f"Kết quả đánh giá bao phủ {total} tài sản, gồm {anomalous} tài sản cần tiếp tục xác minh và "
        f"{clean} tài sản chưa ghi nhận dấu hiệu bất thường từ dữ liệu hiện có.",
    )

    _add_heading(document, "Khuyến nghị ưu tiên", level=2)
    for item in DEFAULT_RECOMMENDATIONS:
        _add_list_paragraph(document, item, num_id=22)

    _add_heading(document, "Các bước tiếp theo", level=2)
    for item in (
        "Xác minh các phát hiện đáng chú ý với chủ sở hữu hệ thống và dữ liệu nhật ký liên quan.",
        "Lập kế hoạch xử lý theo mức độ ưu tiên và theo dõi trạng thái hoàn thành.",
        "Thực hiện rà soát lại sau khắc phục để xác nhận rủi ro đã được kiểm soát.",
    ):
        _add_list_paragraph(document, item, num_id=22)


def _add_technical_overview_section(
    document: Any,
    data: dict[str, Any],
    *,
    organization: str,
    assessment_date: str | None,
) -> None:
    servers, clients = _report_assets(data)
    org_label = _normalize_organization_name(organization)

    _add_heading(document, "Tổng quan", level=1)
    _add_heading(document, "Mục tiêu và phạm vi đánh giá", level=2)
    _add_body_paragraph(
        document,
        f"Thực hiện rà soát kỹ thuật các tài sản do {org_label} cung cấp trong thời gian "
        f"{_format_assessment_period(assessment_date)}.",
    )

    _add_heading(document, "Danh sách tài sản đánh giá", level=2)
    _add_inventory_table(document, servers, "Máy chủ")
    _add_blank_paragraph(document)
    _add_inventory_table(document, clients, "Máy trạm")

    _add_heading(document, "Phương pháp thực hiện", level=2)
    for item, level in METHOD_CHECKLIST_ITEMS:
        _add_list_paragraph(document, item, num_id=25, level=level)

    _add_heading(document, "Công cụ và nguồn dữ liệu", level=2)
    for tool in TOOLS_USED:
        _add_list_paragraph(document, tool, num_id=26)
    _add_body_paragraph(document, "Nguồn dữ liệu phân tích gồm kết quả rà soát và ghi chú kỹ thuật được nhập vào hệ thống.")


_TECHNICAL_CATEGORY_LABELS = {
    "rootkit": "Dấu hiệu rootkit",
    "autorun": "Cơ chế tự khởi động",
    "service_process": "Dịch vụ và tiến trình",
    "scheduled_task": "Tác vụ lập lịch",
    "suspicious_file": "Tệp tin đáng ngờ",
    "network": "Kết nối mạng đáng ngờ",
    "account": "Tài khoản và thông tin xác thực",
    "shared_file": "Tệp/thư mục chia sẻ",
    "tunnel": "Kênh tunnel hoặc proxy",
    "named_pipe": "Named pipe",
    "prefetch": "Dữ liệu Prefetch",
    "webshell": "Dấu hiệu WebShell",
    "web_log": "Nhật ký truy cập Web",
}


def _add_technical_analysis_section(document: Any, data: dict[str, Any]) -> None:
    servers, clients = _report_assets(data)
    assets = [*servers, *clients]

    _add_heading(document, "Phân tích chi tiết", level=1)
    _add_heading(document, "Tổng hợp kết quả", level=2)
    _add_heading(document, "Kết quả tổng hợp máy chủ", level=3)
    _add_summary_table(document, servers, "Máy chủ")
    _add_heading(document, "Kết quả tổng hợp máy trạm", level=3)
    _add_summary_table(document, clients, "Máy trạm")

    _add_heading(document, "Findings từ rule engine", level=2)
    rule_rows: list[list[str]] = []
    for asset in assets:
        for finding in asset.get("findings", []):
            evidence = "; ".join(
                f"{item.get('field')}: {item.get('value')}" for item in finding.get("evidence", [])
            )
            rule_rows.append([
                str(asset.get("hostname", DEFAULT_TEXT_VALUE)),
                str(finding.get("ruleId", "N/A")),
                str(finding.get("severity", "N/A")),
                str(finding.get("classification", "N/A")),
                evidence,
            ])
    table = _create_table(
        document,
        ["Tài sản", "Rule", "Mức độ", "Phân loại", "Bằng chứng"],
        rule_rows or [["", "", "", "", "Chưa có finding đủ bằng chứng."]],
    )
    _style_table(table, left_align_columns={0, 1, 2, 3, 4})

    _add_heading(document, "Phân tích máy chủ", level=2)
    if not servers:
        _add_body_paragraph(document, "Chưa có dữ liệu máy chủ để phân tích chi tiết.")
    for asset in servers:
        _add_heading(document, str(asset.get("hostname", DEFAULT_TEXT_VALUE)), level=3)
        _add_detail_table(document, asset=asset)

    _add_heading(document, "Phân tích máy trạm", level=2)
    if not clients:
        _add_body_paragraph(document, "Chưa có dữ liệu máy trạm để phân tích chi tiết.")
    for asset in clients:
        _add_heading(document, str(asset.get("hostname", DEFAULT_TEXT_VALUE)), level=3)
        _add_detail_table(document, asset=asset)

    _add_heading(document, "Phân tích điều tra", level=2)
    categorized: dict[str, list[dict[str, Any]]] = {}
    for asset in assets:
        if not _is_anomalous_asset(asset):
            continue
        for category in _finding_categories(asset):
            categorized.setdefault(category, []).append(asset)
    if not categorized:
        _add_body_paragraph(document, "Chưa có bằng chứng đủ để tạo các nhóm phân tích điều tra chuyên sâu.")
    for category, category_assets in sorted(
        categorized.items(), key=lambda item: _TECHNICAL_CATEGORY_LABELS.get(item[0], item[0])
    ):
        _add_heading(document, _TECHNICAL_CATEGORY_LABELS.get(category, category.replace("_", " ").title()), level=3)
        for asset in category_assets:
            evidence = str(asset.get("notes") or asset.get("result") or ANOMALY_RESULT_TEXT)
            _add_list_paragraph(
                document,
                f"{asset.get('hostname', DEFAULT_TEXT_VALUE)} ({asset.get('ip', DEFAULT_TEXT_VALUE)}): {evidence}",
                num_id=22,
            )

    _add_heading(document, "Indicators of Compromise (IoCs)", level=2)
    raw_iocs: list[dict[str, Any]] = []
    for asset in assets:
        extras = asset.get("extras") if isinstance(asset.get("extras"), dict) else {}
        iocs = extras.get("iocs", asset.get("iocs", []))
        if not isinstance(iocs, list):
            continue
        for ioc in iocs:
            raw = dict(ioc) if isinstance(ioc, dict) else {"value": ioc}
            raw.setdefault("source", asset.get("hostname", DEFAULT_TEXT_VALUE))
            raw_iocs.append(raw)
    normalized = normalize_iocs(raw_iocs, default_source="asset")
    ioc_rows = [
        [str(index), item["value"], f"{item['type']} · {'valid' if item['valid'] else 'invalid'} · {', '.join(item['sources'])}"]
        for index, item in enumerate(normalized, start=1)
    ]
    table = _create_table(
        document,
        ["STT", "Thông tin", "Chi tiết"],
        ioc_rows or [["1", "", "Chưa ghi nhận IoC có cấu trúc trong dữ liệu đầu vào."]],
        prototype_key="ioc",
        column_widths_mm=[14, 38, 108],
    )
    _style_table(table, left_align_columns={1, 2})

    _add_heading(document, "Hành động xử lý và gỡ bỏ", level=2)
    _add_technical_remediation_scope(document, servers, "Máy chủ")
    _add_technical_remediation_scope(document, clients, "Máy trạm")


def _add_technical_remediation_scope(
    document: Any,
    assets: list[dict[str, Any]],
    asset_label: str,
) -> None:
    _add_heading(document, asset_label, level=3)
    if not assets:
        _add_body_paragraph(document, f"Chưa có {asset_label.lower()} trong phạm vi để ghi nhận hành động xử lý.")
        return
    rows = []
    for index, asset in enumerate(assets, start=1):
        extras = asset.get("extras") if isinstance(asset.get("extras"), dict) else {}
        remediation = extras.get("remediation") or asset.get("remediation")
        status = str(remediation or "Chưa ghi nhận hành động xử lý bổ sung.")
        rows.append([str(index), str(asset.get("hostname", DEFAULT_TEXT_VALUE)), str(asset.get("ip", DEFAULT_TEXT_VALUE)), status])
    table = _create_table(
        document,
        ["STT", asset_label, "Địa chỉ IP", "Hành động/Trạng thái"],
        rows,
        prototype_key="remediation_client" if asset_label == "Máy trạm" else None,
        column_widths_mm=[12, 44, 34, 70],
    )
    _style_table(table, left_align_columns={1, 2, 3})


def _add_technical_conclusion_section(document: Any, data: dict[str, Any]) -> None:
    total, anomalous, clean = _asset_statistics(data)
    _add_heading(document, "Kết luận và khuyến nghị", level=1)
    _add_heading(document, "Kết luận kỹ thuật", level=2)
    _add_body_paragraph(
        document,
        f"Đã tổng hợp {total} tài sản; {anomalous} tài sản có dấu hiệu cần xác minh thêm và "
        f"{clean} tài sản chưa ghi nhận dấu hiệu bất thường từ nguồn dữ liệu hiện có.",
    )

    _add_heading(document, "Khuyến nghị khắc phục", level=2)
    for item in DEFAULT_RECOMMENDATIONS:
        _add_list_paragraph(document, item, num_id=22)

    _add_heading(document, "Khuyến nghị giám sát", level=2)
    for item in (
        "Theo dõi liên tục các chỉ dấu đã ghi nhận trên hệ thống giám sát tập trung.",
        "Rà soát biến động tiến trình, kết nối mạng, tài khoản và cơ chế duy trì truy cập.",
        "Đối chiếu lại nhật ký sau khi hoàn tất các hành động khắc phục.",
    ):
        _add_list_paragraph(document, item, num_id=22)

    _add_heading(document, "Hạn chế và nội dung cần xác minh thêm", level=2)
    _add_body_paragraph(
        document,
        "Kết luận phụ thuộc vào phạm vi, thời điểm thu thập và độ đầy đủ của dữ liệu đầu vào. "
        "Các phát hiện đáng chú ý cần được đối chiếu với nhật ký gốc và xác nhận bởi chủ sở hữu hệ thống.",
    )

def _add_overview_section(
    document: Any,
    data: dict[str, Any],
    *,
    organization: str,
    assessment_date: str | None,
    include_servers: bool = True,
    include_clients: bool = True,
) -> None:
    org_label = _normalize_organization_name(organization)

    _add_heading(document, "Tổng quan", level=1)
    _add_heading(document, "Phạm vi và mục đích thực hiện", level=2)
    _add_body_paragraph(document, f"Thời gian thực hiện đánh giá: {_format_assessment_period(assessment_date)}.")
    if include_servers:
        _add_body_paragraph(document, f"Thực hiện đánh giá các máy chủ do phía {org_label} cung cấp:")
        _add_inventory_table(document, data.get("servers", []), "Máy chủ")
        _add_blank_paragraph(document)
    if include_clients:
        _add_body_paragraph(document, f"Thực hiện đánh giá các máy trạm do phía {org_label} cung cấp:")
        _add_inventory_table(document, data.get("clients", []), "Máy trạm")
        _add_blank_paragraph(document)

    _add_heading(document, "Phương pháp thực hiện", level=2)
    _add_body_paragraph(document, "Checklist công việc thực hiện trong quá trình kiểm tra:")
    for item, level in METHOD_CHECKLIST_ITEMS:
        _add_list_paragraph(document, item, num_id=25, level=level)

    _add_body_paragraph(document, "Công cụ sử dụng:")
    for tool in TOOLS_USED:
        _add_list_paragraph(document, tool, num_id=26)

    _add_body_paragraph(document, "Các cách thức thực hiện:")
    for approach in ASSESSMENT_APPROACHES:
        _add_list_paragraph(document, approach, num_id=27)


def _add_results_section(
    document: Any,
    data: dict[str, Any],
    *,
    include_servers: bool = True,
    include_clients: bool = True,
) -> None:
    servers = data.get("servers", [])
    clients = data.get("clients", [])

    _add_heading(document, "Kết quả", level=1)

    if include_servers:
        _add_heading(document, "Đánh giá chung đối với máy chủ", level=2)
        _add_body_paragraph(document, "Kết quả được mô tả sơ lược trong bảng sau:")
        _add_summary_table(document, servers, "Máy chủ")

    if include_clients:
        _add_heading(document, "Đánh giá chung với các máy trạm", level=2)
        _add_body_paragraph(document, "Kết quả được mô tả sơ lược trong bảng sau:")
        _add_summary_table(document, clients, "Máy trạm")

    if include_servers:
        _add_heading(document, "Chi tiết kết quả CA các máy chủ", level=2)
        if not servers:
            _add_body_paragraph(document, "Chưa có dữ liệu máy chủ để trình bày chi tiết.")
        for asset in servers:
            _add_heading(document, f"Kết quả thực hiện CA trên máy chủ {asset.get('hostname', DEFAULT_TEXT_VALUE)}", level=3)
            _add_detail_table(document, asset=asset)

    if include_clients:
        _add_heading(document, "Chi tiết kết quả CA các máy trạm", level=2)
        if not clients:
            _add_body_paragraph(document, "Chưa có dữ liệu máy trạm để trình bày chi tiết.")
        for asset in clients:
            _add_heading(document, f"Kết quả thực hiện CA trên máy trạm {asset.get('hostname', DEFAULT_TEXT_VALUE)}", level=3)
            _add_detail_table(document, asset=asset)


def _add_investigation_section(
    document: Any,
    *,
    include_servers: bool = True,
    include_clients: bool = True,
) -> None:
    _add_heading(document, "Phân tích điều tra", level=1)
    if include_servers:
        _add_heading(document, "Phân tích điều tra máy chủ", level=2)
        _add_body_paragraph(document, "Chưa ghi nhận nội dung cần bổ sung cho mục phân tích điều tra máy chủ.")
    if include_clients:
        _add_heading(document, "Phân tích điều tra máy trạm", level=2)
        _add_body_paragraph(document, "Chưa ghi nhận nội dung cần bổ sung cho mục phân tích điều tra máy trạm.")


def _add_remediation_section(
    document: Any,
    data: dict[str, Any],
    *,
    include_servers: bool = True,
    include_clients: bool = True,
) -> None:
    _add_heading(document, "Gỡ bỏ mã độc", level=1)
    if include_servers:
        _add_heading(document, "Gỡ bỏ mã độc trên máy chủ", level=2)
        _add_body_paragraph(document, "Chưa có thao tác gỡ bỏ mã độc bổ sung cần ghi nhận cho các máy chủ trong phạm vi.")
    if not include_clients:
        return
    _add_heading(document, "Gỡ bỏ mã độc trên máy trạm", level=2)
    clients = data.get("clients", [])
    if not clients:
        _add_body_paragraph(document, "Chưa có thao tác gỡ bỏ mã độc bổ sung cần ghi nhận cho các máy trạm trong phạm vi.")
        return

    rows = []
    for index, asset in enumerate(clients, start=1):
        rows.append(
            [
                str(index),
                asset.get("hostname", DEFAULT_TEXT_VALUE),
                asset.get("ip", DEFAULT_TEXT_VALUE),
                "Chưa ghi nhận mã độc cần gỡ bỏ.",
            ]
        )

    table = _create_table(
        document,
        ["STT", "Máy trạm", "Địa chỉ IP", "Trạng thái"],
        rows,
        prototype_key="remediation_client",
        column_widths_mm=[12, 50, 34, 64],
    )
    _style_table(table, center_all=True)


def _add_ioc_section(document: Any) -> None:
    _add_heading(document, "Indicators of compromise (IoCs)", level=1)
    table = _create_table(
        document,
        ["STT", "Thông tin", "Chi tiết"],
        [["1", "", ""]],
        prototype_key="ioc",
        column_widths_mm=[14, 38, 108],
    )
    _style_table(table, left_align_columns={1, 2})


def _add_recommendations_section(
    document: Any,
    *,
    include_servers: bool = True,
    include_clients: bool = True,
) -> None:
    _add_heading(document, "Khuyến nghị & khắc phục", level=1)
    for item in DEFAULT_RECOMMENDATIONS:
        normalized = item.lower()
        if not include_servers and "máy chủ" in normalized:
            continue
        if not include_clients and "máy trạm" in normalized:
            continue
        _add_list_paragraph(document, item, num_id=22)


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------

def _add_inventory_table(document: Any, assets: list[dict[str, Any]], asset_label: str) -> None:
    rows = []
    for index, asset in enumerate(assets, start=1):
        rows.append(
            [
                str(index),
                asset.get("hostname", DEFAULT_TEXT_VALUE),
                asset.get("ip", DEFAULT_TEXT_VALUE),
                asset.get("os", DEFAULT_TEXT_VALUE),
            ]
        )

    if not rows:
        rows = [["", "", "", ""]]

    table = _create_table(
        document,
        ["STT", asset_label, "Địa chỉ truy cập", "Phiên bản hệ điều hành"],
        rows,
        prototype_key="inventory_server" if asset_label == "Máy chủ" else "inventory_client",
        column_widths_mm=[12, 52, 37, 59],
    )
    _style_table(table, center_all=True)


def _add_summary_table(document: Any, assets: list[dict[str, Any]], asset_label: str) -> None:
    rows = []
    for index, asset in enumerate(assets, start=1):
        # The primary conclusion is deliberately standardized. Original result/notes
        # remain available in evidence and detail sections for traceability.
        result_text = assessment_text(asset.get("findings"))
        rows.append([str(index), asset.get("hostname", DEFAULT_TEXT_VALUE), result_text])

    if not rows:
        rows = [["", "", ""]]

    table = _create_table(
        document,
        ["STT", asset_label, "Kết quả rà soát đánh giá"],
        rows,
        prototype_key="summary_server" if asset_label == "Máy chủ" else "summary_client",
        column_widths_mm=[12, 52, 96],
    )
    _style_table(table, center_all=True)


def _add_detail_table(document: Any, asset: dict[str, Any] | None = None) -> None:
    # A review-only rule may appear in summary/findings, but must not turn every
    # detailed checklist row into a confirmed anomaly.
    is_anomalous = _has_confirmed_anomaly(asset) if asset else False

    rows = []
    for index, item in enumerate(DETAIL_CHECKLIST_ITEMS, start=1):
        if is_anomalous:
            result_text = _get_anomaly_text(item, asset)
        else:
            result_text = DEFAULT_RESULT_TEXT
        rows.append([str(index), item, result_text])

    table = _create_table(
        document,
        ["STT", "Hạng mục rà soát", "Kết quả rà soát"],
        rows,
        prototype_key="detail",
        column_widths_mm=[12, 74, 74],
    )
    _style_table(table, left_align_columns={1, 2})


def _is_anomalous_asset(asset: dict[str, Any] | None) -> bool:
    """Kiểm tra asset có ghi nhận bất thường hay không dựa trên result + notes."""
    if not asset:
        return False
    if "findings" in asset and isinstance(asset["findings"], list):
        return any(
            finding.get("classification") in {"anomaly", "needs_review"}
            and bool(finding.get("evidence"))
            for finding in asset["findings"] if isinstance(finding, dict)
        )
    result = str(asset.get("result", "")).lower()
    notes = str(asset.get("notes", "")).lower()
    combined = result + " " + notes

    # Kiểm tra phủ định trước — nếu result rõ ràng "Không phát hiện" → clean
    if any(ck in result for ck in _CLEAN_KEYWORDS):
        return False

    return any(kw in combined for kw in _ANOMALY_KEYWORDS)


def _has_confirmed_anomaly(asset: dict[str, Any] | None) -> bool:
    if not asset:
        return False
    findings = asset.get("findings")
    if isinstance(findings, list):
        return any(
            isinstance(finding, dict)
            and finding.get("classification") == "anomaly"
            and bool(finding.get("evidence"))
            for finding in findings
        )
    return _is_anomalous_asset(asset)


# Mapping checklist items → mô tả kết quả cho máy bất thường
_ANOMALY_ITEM_KEYWORDS = {
    "tệp tin bất thường": "Phát hiện tệp tin bất thường",
    "kết nối mạng bất thường": "Phát hiện kết nối mạng bất thường",
    "rootkit": "Phát hiện dấu hiệu bất thường",
    "autorun entry": "Phát hiện dấu hiệu bất thường",
    "service, process": "Phát hiện dấu hiệu bất thường",
    "lập lịch bất thường": "Phát hiện dấu hiệu bất thường",
    "tài khoản": "Phát hiện dấu hiệu bất thường",
    "file/thư mục chia sẻ": "Không ghi nhận dấu hiệu bất thường",
    "tunnel": "Không ghi nhận dấu hiệu bất thường",
    "named pipe": "Phát hiện dấu hiệu bất thường",
    "prefetch": "Phát hiện dấu hiệu bất thường",
    "webshell": "Phát hiện dấu hiệu bất thường",
    "access log": "Phát hiện dấu hiệu bất thường",
}


_FINDING_CATEGORY_RULES: dict[str, tuple[str, ...]] = {
    "rootkit": ("rootkit",),
    "autorun": ("autorun", "run key", "startup"),
    "service_process": (
        "process", "service", "dll sideload", "sideloading", "svchost",
        "lsass", "mimikatz", "injection", "cobalt", "beacon",
    ),
    "scheduled_task": ("scheduled task", "schedule task", "crontab", "cron", "persistence"),
    "suspicious_file": (
        "malware", "mã độc", "trojan", "virus", "ransomware", "dropper",
        "plugx", "shadowpad", "emotet", "payload", "backdoor", "keylogger",
    ),
    "network": ("c2", "callback", "reverse shell", "beacon", "smtp", "kết nối"),
    "account": ("credential", "mimikatz", "account", "tài khoản", "lsass"),
    "shared_file": ("shared folder", "file share", "thư mục chia sẻ"),
    "tunnel": ("tunnel", "proxy",),
    "named_pipe": ("named pipe",),
    "prefetch": ("prefetch",),
    "webshell": ("webshell", "aspxspy",),
    "web_log": ("access log", "web log", "http request",),
}

_CHECKLIST_CATEGORY_RULES: tuple[tuple[str, str], ...] = (
    ("rootkit", "rootkit"),
    ("autorun entry", "autorun"),
    ("service, process", "service_process"),
    ("schedule task", "scheduled_task"),
    ("tệp tin bất thường", "suspicious_file"),
    ("kết nối mạng bất thường", "network"),
    ("tài khoản", "account"),
    ("file/thư mục chia sẻ", "shared_file"),
    ("tunnel", "tunnel"),
    ("named pipe", "named_pipe"),
    ("prefetch", "prefetch"),
    ("webshell", "webshell"),
    ("access log", "web_log"),
)


def _finding_categories(asset: dict[str, Any]) -> set[str]:
    """Return evidence-backed categories; structured findings take priority."""
    categories: set[str] = set()
    extras = asset.get("extras") if isinstance(asset.get("extras"), dict) else {}
    findings = extras.get("findings", asset.get("findings", []))
    if isinstance(findings, list):
        for finding in findings:
            if isinstance(finding, dict) and finding.get("category"):
                categories.add(str(finding["category"]).strip().lower())

    evidence = f"{asset.get('result', '')} {asset.get('notes', '')}".lower()
    for category, keywords in _FINDING_CATEGORY_RULES.items():
        if any(keyword in evidence for keyword in keywords):
            categories.add(category)
    return categories


def _get_anomaly_text(checklist_item: str, asset: dict[str, Any]) -> str:
    """Return a finding only when the checklist row has matching evidence."""
    notes = str(asset.get("notes", "")).strip()
    item_lower = checklist_item.lower()

    if "sử dụng công cụ" in item_lower or "thu thập và phân tích" in item_lower:
        return f"{ANOMALY_RESULT_TEXT}: {notes}" if notes else ANOMALY_RESULT_TEXT

    categories = _finding_categories(asset)
    for keyword, category in _CHECKLIST_CATEGORY_RULES:
        if keyword in item_lower:
            if category not in categories:
                return DEFAULT_RESULT_TEXT
            return f"{ANOMALY_RESULT_TEXT}: {notes}" if notes else ANOMALY_RESULT_TEXT

    return DEFAULT_RESULT_TEXT


def _table_metric_category(headers: list[str], prototype_key: str | None) -> str:
    prototype_categories = {
        "inventory_server": "assetInventory",
        "inventory_client": "assetInventory",
        "summary_server": "assetSummary",
        "summary_client": "assetSummary",
        "detail": "assetDetail",
        "remediation_client": "remediation",
        "ioc": "ioc",
    }
    if prototype_key in prototype_categories:
        return prototype_categories[prototype_key]

    normalized = tuple(str(header).strip().casefold() for header in headers)
    joined = " ".join(normalized)
    if normalized[:2] == ("thuộc tính", "giá trị"):
        return "incidentMetadata"
    if "thời gian" in normalized and "sự kiện" in normalized:
        return "timeline"
    if "rule" in normalized:
        return "findings"
    if "technique" in normalized and "tactic" in normalized:
        return "mitre"
    if "hành động" in normalized and "phụ trách" in normalized:
        return "incidentActions"
    if "tài sản" in normalized and "địa chỉ ip" in normalized:
        return "incidentAssets"
    if "thông tin" in normalized and ("chi tiết" in normalized or "nguồn" in normalized):
        return "ioc"
    if "hạng mục rà soát" in normalized:
        return "assetDetail"
    if "kết quả rà soát đánh giá" in normalized or "phát hiện/ghi chú" in normalized:
        return "assetSummary"
    if "phiên bản hệ điều hành" in normalized:
        return "assetInventory"
    if "trạng thái" in joined or "hành động/" in joined:
        return "remediation"
    return "other"


def _create_table(
    document: Any,
    headers: list[str],
    rows: list[list[str]],
    *,
    prototype_key: str | None = None,
    column_widths_mm: list[int] | None = None,
) -> Any:
    metrics = _CURRENT_PERFORMANCE_METRICS.get()
    metric_started_ns = time.perf_counter_ns() if metrics is not None else 0
    metric_category = (
        _table_metric_category(headers, prototype_key)
        if metrics is not None
        else ""
    )
    prototype = _get_table_prototype(document, prototype_key)
    blueprint = _get_table_blueprint(document, prototype_key)
    table = None
    if (
        prototype is not None
        and blueprint is not None
        and _compact_blueprint_matches_output(blueprint, headers, rows)
    ):
        compact_table = None
        try:
            compact_table = _create_table_from_blueprint(
                document,
                blueprint,
                headers,
                rows,
            )
            if column_widths_mm:
                _set_table_layout(compact_table, column_widths_mm)
            document._body._element._insert_tbl(compact_table._tbl)
            table = compact_table
        except Exception as exc:
            _remove_table_if_attached(document, compact_table)
            # Fast-path failures must be recoverable without exposing template
            # contents or source rows through logs.
            _LOGGER.warning(
                "Compact prototype table build failed (%s); using legacy path.",
                type(exc).__name__,
            )

    if table is None and prototype is not None:
        table = _create_table_from_prototype(document, prototype, headers, rows)
        if column_widths_mm:
            _set_table_layout(table, column_widths_mm)
    elif table is None:
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        _set_table_layout(table, column_widths_mm)
        _ensure_table_borders(table)

        header_cells = table.rows[0].cells
        for index, value in enumerate(headers):
            _set_cell_text(header_cells[index], value, bold=True, centered=True)

        for row_values in rows:
            row_cells = table.add_row().cells
            for index, value in enumerate(row_values):
                _set_cell_text(row_cells[index], value, bold=False, centered=False)
            _advance_build_tracker()

    if metrics is not None:
        setattr(table, "_reporter_metric_category", metric_category)
        metrics.record_aggregate(
            "tableCreate",
            metric_category,
            (time.perf_counter_ns() - metric_started_ns) / 1_000_000,
        )
    return table


def _get_table_prototype(document: Any, prototype_key: str | None) -> Any | None:
    if not prototype_key:
        return None
    prototypes = getattr(document, "_codex_table_prototypes", {})
    return prototypes.get(prototype_key)


def _get_table_blueprint(
    document: Any,
    prototype_key: str | None,
) -> TableBlueprint | None:
    if (
        not prototype_key
        or not getattr(document, "_reporter_compact_prototype_enabled", False)
    ):
        return None
    blueprints = getattr(document, "_reporter_table_blueprints", {})
    return blueprints.get(prototype_key)


def _compact_blueprint_matches_output(
    blueprint: TableBlueprint,
    headers: list[str],
    rows: list[list[str]],
) -> bool:
    """Use the first rollout only for one unambiguous, rectangular row style."""

    try:
        column_count = int(blueprint.column_count)
        if (
            blueprint.schema_version != BLUEPRINT_SCHEMA_VERSION
            or blueprint.data_row_variant_count != 1
            or column_count < 1
            or column_count != len(headers)
        ):
            return False
        return all(len(row) == column_count for row in rows)
    except (AttributeError, TypeError, ValueError):
        return False


def _remove_table_if_attached(document: Any, table: Any | None) -> None:
    if table is None:
        return
    table_element = getattr(table, "_tbl", None)
    body = document._body._element
    if table_element is not None and table_element.getparent() is body:
        body.remove(table_element)


def _create_table_from_blueprint(
    document: Any,
    blueprint: TableBlueprint,
    headers: list[str],
    rows: list[list[str]],
) -> Any:
    """Build a compact prototype table completely before attaching it."""

    modules = _docx_modules()
    Table = modules["Table"]

    tbl = blueprint.to_table_element()
    table = Table(tbl, document._body)
    if len(table.rows) != 2:
        raise ValueError("Compact blueprint must contain one header and one data row.")
    if (
        len(table.rows[0].cells) != len(headers)
        or len(table.rows[1].cells) != len(headers)
    ):
        raise ValueError("Compact blueprint column count does not match output.")

    data_row_prototype = deepcopy(table.rows[1]._tr)
    if fast_cell_enabled():
        _normalize_simple_data_row_prototype(data_row_prototype)
    table._tbl.remove(table.rows[1]._tr)

    for row_values in rows:
        table._tbl.append(deepcopy(data_row_prototype))
        row_cells = table.rows[-1].cells
        if len(row_cells) != len(row_values):
            raise ValueError("Compact blueprint data row is not rectangular.")
        for index, value in enumerate(row_values):
            _set_cell_text(
                row_cells[index],
                value,
                bold=False,
                centered=False,
                strip_list_numbering=index == 0,
            )
        _advance_build_tracker()

    header_cells = table.rows[0].cells
    for index, value in enumerate(headers):
        _set_cell_text(header_cells[index], value, bold=True, centered=True)

    # Preserve the legacy OOXML property order by applying borders before the
    # caller optionally applies column layout. The former second border pass was
    # idempotent and is intentionally omitted.
    _ensure_table_borders(table)
    return table


def _create_table_from_prototype(
    document: Any,
    prototype_table: Any,
    headers: list[str],
    rows: list[list[str]],
) -> Any:
    modules = _docx_modules()
    Table = modules["Table"]

    tbl = deepcopy(prototype_table)
    document._body._element._insert_tbl(tbl)
    table = Table(tbl, document._body)

    data_row_prototype = deepcopy(table.rows[1]._tr) if len(table.rows) > 1 else None
    if data_row_prototype is not None and fast_cell_enabled():
        _normalize_simple_data_row_prototype(data_row_prototype)
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)

    for row_values in rows:
        if data_row_prototype is not None:
            table._tbl.append(deepcopy(data_row_prototype))
        else:
            table.add_row()
        row_cells = table.rows[-1].cells
        for index, value in enumerate(row_values):
            _set_cell_text(
                row_cells[index],
                value,
                bold=False,
                centered=False,
                strip_list_numbering=index == 0,
            )
        _advance_build_tracker()

    header_cells = table.rows[0].cells
    for index, value in enumerate(headers):
        _set_cell_text(header_cells[index], value, bold=True, centered=True)

    # See the compact builder above: one pre-layout border pass keeps byte-level
    # format parity with the legacy sequence without repeating the operation.
    _ensure_table_borders(table)
    return table


def _set_table_layout(table: Any, column_widths_mm: list[int] | None) -> None:
    modules = _docx_modules()
    Mm = modules["Mm"]
    WD_TABLE_ALIGNMENT = modules["WD_TABLE_ALIGNMENT"]
    OxmlElement = modules["OxmlElement"]
    qn = modules["qn"]

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")

    tbl_jc = tbl_pr.find(qn("w:jc"))
    if tbl_jc is None:
        tbl_jc = OxmlElement("w:jc")
        tbl_pr.append(tbl_jc)
    tbl_jc.set(qn("w:val"), "center")

    if not column_widths_mm:
        return

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)

    for width in column_widths_mm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(Mm(width).twips)))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, column_widths_mm):
            cell.width = Mm(width)


def _ensure_table_borders(table: Any) -> None:
    modules = _docx_modules()
    OxmlElement = modules["OxmlElement"]
    qn = modules["qn"]

    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_element = tbl_borders.find(qn(f"w:{edge}"))
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            tbl_borders.append(edge_element)
        edge_element.set(qn("w:val"), "single")
        edge_element.set(qn("w:sz"), "4")
        edge_element.set(qn("w:space"), "0")
        edge_element.set(qn("w:color"), "auto")


def _style_table(table: Any, *, center_all: bool = False, left_align_columns: set[int] | None = None) -> None:
    metrics = _CURRENT_PERFORMANCE_METRICS.get()
    metric_started_ns = time.perf_counter_ns() if metrics is not None else 0
    left_align_columns = left_align_columns or set()
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            centered = row_index == 0 or center_all or col_index == 0
            if col_index in left_align_columns and row_index != 0 and not center_all:
                centered = False
            _format_cell(cell, bold=row_index == 0, centered=centered)
    if metrics is not None:
        metrics.record_aggregate(
            "tableStyle",
            str(getattr(table, "_reporter_metric_category", "other")),
            (time.perf_counter_ns() - metric_started_ns) / 1_000_000,
        )


def _set_cell_text(
    cell: Any,
    text: str,
    *,
    bold: bool,
    centered: bool,
    strip_list_numbering: bool = False,
) -> None:
    simple_candidate = len(cell.paragraphs) == 1
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[0]
        for extra_paragraph in cell.paragraphs[1:]:
            cell._tc.remove(extra_paragraph._p)

    p_pr = paragraph._p.get_or_add_pPr()
    qn = _docx_modules()["qn"]
    if strip_list_numbering:
        # Inventory prototypes may use Word list numbering in the STT cell.
        # A cloned row would then render the list marker plus our literal value
        # (11, 22, ...). Conversely, docx-preview does not render that marker
        # reliably when the literal value is omitted. Make STT deterministic:
        # remove both the numbering reference and its list paragraph style,
        # then write exactly one plain-text sequence value.
        for num_pr in list(p_pr.findall(qn("w:numPr"))):
            p_pr.remove(num_pr)
        p_style = p_pr.find(qn("w:pStyle"))
        if p_style is not None and p_style.get(qn("w:val")) == "ListParagraph":
            p_pr.remove(p_style)

    if not (
        fast_cell_enabled()
        and simple_candidate
        and _replace_simple_cell_text(cell, paragraph, text)
    ):
        _replace_paragraph_text(paragraph, text)
    # All production table builders immediately call ``_style_table`` after
    # construction. Keeping formatting there makes each cell pass through the
    # expensive run/paragraph formatter exactly once while preserving the same
    # final style for legacy, compact and generic tables.


def _advance_build_tracker(rows: int = 1) -> None:
    tracker = _CURRENT_BUILD_TRACKER.get()
    if tracker is not None:
        tracker.advance(rows)


def _replace_simple_cell_text(cell: Any, paragraph: Any, text: str) -> bool:
    """Replace one plain run in place; reject every ambiguous OOXML construct."""

    if any(character in str(text) for character in ("\n", "\r", "\t")):
        return False
    modules = _docx_modules()
    qn = modules["qn"]
    tc_pr = cell._tc.tcPr
    if tc_pr is not None and (
        tc_pr.find(qn("w:gridSpan")) is not None
        or tc_pr.find(qn("w:vMerge")) is not None
    ):
        return False
    children = list(paragraph._p)
    if not children or children[0].tag != qn("w:pPr"):
        return False
    run_elements = [child for child in children if child.tag == qn("w:r")]
    if len(children) != 2 or len(run_elements) != 1:
        return False
    run_element = run_elements[0]
    run_children = list(run_element)
    if not run_children or run_children[0].tag != qn("w:rPr"):
        return False
    if any(child.tag not in {qn("w:rPr"), qn("w:t")} for child in run_children):
        return False
    if len(paragraph.runs) != 1:
        return False
    if not _has_canonical_data_run_properties(run_element):
        return False
    paragraph.runs[0].text = str(text)
    return True


def _has_canonical_data_run_properties(run_element: Any) -> bool:
    """Return true only for the exact rPr emitted by the legacy style pass."""

    modules = _docx_modules()
    qn = modules["qn"]
    if run_element.attrib:
        return False
    r_pr = run_element.find(qn("w:rPr"))
    if r_pr is None or r_pr.attrib:
        return False
    children = list(r_pr)
    if [child.tag for child in children] != [
        qn("w:rFonts"),
        qn("w:b"),
        qn("w:color"),
        qn("w:sz"),
    ]:
        return False
    fonts, bold, color, size = children
    return (
        dict(fonts.attrib) == {
            qn("w:ascii"): "Times New Roman",
            qn("w:hAnsi"): "Times New Roman",
        }
        and dict(bold.attrib) == {qn("w:val"): "0"}
        and dict(color.attrib) == {qn("w:val"): "000000"}
        and dict(size.attrib) == {qn("w:val"): "26"}
    )


def _normalize_simple_data_row_prototype(row_element: Any) -> int:
    """Canonicalize plain prototype runs once before cloning many data rows."""

    modules = _docx_modules()
    OxmlElement = modules["OxmlElement"]
    qn = modules["qn"]
    normalized = 0
    for cell_element in row_element.iter(qn("w:tc")):
        tc_pr = cell_element.find(qn("w:tcPr"))
        if tc_pr is not None and (
            tc_pr.find(qn("w:gridSpan")) is not None
            or tc_pr.find(qn("w:vMerge")) is not None
        ):
            continue
        paragraphs = cell_element.findall(qn("w:p"))
        if len(paragraphs) != 1:
            continue
        paragraph = paragraphs[0]
        children = list(paragraph)
        if not children or children[0].tag != qn("w:pPr"):
            continue
        runs = [child for child in children if child.tag == qn("w:r")]
        if len(children) != 2 or len(runs) != 1:
            continue
        run = runs[0]
        run_children = list(run)
        if not run_children or run_children[0].tag != qn("w:rPr"):
            continue
        if any(child.tag not in {qn("w:rPr"), qn("w:t")} for child in run_children):
            continue

        run.attrib.clear()
        r_pr = run_children[0]
        r_pr.attrib.clear()
        for child in list(r_pr):
            r_pr.remove(child)
        for tag, attributes in (
            ("w:rFonts", {"w:ascii": "Times New Roman", "w:hAnsi": "Times New Roman"}),
            ("w:b", {"w:val": "0"}),
            ("w:color", {"w:val": "000000"}),
            ("w:sz", {"w:val": "26"}),
        ):
            property_element = OxmlElement(tag)
            for name, value in attributes.items():
                property_element.set(qn(name), value)
            r_pr.append(property_element)
        normalized += 1
    return normalized


def _format_cell(cell: Any, *, bold: bool, centered: bool) -> None:
    modules = _docx_modules()
    WD_ALIGN_PARAGRAPH = modules["WD_ALIGN_PARAGRAPH"]
    WD_CELL_VERTICAL_ALIGNMENT = modules["WD_CELL_VERTICAL_ALIGNMENT"]

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
        _format_paragraph_runs(paragraph, bold=bold, preserve_color=False)


# ---------------------------------------------------------------------------
# Heading / paragraph helpers
# ---------------------------------------------------------------------------

def _add_heading(document: Any, text: str, *, level: int) -> Any:
    paragraph = document.add_heading(text, level=level)
    if not _apply_heading_numbering(paragraph, level=level):
        numbering_prefix = _next_heading_prefix(document, level)
        if paragraph.runs:
            paragraph.runs[0].text = f"{numbering_prefix} {paragraph.runs[0].text}"
        else:
            paragraph.add_run(f"{numbering_prefix} {text}")
    return paragraph


def _add_body_paragraph(document: Any, text: str) -> Any:
    paragraph = document.add_paragraph(text)
    _format_paragraph_runs(paragraph, preserve_color=False)
    return paragraph


def _add_blank_paragraph(document: Any) -> None:
    paragraph = document.add_paragraph("")
    _format_paragraph_runs(paragraph, preserve_color=False)


def _add_list_paragraph(document: Any, text: str, *, num_id: int, level: int = 0) -> Any:
    paragraph = document.add_paragraph(style="List Paragraph")
    paragraph.add_run(text)
    _format_paragraph_runs(paragraph, preserve_color=False)

    if not _apply_list_numbering(paragraph, num_id=num_id, level=level):
        paragraph.style = "List Bullet"
        _format_paragraph_runs(paragraph, preserve_color=False)

    return paragraph


def _apply_list_numbering(paragraph: Any, *, num_id: int, level: int) -> bool:
    modules = _docx_modules()
    OxmlElement = modules["OxmlElement"]
    qn = modules["qn"]

    if not _numbering_id_exists(paragraph, num_id):
        return False

    p_pr = paragraph._p.get_or_add_pPr()
    current_num_pr = p_pr.find(qn("w:numPr"))
    if current_num_pr is not None:
        p_pr.remove(current_num_pr)

    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)
    p_pr.append(num_pr)
    return True


def _apply_heading_numbering(paragraph: Any, *, level: int) -> bool:
    numbering_id = _resolve_heading_numbering_id(paragraph)
    if numbering_id is None:
        return False
    return _apply_list_numbering(paragraph, num_id=numbering_id, level=max(level - 1, 0))


def _resolve_heading_numbering_id(paragraph: Any) -> int | None:
    cached = getattr(paragraph.part, "_codex_heading_numbering_id", None)
    if cached is not None:
        return cached

    try:
        numbering_root = paragraph.part.numbering_part.element
    except Exception:
        setattr(paragraph.part, "_codex_heading_numbering_id", None)
        return None

    selected = _create_reporter_heading_numbering(numbering_root)
    setattr(paragraph.part, "_codex_heading_numbering_id", selected)
    return selected


def _create_reporter_heading_numbering(numbering_root: Any) -> int | None:
    """Create a fresh Reporter heading list so every report starts at 1."""
    modules = _docx_modules()
    OxmlElement = modules["OxmlElement"]
    qn = modules["qn"]

    abstract_num = None
    for candidate in numbering_root.findall(qn("w:abstractNum")):
        marker = candidate.find(qn("w:tmpl"))
        if marker is not None and marker.get(qn("w:val")) == REPORTER_HEADING_NUMBERING_MARKER:
            abstract_num = candidate
            break

    if abstract_num is None:
        existing_abstract_ids: list[int] = []
        for candidate in numbering_root.findall(qn("w:abstractNum")):
            value = candidate.get(qn("w:abstractNumId"))
            try:
                existing_abstract_ids.append(int(value))
            except (TypeError, ValueError):
                continue
        abstract_id = max(existing_abstract_ids, default=-1) + 1
        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), str(abstract_id))

        nsid = OxmlElement("w:nsid")
        nsid.set(qn("w:val"), REPORTER_HEADING_NUMBERING_MARKER)
        abstract_num.append(nsid)
        multi_level_type = OxmlElement("w:multiLevelType")
        multi_level_type.set(qn("w:val"), "multilevel")
        abstract_num.append(multi_level_type)
        template_marker = OxmlElement("w:tmpl")
        template_marker.set(qn("w:val"), REPORTER_HEADING_NUMBERING_MARKER)
        abstract_num.append(template_marker)

        for level_index, level_text in enumerate(("%1.", "%1.%2.", "%1.%2.%3.")):
            level = OxmlElement("w:lvl")
            level.set(qn("w:ilvl"), str(level_index))

            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            level.append(start)
            number_format = OxmlElement("w:numFmt")
            number_format.set(qn("w:val"), "decimal")
            level.append(number_format)
            paragraph_style = OxmlElement("w:pStyle")
            paragraph_style.set(qn("w:val"), f"Heading{level_index + 1}")
            level.append(paragraph_style)
            text = OxmlElement("w:lvlText")
            text.set(qn("w:val"), level_text)
            level.append(text)
            justification = OxmlElement("w:lvlJc")
            justification.set(qn("w:val"), "left")
            level.append(justification)

            paragraph_properties = OxmlElement("w:pPr")
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "num")
            tab.set(qn("w:pos"), REPORTER_HEADING_TEXT_LEFT_TWIPS)
            tabs.append(tab)
            paragraph_properties.append(tabs)
            indentation = OxmlElement("w:ind")
            indentation.set(qn("w:left"), REPORTER_HEADING_TEXT_LEFT_TWIPS)
            indentation.set(qn("w:hanging"), REPORTER_HEADING_TEXT_LEFT_TWIPS)
            paragraph_properties.append(indentation)
            level.append(paragraph_properties)
            abstract_num.append(level)

        first_num = numbering_root.find(qn("w:num"))
        if first_num is None:
            numbering_root.append(abstract_num)
        else:
            numbering_root.insert(numbering_root.index(first_num), abstract_num)
    else:
        raw_abstract_id = abstract_num.get(qn("w:abstractNumId"))
        try:
            abstract_id = int(raw_abstract_id)
        except (TypeError, ValueError):
            return None

    existing_num_ids: list[int] = []
    for num in numbering_root.findall(qn("w:num")):
        value = num.get(qn("w:numId"))
        try:
            existing_num_ids.append(int(value))
        except (TypeError, ValueError):
            continue
    num_id = max(existing_num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering_root.append(num)
    return num_id


def _numbering_id_exists(paragraph: Any, num_id: int) -> bool:
    modules = _docx_modules()
    qn = modules["qn"]

    try:
        numbering_root = paragraph.part.numbering_part.element
    except Exception:
        return False

    for num in numbering_root.findall(qn("w:num")):
        if num.get(qn("w:numId")) == str(num_id):
            return True
    return False


def _next_heading_prefix(document: Any, level: int) -> str:
    counters = getattr(document, "_codex_heading_counters", [])
    while len(counters) < level:
        counters.append(0)

    counters[level - 1] += 1
    for index in range(level, len(counters)):
        counters[index] = 0

    setattr(document, "_codex_heading_counters", counters)
    values = [str(value) for value in counters[:level] if value > 0]
    return ".".join(values) + "."


def _replace_paragraph_text(paragraph: Any, text: str) -> None:
    modules = _docx_modules()
    qn = modules["qn"]

    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.add_run(text)


def _format_paragraph_runs(
    paragraph: Any,
    *,
    bold: bool | None = None,
    preserve_color: bool,
) -> None:
    if not paragraph.runs:
        paragraph.add_run("")

    for run in paragraph.runs:
        _format_run(run, bold=bold, preserve_color=preserve_color)


def _format_run(run: Any, *, bold: bool | None, preserve_color: bool, font_size: Any = None) -> None:
    modules = _docx_modules()
    Pt = modules["Pt"]
    RGBColor = modules["RGBColor"]

    run.font.name = "Times New Roman"
    if font_size is not None:
        run.font.size = font_size
    else:
        run.font.size = Pt(13)
    if bold is not None:
        run.font.bold = bold
    if not preserve_color:
        run.font.color.rgb = RGBColor(0, 0, 0)


def _normalize_organization_name(organization: str) -> str:
    value = (organization or "").strip()
    if not value or value == "To chuc chua xac dinh":
        return "khách hàng"
    return value


def _metadata_text(value: Any, fallback: str) -> str:
    if value in (None, "", []):
        return fallback
    if isinstance(value, list):
        return "\n".join(str(item) for item in value if str(item).strip()) or fallback
    return str(value)


def _format_assessment_period(assessment_date: str | None) -> str:
    if not assessment_date:
        target = date.today()
        return f"Tháng {target.month:02d}/{target.year}"

    normalized = assessment_date.strip()
    for pattern in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y/%m/%d"):
        try:
            parsed = datetime.strptime(normalized, pattern)
            return f"Tháng {parsed.month:02d}/{parsed.year}"
        except ValueError:
            continue

    return normalized
