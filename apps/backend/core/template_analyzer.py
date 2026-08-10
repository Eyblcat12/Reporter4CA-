"""Template analyzer — inspect DOCX files for structure and tokens.

Security:
- Validates file magic bytes (PK\x03\x04 for ZIP/DOCX)
- Enforces max file size (20 MB)
- Sanitizes filenames (no path traversal)
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

MAX_TEMPLATE_SIZE = 20 * 1024 * 1024  # 20 MB
DOCX_MAGIC = b"PK\x03\x04"
KNOWN_TOKENS = [
    "{{TITLE}}",
    "{{ORGANIZATION}}",
    "{{ASSESSMENT_DATE}}",
    "{{ASSESSMENT_PERIOD}}",
]
TOKEN_PATTERN = re.compile(r"\{\{\s*[A-Za-z][A-Za-z0-9_]*\s*\}\}")

# ---------------------------------------------------------------------------
# Validation helpers
# ---------------------------------------------------------------------------


def validate_docx_bytes(data: bytes, max_size: int = MAX_TEMPLATE_SIZE) -> None:
    """Validate raw DOCX bytes.

    Raises ValueError on invalid input.
    """
    if len(data) > max_size:
        raise ValueError(
            f"File qua lon: {len(data) / 1024 / 1024:.1f} MB (gioi han {max_size / 1024 / 1024:.0f} MB)"
        )
    if not data[:4] == DOCX_MAGIC:
        raise ValueError("File khong phai dinh dang DOCX hop le (magic bytes khong khop)")


def sanitize_filename(filename: str) -> str:
    """Sanitize filename: remove path separators, special chars."""
    # Strip any directory components
    name = Path(filename).name
    # Remove dangerous characters
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", name)
    # Ensure .docx extension
    if not name.lower().endswith(".docx"):
        name = name + ".docx"
    # Prevent empty or dot-only names
    stem = Path(name).stem.strip(". ")
    if not stem:
        stem = "template"
    return stem + ".docx"


# ---------------------------------------------------------------------------
# Analysis
# ---------------------------------------------------------------------------


def analyze_template(file_path: Path, report_type: str = "full") -> dict[str, Any]:
    """Analyze a DOCX template and return structured metadata.

    Returns dict with:
        template_mode: 'cover' | 'full' | 'none'
        has_tokens: bool
        tokens_found: list[str]
        table_count: int
        heading_count: int
        heading_list: list[str]
        page_estimate: int
        styles_used: list[str]
        has_prototypes: bool
        prototype_tables: list[str]
    """
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed"}

    if not file_path.exists():
        return {"error": f"File not found: {file_path}"}

    doc = Document(str(file_path))

    # ── Collect all text for token search ───────────────────
    all_text = _collect_all_text(doc)

    # ── Find tokens ─────────────────────────────────────────
    tokens_found = [t for t in KNOWN_TOKENS if t in all_text]
    all_tokens = sorted(set(TOKEN_PATTERN.findall(all_text)))
    unknown_tokens = [token for token in all_tokens if token not in KNOWN_TOKENS]

    # ── Count headings ──────────────────────────────────────
    headings: list[dict[str, Any]] = []
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if style_name.startswith("Heading") and text:
            level = 0
            try:
                level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                level = 1
            headings.append({"text": text, "level": level})

    # ── Count tables ────────────────────────────────────────
    table_count = len(doc.tables)

    # ── Detect prototype tables ─────────────────────────────
    prototype_specs = {
        "inventory_server": ("STT", "Máy chủ", "Địa chỉ truy cập", "Phiên bản hệ điều hành"),
        "inventory_client": ("STT", "Máy trạm", "Địa chỉ truy cập", "Phiên bản hệ điều hành"),
        "summary_server": ("STT", "Máy chủ", "Kết quả rà soát đánh giá"),
        "summary_client": ("STT", "Máy trạm", "Kết quả rà soát đánh giá"),
        "detail": ("STT", "Hạng mục rà soát", "Kết quả rà soát"),
        "remediation_client": ("STT", "Máy trạm", "Địa chỉ IP", "Trạng thái"),
        "ioc": ("STT", "Thông tin", "Chi tiết"),
    }
    prototype_tables: list[str] = []
    for key, headers in prototype_specs.items():
        for table in doc.tables:
            if not table.rows:
                continue
            current = tuple(cell.text.strip() for cell in table.rows[0].cells)
            if current == headers:
                prototype_tables.append(key)
                break

    # ── Detect template mode ────────────────────────────────
    template_mode = _detect_mode(doc, headings, table_count)

    # ── Collect unique styles ───────────────────────────────
    styles_used = set()
    for para in doc.paragraphs:
        if para.style and para.style.name:
            styles_used.add(para.style.name)

    # ── Estimate page count ─────────────────────────────────
    total_paragraphs = len(doc.paragraphs)
    page_estimate = max(1, (total_paragraphs + table_count * 5) // 30)

    result = {
        "template_mode": template_mode,
        "has_tokens": len(tokens_found) > 0,
        "tokens_found": tokens_found,
        "unknown_tokens": unknown_tokens,
        "table_count": table_count,
        "heading_count": len(headings),
        "heading_list": [h["text"] for h in headings[:20]],
        "page_estimate": page_estimate,
        "styles_used": sorted(styles_used),
        "has_prototypes": len(prototype_tables) > 0,
        "prototype_tables": prototype_tables,
    }
    from core.template_schema import evaluate_template_compatibility

    result["compatibility"] = evaluate_template_compatibility(result, report_type)
    return result


def _collect_all_text(doc: Any) -> str:
    """Collect all text from paragraphs, tables, headers, footers."""
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for section in doc.sections:
        for para in section.header.paragraphs:
            parts.append(para.text)
        for para in section.footer.paragraphs:
            parts.append(para.text)
    return "\n".join(parts)


def _detect_mode(
    doc: Any,
    headings: list[dict[str, Any]],
    table_count: int,
) -> str:
    """Detect whether template is 'full', 'cover', or 'none'."""
    if not headings:
        # No headings at all
        if len(doc.paragraphs) < 5:
            return "none"
        return "cover"

    if table_count > 0 or len(headings) > 1:
        return "full"

    return "cover"
