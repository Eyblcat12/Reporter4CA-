from __future__ import annotations

import json
import os
import sys
import tempfile
import threading
import time
import unittest
from concurrent.futures import ThreadPoolExecutor
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "apps" / "backend"
TEMPLATE_ROOT = BACKEND / "templates"
GOLDEN_ROOT = ROOT / "tests" / "golden" / "docx-v1"
sys.path.insert(0, str(BACKEND))

from core import report_generator as generator  # noqa: E402
from core.config import prepared_template_enabled  # noqa: E402
from core.performance_metrics import PerformanceMetrics  # noqa: E402
from core.prepared_template import (  # noqa: E402
    PreparedTemplateCache,
    PreparedTemplateError,
)

from tests.test_docx_golden import document_snapshot, fixture_data  # noqa: E402

REPORT_TEMPLATE_PATHS = {
    "full": TEMPLATE_ROOT / "report_template.docx",
    "server_only": TEMPLATE_ROOT / "server_only" / "report_server_only_default.docx",
    "client_only": TEMPLATE_ROOT / "client_only" / "report_client_only_default.docx",
    "summary": TEMPLATE_ROOT / "summary" / "report_summary_default.docx",
    "technical": TEMPLATE_ROOT / "technical" / "report_technical_default.docx",
    "incident_response": TEMPLATE_ROOT / "report_template.docx",
}


def _docx_bytes(label: str) -> bytes:
    document = Document()
    document.add_heading(label, level=1)
    document.add_paragraph(f"Body {label}")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


class PreparedTemplateCacheTests(unittest.TestCase):
    def test_cache_miss_hit_restart_and_source_immutability(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.docx"
            original = _docx_bytes("first")
            source.write_bytes(original)
            calls = 0

            def compiler(payload: bytes):
                nonlocal calls
                calls += 1
                return payload, {"templateMode": "cover"}

            cache = PreparedTemplateCache(root / "cache")
            first = cache.get_or_compile(source, "full", compiler)
            second = cache.get_or_compile(source, "full", compiler)
            restarted = PreparedTemplateCache(root / "cache").get_or_compile(
                source, "full", compiler
            )

            self.assertFalse(first.cache_hit)
            self.assertTrue(second.cache_hit)
            self.assertTrue(restarted.cache_hit)
            self.assertEqual(calls, 1)
            self.assertEqual(source.read_bytes(), original)
            self.assertEqual(first.key, second.key)
            self.assertEqual(first.path, restarted.path)

    def test_same_filename_changed_content_gets_a_new_key(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "same-name.docx"
            cache = PreparedTemplateCache(root / "cache")
            source.write_bytes(_docx_bytes("v1"))
            first = cache.get_or_compile(
                source, "full", lambda payload: (payload, {"templateMode": "cover"})
            )
            source.write_bytes(_docx_bytes("v2"))
            second = cache.get_or_compile(
                source, "full", lambda payload: (payload, {"templateMode": "cover"})
            )
            self.assertNotEqual(first.key, second.key)
            self.assertFalse(second.cache_hit)

    def test_corrupt_artifact_is_recompiled(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.docx"
            source.write_bytes(_docx_bytes("corrupt"))
            calls = 0

            def compiler(payload: bytes):
                nonlocal calls
                calls += 1
                return payload, {"templateMode": "cover"}

            cache = PreparedTemplateCache(root / "cache")
            first = cache.get_or_compile(source, "full", compiler)
            first.path.write_bytes(b"not-a-docx")
            repaired = cache.get_or_compile(source, "full", compiler)
            self.assertFalse(repaired.cache_hit)
            self.assertEqual(calls, 2)
            Document(str(repaired.path))

    def test_concurrent_same_key_compiles_once(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.docx"
            source.write_bytes(_docx_bytes("concurrent"))
            calls = 0
            calls_lock = threading.Lock()

            def compiler(payload: bytes):
                nonlocal calls
                with calls_lock:
                    calls += 1
                time.sleep(0.05)
                return payload, {"templateMode": "cover"}

            cache = PreparedTemplateCache(root / "cache")
            with ThreadPoolExecutor(max_workers=4) as executor:
                results = list(
                    executor.map(
                        lambda _index: cache.get_or_compile(source, "full", compiler),
                        range(4),
                    )
                )
            self.assertEqual(calls, 1)
            self.assertEqual(sum(not item.cache_hit for item in results), 1)
            self.assertEqual(len({item.key for item in results}), 1)

    def test_lru_and_startup_temporary_cleanup_are_bounded(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            cache_root = root / "cache"
            cache_root.mkdir()
            temporary = cache_root / ".orphan.tmp"
            temporary.write_text("partial", encoding="utf-8")
            cache = PreparedTemplateCache(cache_root, max_entries=1)
            self.assertFalse(temporary.exists())

            first_source = root / "first.docx"
            second_source = root / "second.docx"
            first_source.write_bytes(_docx_bytes("first"))
            second_source.write_bytes(_docx_bytes("second"))
            first = cache.get_or_compile(
                first_source, "full", lambda payload: (payload, {"templateMode": "cover"})
            )
            second = cache.get_or_compile(
                second_source, "full", lambda payload: (payload, {"templateMode": "cover"})
            )
            self.assertFalse(first.path.exists())
            self.assertTrue(second.path.exists())

    def test_startup_removes_stale_content_addressed_entry_only(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            cache_root = root / "cache"
            stale = cache_root / ("a" * 64)
            unrelated = cache_root / "user-note"
            stale.mkdir(parents=True)
            unrelated.mkdir()
            (stale / "manifest.json").write_text("{}", encoding="utf-8")
            (stale / "template.docx").write_bytes(b"invalid")
            (unrelated / "keep.txt").write_text("keep", encoding="utf-8")

            PreparedTemplateCache(cache_root)

            self.assertFalse(stale.exists())
            self.assertTrue((unrelated / "keep.txt").is_file())

    def test_invalid_compiler_output_never_becomes_a_cache_entry(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.docx"
            source.write_bytes(_docx_bytes("invalid"))
            cache = PreparedTemplateCache(root / "cache")
            with self.assertRaises(PreparedTemplateError):
                cache.get_or_compile(
                    source,
                    "full",
                    lambda _payload: (b"invalid", {"templateMode": "full"}),
                )
            self.assertFalse(any((root / "cache").glob("*/template.docx")))


class PreparedTemplateIntegrationTests(unittest.TestCase):
    def test_bundled_warmup_resolves_categories_and_defers_individual_failures(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            fallback = root / "report_template.docx"
            fallback.write_bytes(_docx_bytes("fallback"))
            server_dir = root / "server_only"
            server_dir.mkdir()
            server = server_dir / "server.docx"
            server.write_bytes(_docx_bytes("server"))

            def warm(source: Path, report_type: generator.ReportType) -> bool:
                if report_type == generator.ReportType.TECHNICAL:
                    raise PreparedTemplateError("controlled")
                self.assertTrue(source.is_file())
                return True

            with patch.object(generator, "warm_prepared_template", side_effect=warm) as called:
                results = generator.warm_bundled_templates(root)

            self.assertEqual(len(results), len(generator.ReportType))
            self.assertEqual(
                {item["reportType"] for item in results},
                {item.value for item in generator.ReportType},
            )
            technical = next(item for item in results if item["reportType"] == "technical")
            self.assertEqual(technical["outcome"], "deferred")
            self.assertEqual(technical["errorCode"], "PreparedTemplateError")
            server_call = next(
                call
                for call in called.call_args_list
                if call.args[1] == generator.ReportType.SERVER_ONLY
            )
            self.assertTrue(os.path.samefile(server_call.args[0], server))

    def test_flag_defaults_on_and_supports_explicit_rollback(self) -> None:
        with patch.dict(os.environ, {}, clear=True):
            self.assertTrue(prepared_template_enabled())
        for value in ("1", "true", "YES", "on"):
            with (
                self.subTest(value=value),
                patch.dict(os.environ, {"AUTO_REPORT_PREPARED_TEMPLATE": value}),
            ):
                self.assertTrue(prepared_template_enabled())
        for value in ("0", "false", "no", "off", ""):
            with (
                self.subTest(value=value),
                patch.dict(os.environ, {"AUTO_REPORT_PREPARED_TEMPLATE": value}),
            ):
                self.assertFalse(prepared_template_enabled())

    def test_cache_hit_matches_legacy_and_removes_trim_from_hot_path(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            cache = PreparedTemplateCache(Path(directory) / "cache")
            template = REPORT_TEMPLATE_PATHS["full"]
            source_hash_before = template.read_bytes()
            with patch.object(generator, "prepared_template_enabled", return_value=False):
                legacy = generator.generate_report(
                    fixture_data(),
                    title="Prepared",
                    organization="Reporter Team",
                    assessment_date="2026-07-20",
                    template_path=template,
                    report_type="full",
                )
            with (
                patch.object(generator, "prepared_template_enabled", return_value=True),
                patch.object(generator, "_get_prepared_template_cache", return_value=cache),
            ):
                generator.generate_report(
                    fixture_data(),
                    title="Prepared",
                    organization="Reporter Team",
                    assessment_date="2026-07-20",
                    template_path=template,
                    report_type="full",
                )
                metrics = PerformanceMetrics()
                prepared = generator.generate_report(
                    fixture_data(),
                    title="Prepared",
                    organization="Reporter Team",
                    assessment_date="2026-07-20",
                    template_path=template,
                    report_type="full",
                    metrics=metrics,
                )
            self.assertEqual(document_snapshot(legacy), document_snapshot(prepared))
            phases = [item["name"] for item in metrics.public()["phases"]]
            self.assertIn("preparedTemplate", phases)
            self.assertNotIn("templateTrim", phases)
            self.assertEqual(
                metrics.public()["metadata"].get("cacheState"), "cache-warm/prepared-hit"
            )
            self.assertEqual(template.read_bytes(), source_hash_before)

    def test_six_report_types_match_golden_with_prepared_cache(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            cache = PreparedTemplateCache(Path(directory) / "cache")
            with (
                patch.object(generator, "prepared_template_enabled", return_value=True),
                patch.object(generator, "_get_prepared_template_cache", return_value=cache),
            ):
                for report_type, template in REPORT_TEMPLATE_PATHS.items():
                    with self.subTest(report_type=report_type):
                        document = generator.generate_report(
                            fixture_data(),
                            title="Golden Report",
                            organization="Reporter Team",
                            assessment_date="2026-07-20",
                            template_path=template,
                            report_type=report_type,
                        )
                        expected = json.loads(
                            (GOLDEN_ROOT / f"{report_type}.json").read_text(encoding="utf-8")
                        )
                        self.assertEqual(
                            document_snapshot(
                                document,
                                getattr(document, "_reporter_manifest", {}),
                            ),
                            expected,
                        )

    def test_cache_failure_falls_back_without_exposing_source_path(self) -> None:
        class BrokenCache:
            def get_or_compile(self, *_args, **_kwargs):
                raise PreparedTemplateError("C:/secret/customer/template.docx")

        template = REPORT_TEMPLATE_PATHS["summary"]
        with (
            self.assertLogs("core.report_generator", level="WARNING") as logs,
            patch.object(generator, "prepared_template_enabled", return_value=True),
            patch.object(generator, "_get_prepared_template_cache", return_value=BrokenCache()),
        ):
            document = generator.generate_report(
                fixture_data(),
                title="Fallback",
                organization="Reporter Team",
                assessment_date="2026-07-20",
                template_path=template,
                report_type="summary",
            )
        self.assertTrue(document._reporter_integrity["valid"])
        self.assertNotIn("C:/secret", "\n".join(logs.output))


if __name__ == "__main__":
    unittest.main()
