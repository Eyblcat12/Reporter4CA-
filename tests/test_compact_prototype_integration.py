from __future__ import annotations

import os
import sys
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "apps" / "backend"
TEMPLATE_ROOT = BACKEND / "templates"
sys.path.insert(0, str(BACKEND))

from core import report_generator as generator  # noqa: E402
from core.config import compact_prototype_enabled  # noqa: E402
from tests.test_docx_golden import document_snapshot  # noqa: E402


INVENTORY_HEADERS = [
    "STT",
    "Máy chủ",
    "Địa chỉ truy cập",
    "Phiên bản hệ điều hành",
]
OUTPUT_ROWS = [
    ["1", "SRV-01", "10.0.0.1", "Windows Server 2022"],
    ["2", "SRV-02", "10.0.0.2", "Ubuntu 24.04"],
]
REPORT_TEMPLATE_PATHS = {
    "full": TEMPLATE_ROOT / "report_template.docx",
    "server_only": (
        TEMPLATE_ROOT / "server_only" / "report_server_only_default.docx"
    ),
    "client_only": (
        TEMPLATE_ROOT / "client_only" / "report_client_only_default.docx"
    ),
    "summary": TEMPLATE_ROOT / "summary" / "report_summary_default.docx",
    "technical": (
        TEMPLATE_ROOT / "technical" / "report_technical_default.docx"
    ),
    "incident_response": TEMPLATE_ROOT / "report_template.docx",
}


def _append_property(parent, tag: str, attributes: dict[str, str] | None = None):
    element = OxmlElement(tag)
    for name, value in (attributes or {}).items():
        element.set(qn(name), value)
    parent.append(element)
    return element


def _prototype_document(
    *,
    data_rows: int = 1,
    numbered_stt: bool = False,
) -> tuple[Document, object]:
    document = Document()
    table = document.add_table(rows=1 + data_rows, cols=len(INVENTORY_HEADERS))
    table.style = "Table Grid"
    for index, value in enumerate(INVENTORY_HEADERS):
        table.rows[0].cells[index].text = value
        table.rows[0].cells[index].paragraphs[0].runs[0].bold = True
    for row_index, row in enumerate(table.rows[1:], start=1):
        values = [
            str(row_index),
            f"PROTOTYPE-{row_index}",
            f"192.0.2.{row_index}",
            "Prototype OS",
        ]
        for column_index, value in enumerate(values):
            row.cells[column_index].text = value

    shading = _append_property(
        table.rows[1].cells[1]._tc.get_or_add_tcPr(),
        "w:shd",
        {"w:fill": "EAF2F8"},
    )
    self_closing_fill = shading.get(qn("w:fill"))
    if self_closing_fill != "EAF2F8":  # pragma: no cover - defensive OOXML check
        raise AssertionError("Failed to apply prototype shading.")

    if numbered_stt:
        paragraph = table.rows[1].cells[0].paragraphs[0]
        paragraph.style = "List Paragraph"
        num_pr = _append_property(paragraph._p.get_or_add_pPr(), "w:numPr")
        _append_property(num_pr, "w:numId", {"w:val": "1"})
    return document, table


def _capture(document: Document, *, enabled: bool) -> None:
    with patch.object(
        generator,
        "compact_prototype_enabled",
        return_value=enabled,
    ):
        generator._capture_template_prototypes(document)


def _fixture_data() -> dict:
    return {
        "servers": [{
            "hostname": "SRV-COMPACT-01",
            "ip": "10.10.0.10",
            "os": "Windows Server 2022",
            "result": "Malware detected",
            "notes": "SHA256 evidence linked to T1055",
        }],
        "clients": [{
            "hostname": "WS-COMPACT-01",
            "ip": "10.10.0.20",
            "os": "Windows 11",
            "result": "No finding",
            "notes": "Validated by endpoint telemetry",
        }],
        "metadata": {
            "incidentName": "Compact prototype parity",
            "severity": "high",
            "timeline": [{
                "time": "2026-07-20T10:00:00Z",
                "event": "Detection",
                "evidence": "EDR-001",
            }],
        },
    }


class CompactPrototypeConfigTests(unittest.TestCase):
    def test_feature_flag_defaults_off_and_accepts_only_truthy_values(self) -> None:
        with patch.dict(os.environ, {}, clear=True):
            self.assertFalse(compact_prototype_enabled())

        for raw_value in ("1", "true", "TRUE", " yes ", "on"):
            with self.subTest(value=raw_value):
                with patch.dict(
                    os.environ,
                    {"AUTO_REPORT_COMPACT_PROTOTYPE": raw_value},
                ):
                    self.assertTrue(compact_prototype_enabled())

        for raw_value in ("", "0", "false", "disabled", "2"):
            with self.subTest(value=raw_value):
                with patch.dict(
                    os.environ,
                    {"AUTO_REPORT_COMPACT_PROTOTYPE": raw_value},
                ):
                    self.assertFalse(compact_prototype_enabled())


class CompactPrototypeIntegrationTests(unittest.TestCase):
    def test_flag_off_captures_legacy_without_calling_compiler(self) -> None:
        document, _ = _prototype_document()

        with patch.object(generator, "compile_table_blueprint") as compiler:
            _capture(document, enabled=False)

        compiler.assert_not_called()
        self.assertIn(
            "inventory_server",
            getattr(document, "_codex_table_prototypes"),
        )
        self.assertEqual(
            getattr(document, "_reporter_table_blueprints"),
            {},
        )
        self.assertFalse(document._reporter_compact_prototype_enabled)

    def test_flag_on_safe_prototype_uses_detached_fast_path(self) -> None:
        document, _ = _prototype_document()
        _capture(document, enabled=True)

        self.assertIn(
            "inventory_server",
            document._reporter_table_blueprints,
        )
        self.assertEqual(
            document._reporter_table_blueprints[
                "inventory_server"
            ].data_row_variant_count,
            1,
        )

        with patch.object(
            generator,
            "_create_table_from_prototype",
            wraps=generator._create_table_from_prototype,
        ) as legacy_builder:
            generated = generator._create_table(
                document,
                INVENTORY_HEADERS,
                OUTPUT_ROWS,
                prototype_key="inventory_server",
                column_widths_mm=[12, 45, 40, 63],
            )

        legacy_builder.assert_not_called()
        self.assertIs(generated._tbl.getparent(), document._body._element)
        self.assertEqual(len(document.tables), 2)
        self.assertEqual(
            [[cell.text for cell in row.cells] for row in generated.rows],
            [INVENTORY_HEADERS, *OUTPUT_ROWS],
        )

    def test_unsafe_and_multi_variant_prototypes_fall_back_once(self) -> None:
        cases: dict[str, Document] = {}

        unsafe_document, unsafe_table = _prototype_document()
        unsafe_table.rows[1].cells[1].add_paragraph("Second paragraph")
        cases["unsafe"] = unsafe_document

        variant_document, variant_table = _prototype_document(data_rows=2)
        _append_property(
            variant_table.rows[2].cells[1]._tc.get_or_add_tcPr(),
            "w:shd",
            {"w:fill": "FFF2CC"},
        )
        cases["multiple_variants"] = variant_document

        for name, document in cases.items():
            with self.subTest(case=name):
                _capture(document, enabled=True)
                with patch.object(
                    generator,
                    "_create_table_from_blueprint",
                    wraps=generator._create_table_from_blueprint,
                ) as fast_builder, patch.object(
                    generator,
                    "_create_table_from_prototype",
                    wraps=generator._create_table_from_prototype,
                ) as legacy_builder:
                    generator._create_table(
                        document,
                        INVENTORY_HEADERS,
                        OUTPUT_ROWS,
                        prototype_key="inventory_server",
                    )

                fast_builder.assert_not_called()
                legacy_builder.assert_called_once()
                self.assertEqual(len(document.tables), 2)

    def test_compile_error_keeps_legacy_and_does_not_log_template_data(self) -> None:
        document, _ = _prototype_document()
        secret = "SECRET-HOST-FROM-TEMPLATE"

        with self.assertLogs("core.report_generator", level="WARNING") as logs:
            with patch.object(
                generator,
                "compile_table_blueprint",
                side_effect=RuntimeError(secret),
            ):
                _capture(document, enabled=True)

        self.assertNotIn(secret, "\n".join(logs.output))
        self.assertIn(
            "inventory_server",
            document._codex_table_prototypes,
        )
        self.assertNotIn(
            "inventory_server",
            document._reporter_table_blueprints,
        )

        with patch.object(
            generator,
            "_create_table_from_prototype",
            wraps=generator._create_table_from_prototype,
        ) as legacy_builder:
            generator._create_table(
                document,
                INVENTORY_HEADERS,
                OUTPUT_ROWS,
                prototype_key="inventory_server",
            )
        legacy_builder.assert_called_once()

    def test_dimension_mismatch_never_enters_fast_path(self) -> None:
        cases = (
            (INVENTORY_HEADERS[:-1], [row[:-1] for row in OUTPUT_ROWS]),
            (INVENTORY_HEADERS, [OUTPUT_ROWS[0][:-1]]),
        )
        for headers, rows in cases:
            with self.subTest(columns=len(headers), row_width=len(rows[0])):
                document, _ = _prototype_document()
                _capture(document, enabled=True)
                with patch.object(
                    generator,
                    "_create_table_from_blueprint",
                    wraps=generator._create_table_from_blueprint,
                ) as fast_builder, patch.object(
                    generator,
                    "_create_table_from_prototype",
                    wraps=generator._create_table_from_prototype,
                ) as legacy_builder:
                    generator._create_table(
                        document,
                        headers,
                        rows,
                        prototype_key="inventory_server",
                    )

                fast_builder.assert_not_called()
                legacy_builder.assert_called_once()
                self.assertEqual(len(document.tables), 2)

    def test_fast_build_error_falls_back_once_without_dangling_table(self) -> None:
        document, _ = _prototype_document()
        _capture(document, enabled=True)
        initial_table_count = len(document.tables)
        real_set_layout = generator._set_table_layout
        layout_calls = 0
        secret = "SECRET-ASSET-FROM-ROW"

        def fail_first_layout(table, widths):
            nonlocal layout_calls
            layout_calls += 1
            if layout_calls == 1:
                raise RuntimeError(secret)
            return real_set_layout(table, widths)

        with self.assertLogs("core.report_generator", level="WARNING") as logs:
            with patch.object(
                generator,
                "_set_table_layout",
                side_effect=fail_first_layout,
            ), patch.object(
                generator,
                "_create_table_from_prototype",
                wraps=generator._create_table_from_prototype,
            ) as legacy_builder:
                generated = generator._create_table(
                    document,
                    INVENTORY_HEADERS,
                    OUTPUT_ROWS,
                    prototype_key="inventory_server",
                    column_widths_mm=[12, 45, 40, 63],
                )

        legacy_builder.assert_called_once()
        self.assertEqual(layout_calls, 2)
        self.assertEqual(len(document.tables), initial_table_count + 1)
        self.assertIs(generated._tbl.getparent(), document._body._element)
        self.assertNotIn(secret, "\n".join(logs.output))

    def test_fast_path_strips_stt_numbering_above_one_thousand_rows(self) -> None:
        document, _ = _prototype_document(numbered_stt=True)
        _capture(document, enabled=True)
        rows = [
            [
                str(index),
                f"SRV-{index:04d}",
                f"10.0.{index // 256}.{index % 256}",
                "Linux",
            ]
            for index in range(1, 1002)
        ]

        with patch.object(
            generator,
            "_create_table_from_prototype",
            wraps=generator._create_table_from_prototype,
        ) as legacy_builder:
            generated = generator._create_table(
                document,
                INVENTORY_HEADERS,
                rows,
                prototype_key="inventory_server",
            )

        legacy_builder.assert_not_called()
        checkpoints = (1, 2, 9, 10, 99, 100, 999, 1000, 1001)
        self.assertEqual(
            [generated.rows[index].cells[0].text for index in checkpoints],
            [str(index) for index in checkpoints],
        )
        for index in checkpoints:
            p_pr = generated.rows[index].cells[0].paragraphs[0]._p.pPr
            self.assertIsNotNone(p_pr)
            self.assertIsNone(p_pr.find(qn("w:numPr")))
            p_style = p_pr.find(qn("w:pStyle"))
            self.assertTrue(
                p_style is None
                or p_style.get(qn("w:val")) != "ListParagraph"
            )

    def test_compact_and_legacy_table_snapshots_are_identical(self) -> None:
        def build(enabled: bool) -> dict:
            document, _ = _prototype_document()
            _capture(document, enabled=enabled)
            generator._create_table(
                document,
                INVENTORY_HEADERS,
                OUTPUT_ROWS,
                prototype_key="inventory_server",
                column_widths_mm=[12, 45, 40, 63],
            )
            return document_snapshot(document)["tables"][-1]

        self.assertEqual(build(False), build(True))

    def test_table_cells_and_borders_are_formatted_once(self) -> None:
        for enabled in (False, True):
            with self.subTest(compact=enabled):
                document, _ = _prototype_document()
                _capture(document, enabled=enabled)
                with patch.object(
                    generator,
                    "_ensure_table_borders",
                    wraps=generator._ensure_table_borders,
                ) as border_formatter, patch.object(
                    generator,
                    "_format_cell",
                    wraps=generator._format_cell,
                ) as cell_formatter:
                    generated = generator._create_table(
                        document,
                        INVENTORY_HEADERS,
                        OUTPUT_ROWS,
                        prototype_key="inventory_server",
                    )
                    generator._style_table(generated)

                self.assertEqual(border_formatter.call_count, 1)
                self.assertEqual(
                    cell_formatter.call_count,
                    len(generated.rows) * len(INVENTORY_HEADERS),
                )

    def test_six_report_types_match_legacy_structural_golden_snapshot(self) -> None:
        fast_calls = 0
        real_fast_builder = generator._create_table_from_blueprint

        def count_fast_calls(*args, **kwargs):
            nonlocal fast_calls
            fast_calls += 1
            return real_fast_builder(*args, **kwargs)

        for report_type, template_path in REPORT_TEMPLATE_PATHS.items():
            with self.subTest(report_type=report_type):
                with patch.dict(
                    os.environ,
                    {"AUTO_REPORT_COMPACT_PROTOTYPE": "0"},
                ):
                    legacy = generator.generate_report(
                        _fixture_data(),
                        title="Compact Prototype Golden",
                        organization="Reporter Team",
                        assessment_date="2026-07-20",
                        template_path=template_path,
                        report_type=report_type,
                    )
                with patch.dict(
                    os.environ,
                    {"AUTO_REPORT_COMPACT_PROTOTYPE": "1"},
                ), patch.object(
                    generator,
                    "_create_table_from_blueprint",
                    side_effect=count_fast_calls,
                ):
                    compact = generator.generate_report(
                        _fixture_data(),
                        title="Compact Prototype Golden",
                        organization="Reporter Team",
                        assessment_date="2026-07-20",
                        template_path=template_path,
                        report_type=report_type,
                    )

                self.assertEqual(
                    document_snapshot(
                        legacy,
                        getattr(legacy, "_reporter_manifest", {}),
                    ),
                    document_snapshot(
                        compact,
                        getattr(compact, "_reporter_manifest", {}),
                    ),
                )

        self.assertGreater(fast_calls, 0)


if __name__ == "__main__":
    unittest.main()
