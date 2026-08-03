from __future__ import annotations

import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "apps" / "backend"
sys.path.insert(0, str(BACKEND))

from core.report_generator import (  # noqa: E402
    DEFAULT_RESULT_TEXT,
    REPORTER_HEADING_NUMBERING_MARKER,
    REPORTER_HEADING_TEXT_LEFT_TWIPS,
    ReportType,
    _create_table_from_prototype,
    _get_anomaly_text,
    generate_report,
)


class ReportGeneratorTests(unittest.TestCase):
    def setUp(self) -> None:
        self.template = BACKEND / "templates" / "report_template.docx"
        self.data = {
            "servers": [
                {
                    "hostname": "SRV-01",
                    "ip": "10.0.0.1",
                    "os": "Windows Server 2022",
                    "result": "Phát hiện mã độc",
                    "notes": "PlugX DLL sideloading qua svchost.exe",
                }
            ],
            "clients": [
                {
                    "hostname": "PC-01",
                    "ip": "10.0.0.2",
                    "os": "Windows 11",
                    "result": "Không phát hiện",
                    "notes": "",
                }
            ],
            "metadata": {},
        }

    def build(self, report_type: ReportType):
        return generate_report(
            self.data,
            title="Test report",
            organization="Test organization",
            template_path=self.template,
            report_type=report_type,
        )

    @staticmethod
    def all_text(document: Document) -> str:
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join(paragraphs + cells)

    def test_server_only_bundled_template_generates_only_server_assets(self) -> None:
        document = generate_report(
            self.data,
            title="Server-only template test",
            organization="Test organization",
            template_path=BACKEND / "templates" / "server_only" / "report_server_only_default.docx",
            report_type=ReportType.SERVER_ONLY,
        )
        text = self.all_text(document)
        self.assertIn("SRV-01", text)
        self.assertNotIn("PC-01", text)

    def test_client_only_bundled_template_generates_only_client_assets(self) -> None:
        document = generate_report(
            self.data,
            title="Client-only template test",
            organization="Test organization",
            template_path=BACKEND / "templates" / "client_only" / "report_client_only_default.docx",
            report_type=ReportType.CLIENT_ONLY,
        )
        text = self.all_text(document)
        self.assertIn("PC-01", text)
        self.assertNotIn("SRV-01", text)

    def test_server_only_omits_client_sections(self) -> None:
        document = self.build(ReportType.SERVER_ONLY)
        text = "\n".join(p.text for p in document.paragraphs)
        self.assertIn("SRV-01", text)
        self.assertNotIn("PC-01", text)
        self.assertNotIn("máy trạm", text.lower())

    def test_client_only_omits_server_sections(self) -> None:
        document = self.build(ReportType.CLIENT_ONLY)
        text = "\n".join(p.text for p in document.paragraphs)
        self.assertIn("PC-01", text)
        self.assertNotIn("SRV-01", text)
        self.assertNotIn("máy chủ", text.lower())

    def test_summary_uses_concise_three_part_structure(self) -> None:
        document = self.build(ReportType.SUMMARY)
        heading_1 = [p.text for p in document.paragraphs if p.style.name == "Heading 1"]
        heading_2 = [p.text for p in document.paragraphs if p.style.name == "Heading 2"]
        heading_3 = [p.text for p in document.paragraphs if p.style.name == "Heading 3"]

        self.assertEqual(
            heading_1,
            ["Tổng quan", "Kết quả và phân tích tổng hợp", "Kết luận và khuyến nghị"],
        )
        self.assertIn("Các phát hiện đáng chú ý", heading_2)
        self.assertIn("Đánh giá ảnh hưởng chung", heading_2)
        self.assertEqual(heading_3, [])
        self.assertNotIn("Phân tích máy chủ", heading_2)

    def test_summary_uses_standardized_rule_assessment_text(self) -> None:
        document = self.build(ReportType.SUMMARY)
        cell_text = "\n".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
        self.assertIn("Ghi nhận dấu hiệu bất thường", cell_text)
        self.assertIn("Không phát hiện dấu hiệu bất thường", cell_text)

    def test_technical_uses_detailed_three_part_structure(self) -> None:
        document = self.build(ReportType.TECHNICAL)
        heading_1 = [p.text for p in document.paragraphs if p.style.name == "Heading 1"]
        heading_2 = [p.text for p in document.paragraphs if p.style.name == "Heading 2"]
        heading_3 = [p.text for p in document.paragraphs if p.style.name == "Heading 3"]

        self.assertEqual(
            heading_1,
            ["Tổng quan", "Phân tích chi tiết", "Kết luận và khuyến nghị"],
        )
        self.assertIn("Phân tích điều tra", heading_2)
        self.assertIn("Indicators of Compromise (IoCs)", heading_2)
        self.assertIn("Hành động xử lý và gỡ bỏ", heading_2)
        self.assertIn("SRV-01", heading_3)
        self.assertIn("PC-01", heading_3)
        self.assertIn("Tệp tin đáng ngờ", heading_3)
        self.assertIn("Dịch vụ và tiến trình", heading_3)

    def test_template_heading_has_no_direct_format_override(self) -> None:
        document = self.build(ReportType.SERVER_ONLY)
        heading = next(p for p in document.paragraphs if p.style.name == "Heading 1" and p.text)
        self.assertIsNone(heading.alignment)
        for run in heading.runs:
            self.assertIsNone(run.font.name)
            self.assertIsNone(run.font.size)

    def test_generated_headings_use_aligned_reporter_numbering_from_one(self) -> None:
        document = self.build(ReportType.FULL)
        headings = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.style.name in {"Heading 1", "Heading 2", "Heading 3"}
        ]
        self.assertTrue(headings)

        num_ids = {
            paragraph._p.find("./" + qn("w:pPr") + "/" + qn("w:numPr") + "/" + qn("w:numId")).get(qn("w:val"))
            for paragraph in headings
        }
        self.assertEqual(len(num_ids), 1)
        num_id = num_ids.pop()

        numbering_root = document.part.numbering_part.element
        num = next(
            item
            for item in numbering_root.findall(qn("w:num"))
            if item.get(qn("w:numId")) == num_id
        )
        abstract_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
        abstract_num = next(
            item
            for item in numbering_root.findall(qn("w:abstractNum"))
            if item.get(qn("w:abstractNumId")) == abstract_id
        )
        self.assertEqual(
            abstract_num.find(qn("w:tmpl")).get(qn("w:val")),
            REPORTER_HEADING_NUMBERING_MARKER,
        )

        expected_text = {0: "%1.", 1: "%1.%2.", 2: "%1.%2.%3."}
        for level in abstract_num.findall(qn("w:lvl")):
            level_index = int(level.get(qn("w:ilvl")))
            self.assertEqual(level.find(qn("w:start")).get(qn("w:val")), "1")
            self.assertEqual(level.find(qn("w:lvlText")).get(qn("w:val")), expected_text[level_index])
            indentation = level.find("./" + qn("w:pPr") + "/" + qn("w:ind"))
            self.assertEqual(indentation.get(qn("w:left")), REPORTER_HEADING_TEXT_LEFT_TWIPS)
            self.assertEqual(indentation.get(qn("w:hanging")), REPORTER_HEADING_TEXT_LEFT_TWIPS)

        levels = {
            paragraph.style.name: paragraph._p.find(
                "./" + qn("w:pPr") + "/" + qn("w:numPr") + "/" + qn("w:ilvl")
            ).get(qn("w:val"))
            for paragraph in headings
        }
        self.assertEqual(levels["Heading 1"], "0")
        self.assertEqual(levels["Heading 2"], "1")
        self.assertEqual(levels["Heading 3"], "2")

    def test_template_toc_cache_is_removed_before_generation(self) -> None:
        document = self.build(ReportType.SERVER_ONLY)
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "report.docx"
            document.save(path)
            with zipfile.ZipFile(path) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")

        self.assertIn('TOC \\o "1-3"', document_xml)
        self.assertNotIn("PAGEREF _Toc", document_xml)
        self.assertIn('w:fldCharType="begin" w:dirty="true"', document_xml)

    def test_inventory_sequence_is_written_once(self) -> None:
        document = self.build(ReportType.SERVER_ONLY)
        inventory = next(table for table in document.tables if table.rows[0].cells[0].text == "STT")
        self.assertEqual(inventory.rows[1].cells[0].text, "1")
        self.assertNotEqual(inventory.rows[1].cells[0].text, "11")

    def test_inventory_sequence_remains_correct_above_one_thousand_rows(self) -> None:
        document = Document()
        prototype = document.add_table(rows=2, cols=4)
        prototype.rows[0].cells[0].text = "STT"

        stt_paragraph = prototype.rows[1].cells[0].paragraphs[0]
        stt_paragraph.style = "List Paragraph"
        num_pr = OxmlElement("w:numPr")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "1")
        num_pr.append(num_id)
        stt_paragraph._p.get_or_add_pPr().append(num_pr)

        generated = _create_table_from_prototype(
            document,
            prototype._tbl,
            ["STT", "Máy chủ", "Địa chỉ truy cập", "Phiên bản hệ điều hành"],
            [
                [str(index), f"SRV-{index:04d}", f"10.0.{index // 256}.{index % 256}", "Linux"]
                for index in range(1, 1002)
            ],
        )

        checkpoints = (1, 2, 9, 10, 99, 100, 999, 1000, 1001)
        self.assertEqual(
            [generated.rows[index].cells[0].text for index in checkpoints],
            [str(index) for index in checkpoints],
        )
        for index in checkpoints:
            paragraph = generated.rows[index].cells[0].paragraphs[0]
            self.assertIsNone(paragraph._p.find(".//" + qn("w:numPr")))
            self.assertIsNone(paragraph._p.find(".//" + qn("w:pStyle")))

    def test_inventory_removes_template_automatic_numbering(self) -> None:
        numbered_template = BACKEND / "templates" / "report_da_chuan_hoa.docx"
        numbered_data = {
            **self.data,
            "servers": [
                {
                    "hostname": f"SRV-{index:02d}",
                    "ip": f"10.0.0.{index}",
                    "os": "Windows Server 2022",
                    "result": "Không phát hiện",
                    "notes": "",
                }
                for index in range(1, 13)
            ],
            "clients": [],
        }
        document = generate_report(
            numbered_data,
            title="Numbered prototype regression",
            organization="Test organization",
            template_path=numbered_template,
            report_type=ReportType.SERVER_ONLY,
        )
        inventory = next(
            table for table in document.tables
            if [cell.text.strip() for cell in table.rows[0].cells[:2]] == ["STT", "Máy chủ"]
        )
        stt_paragraphs = [row.cells[0].paragraphs[0] for row in inventory.rows[1:13]]
        # This template deliberately uses Word numbering. Generated STT cells
        # must be renderer-independent plain text, with no inherited list
        # numbering or ListParagraph style left to duplicate/hide the value.
        self.assertEqual([paragraph.text for paragraph in stt_paragraphs], [str(index) for index in range(1, 13)])
        self.assertTrue(
            all(paragraph._p.find(".//" + qn("w:numPr")) is None for paragraph in stt_paragraphs)
        )
        self.assertTrue(
            all(paragraph._p.find(".//" + qn("w:pStyle")) is None for paragraph in stt_paragraphs)
        )

    def test_anomaly_only_marks_evidence_backed_rows(self) -> None:
        asset = self.data["servers"][0]
        self.assertEqual(_get_anomaly_text("Kiểm tra rootkit", asset), DEFAULT_RESULT_TEXT)
        self.assertIn("Phát hiện", _get_anomaly_text("Xác định các Service, process, loaded DLL bất thường", asset))
        self.assertIn("Phát hiện", _get_anomaly_text("Xác định, phân tích các tệp tin bất thường", asset))

    def test_incident_response_has_dedicated_sections(self) -> None:
        self.data["metadata"] = {
            "incident_id": "IR-001",
            "severity": "High",
            "timeline": [{"time": "10:00", "event": "Detected", "evidence": "EDR-001", "relatedIocs": "10.0.0.1"}],
            "iocs": [{"type": "ip", "value": "10.0.0.1", "source": "EDR-001"}],
            "containmentActions": [{"action": "Isolate host", "status": "Done", "owner": "SOC", "evidence": "EDR-002"}],
        }
        document = self.build(ReportType.INCIDENT_RESPONSE)
        headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
        self.assertIn("Thông tin sự cố", headings)
        self.assertIn("Dòng thời gian", headings)
        self.assertIn("MITRE ATT&CK", headings)
        table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        self.assertIn("EDR-001", table_text)
        self.assertIn("IoC liên quan", table_text)
        self.assertIn("Isolate host", table_text)
        self.assertIn("SOC", table_text)
        self.assertNotIn("Chi tiết kết quả CA các máy chủ", headings)


if __name__ == "__main__":
    unittest.main()
