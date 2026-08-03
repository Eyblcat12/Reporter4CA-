from __future__ import annotations

import base64
import io
import sys
import tempfile
import time
import unittest
from pathlib import Path
from unittest.mock import Mock, patch

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "apps" / "backend"
sys.path.insert(0, str(BACKEND))

from api.routes import router  # noqa: E402
from core.database import Database  # noqa: E402
from core.report_jobs import ReportJobManager  # noqa: E402
from core.preview_artifacts import PreviewArtifactRegistry  # noqa: E402


class ApiIntegrationTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary.name)
        self.database = Database(self.root / "data" / "reporter.db")
        self.database.initialize()
        self.template_root = self.root / "templates"
        self.template_root.mkdir(parents=True)
        self.jobs = ReportJobManager(max_workers=1, max_pending=2)
        self.preview_registry = PreviewArtifactRegistry(self.root / "preview-cache")
        self.patches = [
            patch("api.routes.get_db", return_value=self.database),
            patch("api.routes.TEMPLATES_DIR", self.template_root),
            patch("api.routes._report_jobs", self.jobs),
            patch("api.routes._preview_artifacts", self.preview_registry),
            patch("api.routes.preview_jobs_enabled", return_value=True),
            patch("api.routes.GENERATED_REPORTS_DIR", self.root / "generated"),
        ]
        for active in self.patches:
            active.start()
        app = FastAPI()
        app.include_router(router)
        self.client = TestClient(app)

    def tearDown(self) -> None:
        self.client.close()
        self.jobs.shutdown()
        for active in reversed(self.patches):
            active.stop()
        self.database.close()
        self.temporary.cleanup()

    def test_import_and_validate_contracts(self) -> None:
        content = base64.b64encode(
            b"type,hostname,ip,os,result\nserver,srv-01,10.0.0.1,Linux,Clean\n"
        ).decode()
        imported = self.client.post("/api/import-file", json={
            "filename": "assets.csv", "contentBase64": content, "defaultType": "server",
        })
        self.assertEqual(imported.status_code, 200)
        self.assertEqual(imported.json()["rows"][0]["hostname"], "srv-01")

        validated = self.client.post("/api/validate-rows", json={"rows": imported.json()["rows"]})
        self.assertEqual(validated.status_code, 200)
        self.assertTrue(validated.json()["valid"])
        self.assertEqual(validated.json()["summary"]["totalRows"], 1)

        rules = self.client.get("/api/rules")
        self.assertEqual(rules.status_code, 200)
        self.assertTrue(any(rule["id"] == "PROXY_TOOL_REVIEW" for rule in rules.json()["rules"]))

    def test_workspace_backup_dry_run_and_confirmed_restore_contract(self) -> None:
        template_path = self.template_root / "summary" / "summary.docx"
        template_path.parent.mkdir(parents=True)
        document = Document()
        document.add_heading("Summary", level=1)
        document.save(template_path)
        self.database.add_template(
            name="Summary",
            filename=template_path.name,
            file_path=str(template_path),
            report_type="summary",
        )

        downloaded = self.client.get("/api/system/backup")
        self.assertEqual(downloaded.status_code, 200)
        self.assertEqual(downloaded.headers["x-backup-templates"], "1")
        files = {"backup": ("workspace.zip", downloaded.content, "application/zip")}
        preview = self.client.post("/api/system/restore/preview", files=files)
        self.assertEqual(preview.status_code, 200)
        self.assertTrue(preview.json()["dryRun"])
        self.assertEqual(preview.json()["templateCount"], 1)

        restored = self.client.post(
            "/api/system/restore",
            files={"backup": ("workspace.zip", downloaded.content, "application/zip")},
            data={"confirmationToken": preview.json()["confirmationToken"]},
        )
        self.assertEqual(restored.status_code, 200)
        self.assertTrue(restored.json()["restored"])

    def test_incident_validation_contract_and_preview_guard(self) -> None:
        invalid = self.client.post("/api/validate-incident", json={"metadata": {}})
        self.assertEqual(invalid.status_code, 200)
        self.assertFalse(invalid.json()["valid"])
        self.assertEqual(invalid.json()["summary"]["errors"], 3)

        blocked = self.client.post("/api/preview-docx", json={
            "rows": [{"type": "server", "hostname": "srv-ir", "result": "Clean"}],
            "reportType": "incident_response",
            "metadata": {},
        })
        self.assertEqual(blocked.status_code, 422)
        self.assertIn("incidentQuality", blocked.json()["detail"])

        ready = self.client.post("/api/validate-incident", json={"metadata": {
            "incidentId": "IR-001", "detectedAt": "2026-07-21T09:00",
            "timeline": [{"time": "09:00", "event": "Alert", "evidence": "E-1"}],
        }})
        self.assertEqual(ready.status_code, 200)
        self.assertTrue(ready.json()["valid"])

    def test_custom_rule_can_be_tested_saved_and_updated(self) -> None:
        draft = {
            "name": "Proxy nội bộ", "severity": "medium",
            "classification": "needs_review", "remediation": "Xác minh phê duyệt",
            "conditions": {
                "fields": ["notes"], "containsAny": ["Acme Relay"],
                "excludeContainsAny": ["đã phê duyệt"],
            },
        }
        tested = self.client.post("/api/rules/evaluate", json={
            "rule": draft,
            "rows": [
                {"hostname": "PC-01", "notes": "Có Acme Relay"},
                {"hostname": "PC-02", "notes": "Acme Relay đã phê duyệt"},
            ],
        })
        self.assertEqual(tested.status_code, 200)
        self.assertEqual(tested.json()["matchedRows"], 1)
        self.assertEqual(tested.json()["changedRows"], 1)
        self.assertEqual(tested.json()["impact"]["clients"], 1)
        self.assertEqual(tested.json()["matches"][0]["classificationAfter"], "needs_review")
        self.assertEqual(tested.json()["matches"][0]["hostname"], "PC-01")

        created = self.client.post("/api/rules", json=draft)
        self.assertEqual(created.status_code, 201)
        rule_id = created.json()["id"]
        self.assertTrue(rule_id.startswith("CUSTOM_"))

        updated_draft = {**draft, "classification": "anomaly"}
        updated = self.client.patch(f"/api/rules/{rule_id}", json=updated_draft)
        self.assertEqual(updated.status_code, 200)
        self.assertEqual(updated.json()["classification"], "anomaly")
        retested = self.client.post("/api/rules/evaluate", json={
            "rule": updated_draft,
            "editingRuleId": rule_id,
            "rows": [{"type": "client", "hostname": "PC-01", "notes": "Có Acme Relay"}],
        })
        self.assertEqual(retested.status_code, 200)
        self.assertEqual(retested.json()["changedRows"], 1)
        self.assertEqual(retested.json()["matches"][0]["classificationAfter"], "anomaly")
        versions = self.client.get(f"/api/rules/{rule_id}/versions")
        self.assertEqual(versions.status_code, 200)
        self.assertEqual(
            [item["versionNumber"] for item in versions.json()["versions"]], [2, 1]
        )
        restored = self.client.post(f"/api/rules/{rule_id}/versions/1/rollback")
        self.assertEqual(restored.status_code, 200)
        self.assertEqual(restored.json()["classification"], "needs_review")
        self.assertEqual(restored.json()["version"], "3")
        listed = self.client.get("/api/rules").json()["rules"]
        self.assertTrue(any(rule["id"] == rule_id and rule["editable"] for rule in listed))

    def test_rule_bundle_clone_import_export_and_conflict_contracts(self) -> None:
        draft = {
            "name": "Team proxy anomaly", "severity": "high", "classification": "anomaly",
            "conditions": {"fields": ["software"], "containsAny": ["proxifier"]},
        }
        created = self.client.post("/api/rules", json=draft)
        self.assertEqual(created.status_code, 201)
        rule_id = created.json()["id"]

        conflicts = self.client.get("/api/rules/conflicts")
        self.assertEqual(conflicts.status_code, 200)
        self.assertTrue(any(
            rule_id in item["ruleIds"] and "PROXY_TOOL_REVIEW" in item["ruleIds"]
            for item in conflicts.json()["conflicts"]
        ))

        cloned = self.client.post(f"/api/rules/{rule_id}/clone")
        self.assertEqual(cloned.status_code, 201)
        self.assertIn("(copy)", cloned.json()["name"])

        exported = self.client.get("/api/rules/export")
        self.assertEqual(exported.status_code, 200)
        bundle = exported.json()
        self.assertGreaterEqual(len(bundle["rules"]), 2)
        self.assertNotIn("createdAt", bundle["rules"][0])

        skipped = self.client.post("/api/rules/import", json={"rules": bundle["rules"], "strategy": "skip"})
        self.assertEqual(skipped.status_code, 200)
        self.assertEqual(len(skipped.json()["imported"]), 0)
        self.assertEqual(len(skipped.json()["skipped"]), len(bundle["rules"]))

        renamed = self.client.post("/api/rules/import", json={"rules": [bundle["rules"][0]], "strategy": "rename"})
        self.assertEqual(renamed.status_code, 200)
        self.assertEqual(len(renamed.json()["imported"]), 1)
        self.assertIn("(imported)", renamed.json()["imported"][0]["name"])

    def test_generate_contract_returns_download_and_history_id(self) -> None:
        output = self.root / "generated.docx"
        output.write_bytes(b"docx-result")
        field_result = Mock(engine="deferred")
        with (
            patch("api.routes.generate_report", return_value=Mock()),
            patch("api.routes._apply_document_plugins", side_effect=lambda document, *_: document),
            patch("api.routes._save_finalized_report", return_value=(output, field_result)),
        ):
            response = self.client.post("/api/generate", json={
                "rows": [{
                    "type": "server", "hostname": "srv-01", "ip": "10.0.0.1",
                    "os": "Linux", "result": "Clean",
                }],
                "title": "Integration report", "reportType": "full",
            })
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.content, b"docx-result")
        self.assertTrue(response.headers["x-report-id"])
        self.assertEqual(self.database.list_reports(1)[0]["status"], "success")

    def test_preview_and_generate_share_backend_content_signature(self) -> None:
        document = Document()
        document.add_heading("Parity", level=1)
        document._reporter_manifest = {}
        document._reporter_integrity = {"valid": True, "verifiedAssets": 1}

        def save_document(current, path, **_kwargs):
            output = Path(path)
            current.save(output)
            return output, Mock(engine="deferred")

        request = {
            "rows": [{
                "type": "server", "hostname": "srv-parity", "ip": "10.0.0.7",
                "os": "Linux", "result": "Clean",
            }],
            "organization": "Reporter Pro",
            "reportType": "full",
        }
        with (
            patch("api.routes.generate_report", return_value=document),
            patch("api.routes._apply_document_plugins", side_effect=lambda current, *_: current),
            patch("api.routes._save_finalized_report", side_effect=save_document),
        ):
            preview = self.client.post("/api/preview-docx", json=request)
            generated = self.client.post(
                "/api/generate",
                json={**request, "outputName": "name-does-not-affect-content.docx"},
            )

        self.assertEqual(preview.status_code, 200)
        self.assertEqual(generated.status_code, 200)
        self.assertEqual(
            preview.headers["x-request-signature"],
            generated.headers["x-request-signature"],
        )
        self.assertEqual(
            preview.headers["x-content-signature"],
            generated.headers["x-content-signature"],
        )

    def test_preview_job_is_pollable_downloadable_and_reused_by_signature(self) -> None:
        document = Document()
        document.add_heading("Async Preview", level=1)
        document._reporter_manifest = {}
        document._reporter_integrity = {"valid": True, "verifiedAssets": 1}

        def save_document(current, path, **_kwargs):
            output = Path(path)
            current.save(output)
            return output, Mock(engine="deferred")

        request = {
            "rows": [{
                "type": "server", "hostname": "srv-preview", "ip": "10.0.0.9",
                "os": "Linux", "result": "Clean",
            }],
            "reportType": "full",
            "clientRequestId": "preview-sequence-1",
            "disablePlugins": True,
        }
        with (
            patch("api.routes.generate_report", return_value=document),
            patch("api.routes._apply_document_plugins", side_effect=lambda current, *_: current),
            patch("api.routes._save_finalized_report", side_effect=save_document),
        ):
            created = self.client.post("/api/preview-jobs", json=request)
            self.assertEqual(created.status_code, 202)
            preview_id = created.json()["previewId"]
            deadline = time.time() + 5
            state = None
            while time.time() < deadline:
                state = self.client.get(f"/api/preview-jobs/{preview_id}")
                if state.json().get("status") == "ready":
                    break
                time.sleep(0.02)

            self.assertEqual(state.status_code, 200)
            self.assertEqual(state.json()["status"], "ready")
            content = self.client.get(f"/api/preview-jobs/{preview_id}/content")
            self.assertEqual(content.status_code, 200)
            self.assertGreater(len(content.content), 100)
            self.assertEqual(content.headers["x-preview-id"], preview_id)

            reused = self.client.post("/api/preview-jobs", json=request)
            self.assertEqual(reused.status_code, 202)
            self.assertTrue(reused.json()["deduplicated"])
            self.assertEqual(reused.json()["previewId"], preview_id)

    def test_report_job_promotes_preview_byte_for_byte(self) -> None:
        document = Document()
        document.add_heading("Promoted Preview", level=1)
        document._reporter_manifest = {}
        document._reporter_integrity = {"valid": True, "verifiedAssets": 1}

        def save_document(current, path, **_kwargs):
            output = Path(path)
            current.save(output)
            return output, Mock(engine="deferred")

        request = {
            "rows": [{
                "type": "server", "hostname": "srv-promote", "ip": "10.0.0.10",
                "os": "Linux", "result": "Clean",
            }],
            "reportType": "full",
            "disablePlugins": True,
        }
        with (
            patch("api.routes.generate_report", return_value=document),
            patch("api.routes._apply_document_plugins", side_effect=lambda current, *_: current),
            patch("api.routes._save_finalized_report", side_effect=save_document),
        ):
            created = self.client.post("/api/preview-jobs", json=request)
            preview_id = created.json()["previewId"]
            deadline = time.time() + 5
            while time.time() < deadline:
                state = self.client.get(f"/api/preview-jobs/{preview_id}").json()
                if state.get("status") == "ready":
                    break
                time.sleep(0.02)
            preview_bytes = self.client.get(
                f"/api/preview-jobs/{preview_id}/content"
            ).content

        with patch("api.routes.preview_cache_enabled", return_value=True):
            submitted = self.client.post(
                "/api/report-jobs",
                json={**request, "previewId": preview_id, "outputName": "promoted.docx"},
            )
            self.assertEqual(submitted.status_code, 202)
            job_id = submitted.json()["job"]["id"]
            deadline = time.time() + 5
            while time.time() < deadline:
                job = self.client.get(f"/api/report-jobs/{job_id}").json()["job"]
                if job["status"] in {"completed", "failed", "cancelled"}:
                    break
                time.sleep(0.02)

        self.assertEqual(job["status"], "completed")
        self.assertEqual(job["cacheMode"], "preview_cache_hit")
        downloaded = self.client.get(f"/api/report-jobs/{job_id}/download")
        self.assertEqual(downloaded.status_code, 200)
        self.assertEqual(downloaded.content, preview_bytes)
        history = self.database.list_reports(None)
        self.assertEqual(len(history), 1)
        self.assertEqual(history[0]["job_id"], job_id)
        self.assertEqual(history[0]["source_artifact_id"], preview_id)
        self.assertEqual(history[0]["cache_status"], "preview_cache_hit")

    def test_generate_failure_removes_temporary_docx(self) -> None:
        created_paths: list[Path] = []
        real_named_temporary_file = tempfile.NamedTemporaryFile

        def tracked_temporary_file(*args, **kwargs):
            temporary = real_named_temporary_file(*args, dir=self.root, **kwargs)
            created_paths.append(Path(temporary.name))
            return temporary

        with (
            patch("api.routes.generate_report", return_value=Mock()),
            patch("api.routes._apply_document_plugins", side_effect=lambda document, *_: document),
            patch("api.routes.tempfile.NamedTemporaryFile", side_effect=tracked_temporary_file),
            patch("api.routes._save_finalized_report", side_effect=RuntimeError("save failed")),
        ):
            response = self.client.post("/api/generate", json={
                "rows": [{"type": "server", "hostname": "srv-01", "result": "Clean"}],
                "reportType": "full",
            })

        self.assertEqual(response.status_code, 500)
        self.assertEqual(len(created_paths), 1)
        self.assertFalse(created_paths[0].exists())
        self.assertEqual(self.database.list_reports(1)[0]["status"], "failed")

    def test_template_and_preset_crud_contracts(self) -> None:
        buffer = io.BytesIO()
        document = Document()
        document.add_heading("Template", level=1)
        document.save(buffer)
        uploaded = self.client.post("/api/templates/upload", json={
            "filename": "team.docx",
            "contentBase64": base64.b64encode(buffer.getvalue()).decode(),
            "name": "Team template", "reportType": "technical",
        })
        self.assertEqual(uploaded.status_code, 200)
        template_id = uploaded.json()["id"]
        patched = self.client.patch(f"/api/templates/{template_id}", json={"description": "Reviewed"})
        self.assertEqual(patched.status_code, 200)

        preset = self.client.post("/api/presets", json={
            "name": "Team preset", "settings": {"reportType": "technical"},
            "templateId": template_id,
        })
        self.assertEqual(preset.status_code, 200)
        preset_id = preset.json()["id"]
        self.assertEqual(self.client.get(f"/api/presets/{preset_id}").json()["name"], "Team preset")
        self.assertEqual(len(self.client.get("/api/presets").json()["presets"]), 1)

    def test_incompatible_template_cannot_become_default(self) -> None:
        buffer = io.BytesIO()
        Document().save(buffer)
        uploaded = self.client.post("/api/templates/upload", json={
            "filename": "empty.docx",
            "contentBase64": base64.b64encode(buffer.getvalue()).decode(),
            "name": "Empty", "reportType": "incident_response", "isDefault": True,
        })
        self.assertEqual(uploaded.status_code, 200)
        payload = uploaded.json()
        self.assertEqual(payload["analysis"]["compatibility"]["status"], "incompatible")
        self.assertTrue(payload["defaultRejected"])

        template_id = payload["id"]
        make_default = self.client.patch(
            f"/api/templates/{template_id}", json={"isDefault": True}
        )
        self.assertEqual(make_default.status_code, 422)
        listed = self.client.get("/api/templates").json()["templates"]
        saved = next(item for item in listed if item["id"] == template_id)
        self.assertEqual(saved["compatibilityStatus"], "incompatible")
        self.assertFalse(saved["isDefault"])

    def test_template_versions_can_compare_and_rollback(self) -> None:
        original = io.BytesIO()
        document = Document()
        document.add_heading("Cover", level=1)
        document.save(original)
        uploaded = self.client.post("/api/templates/upload", json={
            "filename": "versioned.docx", "contentBase64": base64.b64encode(original.getvalue()).decode(),
            "name": "Versioned", "reportType": "full",
        })
        template_id = uploaded.json()["id"]

        revised = io.BytesIO()
        document = Document()
        document.add_heading("Cover", level=1)
        document.add_paragraph("{{TITLE}}")
        document.save(revised)
        created = self.client.post(f"/api/templates/{template_id}/versions", json={
            "contentBase64": base64.b64encode(revised.getvalue()).decode(), "note": "Add title token",
        })
        self.assertEqual(created.status_code, 200)
        self.assertTrue(created.json()["activated"])
        self.assertEqual(created.json()["version"]["version"], 2)

        compared = self.client.get(
            f"/api/templates/{template_id}/versions/compare?fromVersion=1&toVersion=2"
        )
        self.assertEqual(compared.status_code, 200)
        self.assertIn("{{TITLE}}", compared.json()["diff"]["tokens"]["added"])
        rolled_back = self.client.post(f"/api/templates/{template_id}/versions/1/rollback")
        self.assertEqual(rolled_back.status_code, 200)
        self.assertEqual(rolled_back.json()["activeVersion"], 1)

    def test_history_dashboard_and_backup_contracts(self) -> None:
        self.database.add_report(title="Recorded", row_count=5, status="success", duration_ms=20)
        history = self.client.get("/api/reports/history")
        self.assertEqual(history.status_code, 200)
        self.assertEqual(history.json()["reports"][0]["title"], "Recorded")

        dashboard = self.client.get("/api/dashboard/summary?days=30")
        self.assertEqual(dashboard.status_code, 200)
        self.assertEqual(dashboard.json()["metrics"]["reports"], 1)

        backup = self.client.get("/api/system/backup")
        self.assertEqual(backup.status_code, 200)
        self.assertEqual(backup.headers["content-type"], "application/zip")
        self.assertGreater(len(backup.content), 100)

    def test_report_job_lifecycle_and_download_contract(self) -> None:
        output = self.root / "job-output.docx"
        output.write_bytes(b"job-result")

        def runner(job):
            self.jobs.update(job, phase="generating", progress=60, message="Generating")
            return {"outputPath": str(output), "filename": "job.docx", "reportId": "job-report"}

        with patch("api.routes._run_report_job", side_effect=runner):
            created = self.client.post("/api/report-jobs", json={
                "rows": [{"type": "server", "hostname": "srv-job", "result": "Clean"}],
                "reportType": "full",
            })
            self.assertEqual(created.status_code, 202)
            self.assertEqual(len(created.json()["job"]["requestSignature"]), 64)
            job_id = created.json()["job"]["id"]
            deadline = time.time() + 2
            status = ""
            while time.time() < deadline:
                status_response = self.client.get(f"/api/report-jobs/{job_id}")
                status = status_response.json()["job"]["status"]
                if status == "completed":
                    break
                time.sleep(0.01)

        self.assertEqual(status, "completed")
        downloaded = self.client.get(f"/api/report-jobs/{job_id}/download")
        self.assertEqual(downloaded.status_code, 200)
        self.assertEqual(downloaded.content, b"job-result")
        self.assertEqual(downloaded.headers["x-report-id"], "job-report")

    def test_running_report_job_can_be_cancelled(self) -> None:
        entered = __import__("threading").Event()

        def runner(job):
            entered.set()
            while True:
                self.jobs.check_cancelled(job)
                time.sleep(0.01)

        with patch("api.routes._run_report_job", side_effect=runner):
            created = self.client.post("/api/report-jobs", json={
                "rows": [{"type": "server", "hostname": "srv-cancel"}],
            })
            job_id = created.json()["job"]["id"]
            self.assertTrue(entered.wait(1))
            cancelled = self.client.delete(f"/api/report-jobs/{job_id}")
            self.assertEqual(cancelled.status_code, 200)
            deadline = time.time() + 2
            status = ""
            while time.time() < deadline:
                status = self.client.get(f"/api/report-jobs/{job_id}").json()["job"]["status"]
                if status == "cancelled":
                    break
                time.sleep(0.01)
        self.assertEqual(status, "cancelled")
        history = self.database.list_reports(None)
        self.assertEqual(len(history), 1)
        self.assertEqual(history[0]["job_id"], job_id)
        self.assertEqual(history[0]["status"], "cancelled")


if __name__ == "__main__":
    unittest.main()
