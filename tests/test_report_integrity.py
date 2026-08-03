from __future__ import annotations

import sys
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "apps" / "backend"
sys.path.insert(0, str(BACKEND))

from core.report_integrity import (  # noqa: E402
    ReportIntegrityError,
    build_report_manifest,
    verify_report_document,
)
from core import report_integrity  # noqa: E402
from core.rule_engine import evaluate_payload  # noqa: E402


class ReportIntegrityTests(unittest.TestCase):
    def _add_required_sections(self, document, manifest, *, omit: str | None = None) -> None:
        for title in manifest["requiredSections"]:
            if title != omit:
                document.add_heading(title, level=1)

    def _add_summary_rows(self, document, manifest, *, force_type: str | None = None) -> None:
        for asset_type, label in (("server", "Máy chủ"), ("client", "Máy trạm")):
            assets = [
                asset for asset in manifest["assets"]
                if asset["assetType"] == asset_type
            ]
            if not assets:
                continue
            rendered_label = (
                "Máy chủ" if force_type == "server"
                else "Máy trạm" if force_type == "client"
                else label
            )
            table = document.add_table(rows=len(assets) + 1, cols=3)
            for cell, value in zip(
                table.rows[0].cells,
                ("STT", rendered_label, "Kết quả rà soát đánh giá"),
            ):
                cell.text = value
            for index, asset in enumerate(assets, start=1):
                table.rows[index].cells[0].text = str(index)
                table.rows[index].cells[1].text = asset["hostname"]
                table.rows[index].cells[2].text = asset["assessmentLabel"]

    def _add_technical_findings(self, document, manifest, *, evidence: str | None = None) -> None:
        findings = [
            (asset, finding)
            for asset in manifest["assets"]
            for finding in asset["findings"]
        ]
        table = document.add_table(rows=len(findings) + 1, cols=5)
        for cell, value in zip(
            table.rows[0].cells,
            ("Tài sản", "Rule", "Mức độ", "Phân loại", "Bằng chứng"),
        ):
            cell.text = value
        for index, (asset, finding) in enumerate(findings, start=1):
            table.rows[index].cells[0].text = asset["hostname"]
            table.rows[index].cells[1].text = finding["ruleId"]
            table.rows[index].cells[2].text = "high"
            table.rows[index].cells[3].text = "anomaly"
            table.rows[index].cells[4].text = (
                finding["evidenceText"] if evidence is None else evidence
            )

    def test_manifest_counts_every_evaluated_asset_and_conclusion(self) -> None:
        data = evaluate_payload({
            "servers": [
                {"hostname": "SRV-01", "result": "Ghi nhận dấu hiệu bất thường"},
                {"hostname": "SRV-02", "result": "Không phát hiện"},
            ],
            "clients": [{"hostname": "PC-01", "result": "Chưa kết luận"}],
            "metadata": {},
        })
        manifest = build_report_manifest(data, "full")

        self.assertEqual(manifest["assetCount"], 3)
        self.assertEqual(manifest["assetTypeCounts"], {"server": 2, "client": 1})
        self.assertEqual(
            manifest["assessmentCounts"],
            {"anomaly": 1, "clean": 1, "insufficient_data": 1},
        )

    def test_manifest_counts_only_evidence_backed_rules(self) -> None:
        data = {
            "servers": [{
                "hostname": "SRV-01",
                "findings": [
                    {
                        "ruleId": "MALWARE_EVIDENCE",
                        "classification": "anomaly",
                        "evidence": [{"field": "notes", "value": "suspicious.exe"}],
                    },
                    {
                        "ruleId": "NO_EVIDENCE",
                        "classification": "anomaly",
                        "evidence": [],
                    },
                ],
            }],
            "clients": [],
        }

        manifest = build_report_manifest(data, "full")

        self.assertEqual(manifest["findingCount"], 1)
        self.assertEqual(manifest["evidenceCount"], 1)
        self.assertEqual(manifest["ruleCounts"], {"MALWARE_EVIDENCE": 1})

    def test_manifest_declares_distinct_required_sections_for_all_report_types(self) -> None:
        expected_distinctive_section = {
            "full": "Đánh giá chung đối với máy chủ",
            "server_only": "Chi tiết kết quả CA các máy chủ",
            "client_only": "Chi tiết kết quả CA các máy trạm",
            "summary": "Kết quả và phân tích tổng hợp",
            "technical": "Findings từ rule engine",
            "incident_response": "Ứng phó sự cố",
        }

        for report_type, section in expected_distinctive_section.items():
            with self.subTest(report_type=report_type):
                manifest = build_report_manifest(
                    {"servers": [], "clients": []},
                    report_type,
                )
                self.assertIn(section, manifest["requiredSections"])

    def test_verifier_rejects_a_missing_asset_conclusion(self) -> None:
        data = evaluate_payload({
            "servers": [{"hostname": "SRV-01", "result": "Ghi nhận dấu hiệu bất thường"}],
            "clients": [{"hostname": "PC-01", "result": "Không phát hiện"}],
            "metadata": {},
        })
        manifest = build_report_manifest(data, "full")
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.rows[1].cells[0].text = "SRV-01"
        table.rows[1].cells[1].text = "Ghi nhận dấu hiệu bất thường"

        with self.assertRaisesRegex(ReportIntegrityError, "PC-01"):
            verify_report_document(document, manifest)

    def test_verifier_requires_hostname_and_conclusion_in_the_same_row(self) -> None:
        data = evaluate_payload({
            "servers": [{"hostname": "SRV-01", "result": "Ghi nhận dấu hiệu bất thường"}],
            "clients": [],
            "metadata": {},
        })
        manifest = build_report_manifest(data, "full")
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "SRV-01"
        table.rows[1].cells[1].text = manifest["assets"][0]["assessmentLabel"]

        with self.assertRaisesRegex(ReportIntegrityError, "SRV-01"):
            verify_report_document(document, manifest)

    def test_technical_verifier_reports_all_applicable_counters(self) -> None:
        data = {
            "servers": [{
                "hostname": "SRV-01",
                "findings": [{
                    "ruleId": "MALWARE_EVIDENCE",
                    "classification": "anomaly",
                    "severity": "high",
                    "evidence": [{"field": "notes", "value": "suspicious.exe"}],
                }],
            }],
            "clients": [],
        }
        manifest = build_report_manifest(data, "technical")
        document = Document()
        self._add_required_sections(document, manifest)
        self._add_summary_rows(document, manifest)
        self._add_technical_findings(document, manifest)

        result = verify_report_document(document, manifest)

        self.assertTrue(result["valid"])
        self.assertEqual(result["expectedAssets"], 1)
        self.assertEqual(result["actualAssets"], 1)
        self.assertEqual(result["verifiedAssets"], 1)
        self.assertEqual(result["actualAssetTypes"], {"server": 1})
        self.assertEqual(result["verifiedAssetTypes"], {"server": 1})
        self.assertEqual(result["expectedFindings"], 1)
        self.assertEqual(result["actualFindings"], 1)
        self.assertEqual(result["verifiedFindings"], 1)
        self.assertEqual(result["expectedRules"], {"MALWARE_EVIDENCE": 1})
        self.assertEqual(result["actualRules"], {"MALWARE_EVIDENCE": 1})
        self.assertEqual(result["verifiedRules"], {"MALWARE_EVIDENCE": 1})
        self.assertEqual(result["expectedEvidence"], 1)
        self.assertIsNone(result["actualEvidence"])
        self.assertEqual(result["verifiedEvidence"], 1)
        self.assertEqual(result["expectedSections"], len(manifest["requiredSections"]))
        self.assertEqual(result["verifiedSections"], len(manifest["requiredSections"]))
        self.assertEqual(result["errorCodes"], [])

    def test_verifier_exposes_structured_section_and_evidence_errors(self) -> None:
        data = {
            "servers": [{
                "hostname": "SRV-01",
                "findings": [{
                    "ruleId": "MALWARE_EVIDENCE",
                    "classification": "anomaly",
                    "severity": "high",
                    "evidence": [{"field": "notes", "value": "suspicious.exe"}],
                }],
            }],
            "clients": [],
        }
        manifest = build_report_manifest(data, "technical")
        document = Document()
        self._add_required_sections(document, manifest, omit="Kết luận và khuyến nghị")
        self._add_summary_rows(document, manifest)
        self._add_technical_findings(document, manifest, evidence="tampered evidence")

        with self.assertRaises(ReportIntegrityError) as raised:
            verify_report_document(document, manifest)

        error = raised.exception
        self.assertEqual(error.code, "REQUIRED_SECTION_MISSING")
        self.assertEqual(
            error.result["errorCodes"],
            ["REQUIRED_SECTION_MISSING", "EVIDENCE_COUNT_MISMATCH"],
        )
        self.assertEqual(error.result["missingSections"], ["Kết luận và khuyến nghị"])
        self.assertEqual(error.result["verifiedFindings"], 1)
        self.assertEqual(error.result["verifiedEvidence"], 0)

    def test_asset_in_wrong_scope_is_not_counted_as_verified(self) -> None:
        data = evaluate_payload({
            "servers": [{"hostname": "SRV-01", "result": "Không phát hiện"}],
            "clients": [],
            "metadata": {},
        })
        manifest = build_report_manifest(data, "server_only")
        document = Document()
        self._add_required_sections(document, manifest)
        self._add_summary_rows(document, manifest, force_type="client")

        with self.assertRaises(ReportIntegrityError) as raised:
            verify_report_document(document, manifest)

        result = raised.exception.result
        self.assertIn("ASSET_TYPE_COUNT_MISMATCH", result["errorCodes"])
        self.assertIn("ASSET_CONCLUSION_MISSING", result["errorCodes"])
        self.assertEqual(result["actualAssets"], 1)
        self.assertEqual(result["verifiedAssets"], 0)
        self.assertEqual(result["actualAssetTypes"], {"client": 1})
        self.assertEqual(result["verifiedAssetTypes"], {})

    def test_wrong_rule_row_returns_finding_and_rule_error_codes(self) -> None:
        data = {
            "servers": [{
                "hostname": "SRV-01",
                "findings": [{
                    "ruleId": "MALWARE_EVIDENCE",
                    "classification": "anomaly",
                    "severity": "high",
                    "evidence": [{"field": "notes", "value": "suspicious.exe"}],
                }],
            }],
            "clients": [],
        }
        manifest = build_report_manifest(data, "technical")
        document = Document()
        self._add_required_sections(document, manifest)
        self._add_summary_rows(document, manifest)
        self._add_technical_findings(document, manifest)
        finding_table = document.tables[-1]
        finding_table.rows[1].cells[1].text = "UNEXPECTED_RULE"

        with self.assertRaises(ReportIntegrityError) as raised:
            verify_report_document(document, manifest)

        result = raised.exception.result
        self.assertIn("FINDING_COUNT_MISMATCH", result["errorCodes"])
        self.assertIn("RULE_COUNT_MISMATCH", result["errorCodes"])
        self.assertEqual(result["actualFindings"], 1)
        self.assertEqual(result["verifiedFindings"], 0)
        self.assertEqual(result["actualRules"], {"UNEXPECTED_RULE": 1})
        self.assertEqual(result["verifiedRules"], {})

    def test_duplicate_manifest_assets_require_distinct_rendered_rows(self) -> None:
        data = evaluate_payload({
            "servers": [
                {"hostname": "SRV-DUP", "result": "Không phát hiện"},
                {"hostname": "SRV-DUP", "result": "Không phát hiện"},
            ],
            "clients": [],
            "metadata": {},
        })
        manifest = build_report_manifest(data, "server_only")
        document = Document()
        self._add_required_sections(document, manifest)

        first_only_manifest = dict(manifest)
        first_only_manifest["assets"] = manifest["assets"][:1]
        self._add_summary_rows(document, first_only_manifest)

        with self.assertRaises(ReportIntegrityError) as raised:
            verify_report_document(document, manifest)

        result = raised.exception.result
        self.assertIn("ASSET_COUNT_MISMATCH", result["errorCodes"])
        self.assertIn("ASSET_CONCLUSION_MISSING", result["errorCodes"])
        self.assertEqual(result["expectedAssets"], 2)
        self.assertEqual(result["actualAssets"], 1)
        self.assertEqual(result["verifiedAssets"], 1)

    def test_document_index_reads_each_1000_and_3000_asset_cell_once(self) -> None:
        class CountingCell:
            def __init__(self, value: str, counter: list[int]) -> None:
                self.value = value
                self.counter = counter

            @property
            def text(self) -> str:
                self.counter[0] += 1
                return self.value

        class Row:
            def __init__(self, values: tuple[str, ...], counter: list[int]) -> None:
                self.cells = [CountingCell(value, counter) for value in values]

        class Table:
            def __init__(self, size: int, counter: list[int]) -> None:
                self.rows = [
                    Row(("STT", "Máy chủ", "Kết quả rà soát đánh giá"), counter),
                    *[
                        Row((str(index), f"SRV-{index:05d}", "Sạch"), counter)
                        for index in range(1, size + 1)
                    ],
                ]

        class DocumentStub:
            def __init__(self, size: int, counter: list[int]) -> None:
                self.tables = [Table(size, counter)]
                self.paragraphs = []

        observed_reads: list[int] = []
        for size in (1_000, 3_000):
            counter = [0]
            index = report_integrity._index_document(DocumentStub(size, counter))
            self.assertEqual(sum(index["summaryRows"].values()), size)
            self.assertEqual(counter[0], 3 * (size + 1))
            observed_reads.append(counter[0])

        # Three times more assets cause almost exactly three times more cell
        # reads, rather than the ninefold growth expected from a nested scan.
        self.assertLess(observed_reads[1] / observed_reads[0], 3.01)


if __name__ == "__main__":
    unittest.main()
