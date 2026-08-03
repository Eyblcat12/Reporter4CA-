from __future__ import annotations

import json
import os
import sys
import tempfile
import threading
import time
import unittest
from pathlib import Path
from unittest.mock import Mock, patch

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "apps" / "backend"
SCRIPTS = ROOT / "scripts"
FIXTURES = ROOT / "tests" / "fixtures" / "performance"
sys.path.insert(0, str(BACKEND))
sys.path.insert(0, str(SCRIPTS))

from benchmark_report_generation import (  # noqa: E402
    BenchmarkConfigurationError,
    aggregate_trials,
    assert_expected_manifest,
    load_manifest,
    select_fixture,
    validate_fixture,
)
from api.models import GenerateRequest, ReportType as ApiReportType  # noqa: E402
from core.config import (  # noqa: E402
    performance_metrics_enabled,
    preview_cache_enabled,
    preview_jobs_enabled,
)
from core.docx_field_updater import FieldUpdateResult  # noqa: E402
from core.input_parser import parse_input  # noqa: E402
from core.performance_metrics import (  # noqa: E402
    PerformanceMetrics,
    emit_performance_metrics,
    sanitize_metadata,
)
from core.report_generator import ReportType, generate_report  # noqa: E402
from core.report_integrity import build_report_manifest  # noqa: E402
from core.report_jobs import ReportJobManager  # noqa: E402
from core.rule_engine import evaluate_payload  # noqa: E402
from generate_tracking_fixture import (  # noqa: E402
    PROFILES,
    REPORT_TYPES,
    build_fixture_set,
    verify_fixture_set,
    write_fixture_set,
)


class PerformanceMetricsTests(unittest.TestCase):
    def test_collector_separates_product_and_audit_latency(self) -> None:
        clock = iter((0, 1_000_000, 4_000_000, 5_000_000, 7_000_000, 8_000_000))
        metrics = PerformanceMetrics(
            run_id="controlled-run",
            metadata={"assetCount": 50, "hostname": "must-not-leak"},
            clock_ns=lambda: next(clock),
        )
        with metrics.phase("build"):
            pass
        with metrics.phase("reopen", latency_class="audit"):
            pass
        payload = metrics.public()

        self.assertEqual(payload["runId"], "controlled-run")
        self.assertEqual(payload["metadata"], {"assetCount": 50})
        self.assertEqual(payload["productLatencyMs"], 3.0)
        self.assertEqual(payload["auditLatencyMs"], 2.0)
        self.assertEqual([phase["name"] for phase in payload["phases"]], ["build", "reopen"])

    def test_metadata_is_opt_in_and_drops_source_content(self) -> None:
        sanitized = sanitize_metadata({
            "fixtureId": "mixed-50",
            "featureFlags": {"AUTO_REPORT_FAST_CELL": "0"},
            "hostname": "PC-SECRET",
            "notes": "secret analyst note",
            "payload": {"servers": [{"hostname": "PC-SECRET"}]},
            "unknown": "not explicitly approved",
        })
        self.assertEqual(
            sanitized,
            {
                "fixtureId": "mixed-50",
                "featureFlags": {"AUTO_REPORT_FAST_CELL": "0"},
            },
        )

    def test_phase_is_recorded_when_operation_raises(self) -> None:
        metrics = PerformanceMetrics()
        with self.assertRaisesRegex(RuntimeError, "planned"):
            with metrics.phase("documentBuild"):
                raise RuntimeError("planned")
        self.assertEqual(metrics.public()["phases"][0]["name"], "documentBuild")

    def test_phase_attributes_accept_counts_but_not_source_rows(self) -> None:
        metrics = PerformanceMetrics()
        with metrics.phase(
            "saveZip",
            attributes={"bytes": 2048, "rows": "raw,row,content", "outcome": "passed"},
        ):
            pass
        attributes = metrics.public()["phases"][0]["attributes"]
        self.assertEqual(attributes, {"bytes": 2048, "outcome": "passed"})

    def test_nested_phases_do_not_double_count_product_latency(self) -> None:
        clock = iter((0, 1_000_000, 2_000_000, 5_000_000, 8_000_000, 9_000_000))
        metrics = PerformanceMetrics(clock_ns=lambda: next(clock))
        with metrics.phase("documentBuild"):
            with metrics.phase("templateLoad"):
                pass
        payload = metrics.public()
        phases = {phase["name"]: phase for phase in payload["phases"]}
        self.assertEqual(payload["productLatencyMs"], 7.0)
        self.assertNotIn("nested", phases["documentBuild"])
        self.assertTrue(phases["templateLoad"]["nested"])

    def test_runtime_flag_defaults_off_and_accepts_explicit_truthy_values(self) -> None:
        with patch.dict(os.environ, {}, clear=True):
            self.assertFalse(performance_metrics_enabled())
        with patch.dict(os.environ, {"AUTO_REPORT_PERF_METRICS": "yes"}, clear=True):
            self.assertTrue(performance_metrics_enabled())

    def test_preview_job_and_cache_flags_default_on_with_zero_rollback(self) -> None:
        with patch.dict(os.environ, {}, clear=True):
            self.assertTrue(preview_jobs_enabled())
            self.assertTrue(preview_cache_enabled())
        with patch.dict(
            os.environ,
            {"AUTO_REPORT_PREVIEW_JOBS": "0", "AUTO_REPORT_PREVIEW_CACHE": "0"},
            clear=True,
        ):
            self.assertFalse(preview_jobs_enabled())
            self.assertFalse(preview_cache_enabled())

    def test_emitted_json_contains_only_sanitized_metadata(self) -> None:
        metrics = PerformanceMetrics(metadata={
            "operation": "generate-job",
            "assetCount": 2,
            "hostname": "SECRET-HOST",
        })
        metrics.update_metadata({
            "pluginCount": 1,
            "notes": "secret note",
        })
        with self.assertLogs("reporter.performance", level="INFO") as captured:
            payload = emit_performance_metrics(metrics, outcome="passed")
        self.assertEqual(
            payload["metadata"],
            {"operation": "generate-job", "assetCount": 2, "pluginCount": 1},
        )
        self.assertNotIn("SECRET-HOST", captured.output[0])
        self.assertNotIn("secret note", captured.output[0])

    def test_table_timings_are_aggregated_by_allowlisted_category(self) -> None:
        metrics = PerformanceMetrics()
        metrics.record_aggregate("tableCreate", "assetDetail", 3.0)
        metrics.record_aggregate("tableCreate", "assetDetail", 5.0)
        metrics.record_aggregate("tableStyle", "hostname-secret", 2.0)
        aggregates = {
            (item["name"], item["category"]): item
            for item in metrics.public()["aggregates"]
        }
        self.assertEqual(
            aggregates[("tableCreate", "assetDetail")],
            {
                "name": "tableCreate",
                "category": "assetDetail",
                "count": 2,
                "totalDurationMs": 8.0,
                "maxDurationMs": 5.0,
            },
        )
        self.assertIn(("tableStyle", "other"), aggregates)


class FixtureGeneratorTests(unittest.TestCase):
    def test_fixture_generation_is_byte_for_byte_deterministic(self) -> None:
        manifest_a, files_a = build_fixture_set(count=12, seed=1234)
        manifest_b, files_b = build_fixture_set(count=12, seed=1234)
        manifest_c, files_c = build_fixture_set(count=12, seed=1235)

        self.assertEqual(manifest_a, manifest_b)
        self.assertEqual(files_a, files_b)
        self.assertNotEqual(
            files_a["tracking_clean_12.csv"],
            files_c["tracking_clean_12.csv"],
        )
        self.assertEqual(len(manifest_a["fixtures"]), len(PROFILES))

    def test_fixture_verifier_detects_modified_csv(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output_dir = Path(directory)
            manifest = write_fixture_set(output_dir, count=8, seed=42)
            self.assertEqual(verify_fixture_set(output_dir), [])

            fixture_path = output_dir / manifest["fixtures"][0]["file"]
            fixture_path.write_bytes(fixture_path.read_bytes() + b"\n")
            errors = verify_fixture_set(output_dir)
            self.assertTrue(any("differs" in error for error in errors))

    def test_committed_fixture_manifest_and_hashes_are_reproducible(self) -> None:
        self.assertEqual(verify_fixture_set(FIXTURES), [])

    def test_rule_engine_matches_independent_fixture_expectations(self) -> None:
        manifest = load_manifest(FIXTURES / "manifest.json")
        for fixture in manifest["fixtures"]:
            with self.subTest(profile=fixture["profile"]):
                payload = parse_input(FIXTURES / fixture["file"], default_section="clients")
                evaluated = evaluate_payload(payload)
                for report_type in REPORT_TYPES:
                    actual = build_report_manifest(evaluated, report_type)
                    expected = fixture["expectedByReportType"][report_type]
                    assert_expected_manifest(actual, expected)

    def test_manifest_validation_rejects_tampered_fixture(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output_dir = Path(directory)
            write_fixture_set(output_dir, count=4, seed=99, profiles=("mixed",))
            manifest_path = output_dir / "manifest.json"
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            fixture = manifest["fixtures"][0]
            fixture_path = output_dir / fixture["file"]
            fixture_path.write_bytes(fixture_path.read_bytes() + b"tampered")

            with self.assertRaisesRegex(BenchmarkConfigurationError, "size mismatch"):
                validate_fixture(manifest_path, fixture, report_type="full")

    def test_fixture_selector_accepts_profile_or_id(self) -> None:
        manifest = load_manifest(FIXTURES / "manifest.json")
        self.assertEqual(select_fixture(manifest, "mixed")["id"], "mixed-50")
        self.assertEqual(select_fixture(manifest, "mixed-50")["profile"], "mixed")


class BenchmarkAggregationTests(unittest.TestCase):
    @staticmethod
    def _trial(index: int) -> dict:
        return {
            "status": "passed",
            "metrics": {
                "productLatencyMs": float(index),
                "auditLatencyMs": float(index) / 2,
                "phases": [{"name": "documentBuild", "durationMs": float(index)}],
            },
            "resources": {"peakRssMiB": 100.0 + index},
            "artifact": {"bytes": 1000 + index},
        }

    def test_p95_is_not_published_for_too_few_samples(self) -> None:
        summary = aggregate_trials([self._trial(index) for index in range(1, 4)])
        self.assertEqual(summary["passedTrials"], 3)
        self.assertEqual(summary["productLatencyMs"]["p50"], 2.0)
        self.assertIsNone(summary["productLatencyMs"]["p95"])
        self.assertFalse(summary["productLatencyMs"]["p95Published"])

    def test_p95_is_published_at_ten_samples(self) -> None:
        summary = aggregate_trials([self._trial(index) for index in range(1, 11)])
        self.assertEqual(summary["productLatencyMs"]["p95"], 9.55)
        self.assertTrue(summary["productLatencyMs"]["p95Published"])

    def test_table_aggregates_are_summarized_across_trials(self) -> None:
        trials = [self._trial(1), self._trial(2)]
        trials[0]["metrics"]["aggregates"] = [{
            "name": "tableCreate",
            "category": "assetDetail",
            "count": 2,
            "totalDurationMs": 10.0,
            "maxDurationMs": 6.0,
        }]
        trials[1]["metrics"]["aggregates"] = [{
            "name": "tableCreate",
            "category": "assetDetail",
            "count": 4,
            "totalDurationMs": 20.0,
            "maxDurationMs": 8.0,
        }]
        aggregate = aggregate_trials(trials)["aggregates"]["tableCreate:assetDetail"]
        self.assertEqual(aggregate["count"]["p50"], 3.0)
        self.assertEqual(aggregate["totalDurationMs"]["p50"], 15.0)


class RuntimeOrchestrationMetricsTests(unittest.TestCase):
    def test_job_manager_records_queue_wait_without_changing_public_contract(self) -> None:
        manager = ReportJobManager(max_workers=1, max_pending=1)
        release = threading.Event()
        metrics = PerformanceMetrics(run_id="queued-job")

        def blocking_runner(_job):
            release.wait(1)
            return {}

        try:
            first, _ = manager.submit({"id": "first"}, blocking_runner)
            second, _ = manager.submit(
                {"id": "second"},
                lambda _job: {},
                metrics=metrics,
            )
            time.sleep(0.03)
            release.set()
            deadline = time.time() + 2
            while time.time() < deadline and second.status != "completed":
                time.sleep(0.01)

            self.assertEqual(first.status, "completed")
            self.assertEqual(second.status, "completed")
            phases = metrics.public()["phases"]
            queue_wait = next(phase for phase in phases if phase["name"] == "queueWait")
            self.assertGreater(queue_wait["durationMs"], 20)
            self.assertNotIn("metrics", second.public())
        finally:
            release.set()
            manager.shutdown()

    def test_report_artifact_records_runtime_pipeline_phases(self) -> None:
        from api import routes

        request = GenerateRequest(
            rows=[{
                "type": "server",
                "hostname": "SRV-ROUTE-01",
                "ip": "10.0.0.1",
                "os": "Windows Server 2022",
                "result": "Không phát hiện",
            }],
            title="Private report title",
            reportType=ApiReportType.FULL,
        )
        document = Mock()
        document._reporter_manifest = {"reportType": "full", "assetCount": 1}
        document._reporter_integrity = {"valid": True, "verifiedAssets": 1}
        database = Mock()
        database.get_template_by_path.return_value = None
        database.add_report.return_value = "metrics-report"
        metrics = PerformanceMetrics(
            run_id="runtime-pipeline",
            metadata={"operation": "generate-job", "assetCount": 1},
        )

        def fake_save(_document, path):
            output = Path(path)
            output.write_bytes(b"PK-controlled-docx")
            return output

        with tempfile.TemporaryDirectory() as directory:
            template_path = Path(directory) / "template.docx"
            template_path.write_bytes(b"controlled-template")
            with (
                patch("api.routes.assess_rows", return_value={
                    "valid": True,
                    "summary": {"validRows": 1},
                }),
                patch("api.routes._metadata_with_custom_rules", side_effect=lambda value: value),
                patch("api.routes._load_plugins", return_value=[Mock(), Mock()]),
                patch("api.routes._apply_input_plugins", side_effect=lambda value, _plugins: value),
                patch("api.routes._apply_document_plugins", side_effect=lambda value, *_args: value),
                patch("api.routes._default_template_path", return_value=str(template_path)),
                patch("api.routes._assert_template_compatible"),
                patch("api.routes.generate_report", return_value=document) as generate,
                patch("api.routes.verify_report_document", return_value={
                    "valid": True,
                    "verifiedAssets": 1,
                }),
                patch("api.routes.save_report", side_effect=fake_save),
                patch(
                    "api.routes.refresh_docx_fields",
                    return_value=FieldUpdateResult(False, "deferred", "controlled"),
                ),
                patch("api.routes.get_db", return_value=database),
                patch("api.routes.emit_performance_metrics") as emit,
            ):
                artifact = routes._create_report_artifact(
                    request,
                    metrics=metrics,
                    queued_at_ns=time.perf_counter_ns() - 2_000_000,
                )

            output_path = Path(artifact["outputPath"])
            try:
                phase_names = {phase["name"] for phase in metrics.public()["phases"]}
                self.assertTrue({
                    "queueWait",
                    "artifactBuildTotal",
                    "snapshotValidation",
                    "pluginLoad",
                    "pluginInput",
                    "templatePreparation",
                    "pluginDocument",
                    "postPluginIntegrityVerify",
                    "saveZip",
                    "wordFieldUpdate",
                }.issubset(phase_names))
                self.assertEqual(metrics.public()["metadata"]["pluginCount"], 2)
                self.assertNotIn("Private report title", json.dumps(metrics.public()))
                self.assertIs(generate.call_args.kwargs["metrics"], metrics)
                emit.assert_called_once_with(metrics, outcome="passed")
            finally:
                output_path.unlink(missing_ok=True)
                routes._last_generated.pop("metrics-report", None)

    def test_metrics_logging_failure_does_not_change_report_result(self) -> None:
        from api import routes

        metrics = PerformanceMetrics()
        with patch(
            "api.routes.emit_performance_metrics",
            side_effect=RuntimeError("logging unavailable"),
        ):
            routes._emit_runtime_metrics(metrics, outcome="passed")


class GeneratorInstrumentationTests(unittest.TestCase):
    @staticmethod
    def _fixture_data() -> dict:
        return {
            "servers": [{
                "hostname": "SRV-METRICS-01",
                "ip": "10.0.0.10",
                "os": "Windows Server 2022",
                "result": "Phát hiện mã độc",
                "notes": "PlugX evidence from controlled test",
            }],
            "clients": [{
                "hostname": "PC-METRICS-01",
                "ip": "10.0.0.20",
                "os": "Windows 11",
                "result": "Không phát hiện",
                "notes": "Controlled clean endpoint",
            }],
            "metadata": {},
        }

    @staticmethod
    def _semantic_projection(document) -> dict:
        return {
            "paragraphs": [
                (paragraph.style.name, paragraph.text)
                for paragraph in document.paragraphs
            ],
            "tables": [
                [[cell.text for cell in row.cells] for row in table.rows]
                for table in document.tables
            ],
            "sectionCount": len(document.sections),
            "manifest": document._reporter_manifest,
            "integrity": document._reporter_integrity,
        }

    def test_generator_emits_required_internal_phase_names_without_changing_output(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            template_path = Path(directory) / "minimal-template.docx"
            template = Document()
            template.add_paragraph("{{TITLE}}")
            template.add_paragraph("Template body", style="Heading 1")
            template.save(template_path)

            metrics = PerformanceMetrics(run_id="instrumented")
            with patch(
                "core.report_generator.prepared_template_enabled",
                return_value=False,
            ), patch(
                "core.report_generator._detect_template_mode",
                return_value="full",
            ):
                baseline = generate_report(
                    self._fixture_data(),
                    title="Metrics parity",
                    organization="Reporter Pro",
                    assessment_date="2026-07-30",
                    template_path=template_path,
                    report_type=ReportType.FULL,
                )
                instrumented = generate_report(
                    self._fixture_data(),
                    title="Metrics parity",
                    organization="Reporter Pro",
                    assessment_date="2026-07-30",
                    template_path=template_path,
                    report_type=ReportType.FULL,
                    metrics=metrics,
                )

        required = {
            "templateResolve",
            "templateLoad",
            "templateDetect",
            "prototypeCapture",
            "tocCleanup",
            "templateTrim",
            "documentConfigure",
            "tokenOrCover",
            "ruleEvaluation",
            "reportBodyBuild",
            "manifestBuild",
            "integrityVerify",
        }
        phases = metrics.public()["phases"]
        self.assertTrue(required.issubset({phase["name"] for phase in phases}))
        self.assertTrue(all(phase["durationMs"] >= 0 for phase in phases))
        aggregates = metrics.public()["aggregates"]
        self.assertTrue(aggregates)
        self.assertTrue(all(
            item["category"] in {
                "assetDetail", "assetInventory", "assetSummary", "ioc",
                "other", "remediation", "rows",
            }
            for item in aggregates
        ))
        self.assertTrue(any(
            item["name"] == "tableCreate" and item["category"] == "assetDetail"
            for item in aggregates
        ))
        self.assertEqual(
            self._semantic_projection(baseline),
            self._semantic_projection(instrumented),
        )


if __name__ == "__main__":
    unittest.main()
