from __future__ import annotations

import json
import sqlite3
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "apps" / "backend"
sys.path.insert(0, str(BACKEND))

from core.database import Database  # noqa: E402
from core.workspace_backup import (  # noqa: E402
    WorkspaceBackupError,
    create_workspace_backup,
    inspect_workspace_backup,
    restore_workspace_backup,
)


class WorkspaceBackupTests(unittest.TestCase):
    @staticmethod
    def _add_workspace_data(
        database: Database,
        template_path: Path,
        *,
        preset_name: str,
        title: str,
    ) -> None:
        template_id = database.add_template(
            name="Summary",
            filename=template_path.name,
            file_path=str(template_path),
            report_type="summary",
        )
        preset_id = database.save_preset(
            name=preset_name,
            settings={"organization": preset_name},
            template_id=template_id,
        )
        database.add_report(
            title=title,
            report_type="summary",
            row_count=3,
            preset_id=preset_id,
            template_id=template_id,
        )

    @staticmethod
    def _write_template(path: Path, heading: str) -> bytes:
        path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        document.add_heading(heading, level=1)
        document.save(path)
        return path.read_bytes()

    def test_backup_contains_consistent_database_templates_and_manifest(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            database = Database(root / "data" / "reporter.db")
            database.initialize()

            template_dir = root / "templates"
            template_path = template_dir / "summary" / "summary.docx"
            template_path.parent.mkdir(parents=True)
            document = Document()
            document.add_heading("Summary", level=1)
            document.save(template_path)
            (template_dir / "~$summary.docx").write_bytes(b"word-lock")

            template_id = database.add_template(
                name="Summary",
                filename="summary.docx",
                file_path=str(template_path),
                report_type="summary",
            )
            preset_id = database.save_preset(
                name="Team preset",
                settings={"organization": "Example"},
                template_id=template_id,
            )
            database.add_report(
                title="Assessment",
                report_type="summary",
                row_count=3,
                preset_id=preset_id,
                template_id=template_id,
            )

            output = root / "backup.zip"
            manifest = create_workspace_backup(
                database,
                template_dir,
                output,
                app_version="test-version",
            )

            self.assertTrue(output.exists())
            self.assertEqual(manifest["schemaVersion"], 1)
            self.assertEqual(manifest["appVersion"], "test-version")
            self.assertEqual(manifest["database"]["schemaVersion"], 9)
            self.assertEqual(
                manifest["database"]["records"],
                {
                    "templates": 1,
                    "template_versions": 0,
                    "presets": 1,
                    "report_history": 1,
                    "detection_rules": 0,
                    "detection_rule_versions": 0,
                },
            )
            self.assertEqual(len(manifest["templates"]), 1)

            with zipfile.ZipFile(output) as archive:
                names = set(archive.namelist())
                self.assertEqual(
                    names,
                    {"database/reporter.db", "templates/summary/summary.docx", "manifest.json"},
                )
                archived_manifest = json.loads(archive.read("manifest.json"))
                self.assertEqual(archived_manifest, manifest)

                snapshot = root / "snapshot.db"
                snapshot.write_bytes(archive.read("database/reporter.db"))
                connection = sqlite3.connect(snapshot)
                try:
                    self.assertEqual(
                        connection.execute("SELECT name FROM presets").fetchone()[0],
                        "Team preset",
                    )
                finally:
                    connection.close()

            database.close()

    def test_dry_run_previews_then_restore_replaces_workspace(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source_db = Database(root / "source" / "reporter.db")
            source_db.initialize()
            source_templates = root / "source-templates"
            source_template = source_templates / "summary" / "source.docx"
            source_bytes = self._write_template(source_template, "Source template")
            self._add_workspace_data(
                source_db,
                source_template,
                preset_name="Restored preset",
                title="Restored report",
            )
            archive = root / "backup.zip"
            create_workspace_backup(source_db, source_templates, archive)

            target_db = Database(root / "target" / "reporter.db")
            target_db.initialize()
            target_templates = root / "target-templates"
            old_template = target_templates / "summary" / "old.docx"
            self._write_template(old_template, "Old template")
            self._add_workspace_data(
                target_db,
                old_template,
                preset_name="Old preset",
                title="Old report",
            )

            preview = inspect_workspace_backup(archive, target_db, target_templates)
            self.assertTrue(preview["valid"])
            self.assertTrue(preview["dryRun"])
            self.assertEqual(preview["database"]["records"]["presets"], 1)
            self.assertEqual(preview["templateCount"], 1)
            # Dry-run must not mutate current data.
            self.assertEqual(target_db.list_presets()[0]["name"], "Old preset")
            self.assertTrue(old_template.exists())

            result = restore_workspace_backup(
                archive,
                target_db,
                target_templates,
                confirmation_token=preview["confirmationToken"],
            )
            self.assertTrue(result["restored"])
            self.assertEqual(target_db.list_presets()[0]["name"], "Restored preset")
            restored_template = target_templates / "summary" / "source.docx"
            self.assertEqual(restored_template.read_bytes(), source_bytes)
            self.assertFalse(old_template.exists())
            self.assertEqual(
                Path(target_db.list_templates()[0]["file_path"]).resolve(),
                restored_template.resolve(),
            )
            source_db.close()
            target_db.close()

    def test_restore_rejects_changed_archive_after_dry_run(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            database = Database(root / "reporter.db")
            database.initialize()
            templates = root / "templates"
            self._write_template(templates / "summary.docx", "Summary")
            archive = root / "backup.zip"
            create_workspace_backup(database, templates, archive)
            preview = inspect_workspace_backup(archive, database, templates)

            archive.write_bytes(archive.read_bytes() + b"changed-after-preview")
            with self.assertRaisesRegex(WorkspaceBackupError, "changed after dry-run"):
                restore_workspace_backup(
                    archive,
                    database,
                    templates,
                    confirmation_token=preview["confirmationToken"],
                )
            database.close()

    def test_dry_run_rejects_checksum_mismatch(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            database = Database(root / "reporter.db")
            database.initialize()
            archive = root / "backup.zip"
            create_workspace_backup(database, root / "templates", archive)
            broken = root / "broken.zip"
            with zipfile.ZipFile(archive) as source, zipfile.ZipFile(broken, "w") as target:
                for info in source.infolist():
                    payload = source.read(info.filename)
                    if info.filename == "database/reporter.db":
                        payload = payload[:-1] + bytes([payload[-1] ^ 0xFF])
                    target.writestr(info.filename, payload)
            with self.assertRaisesRegex(WorkspaceBackupError, "Checksum mismatch"):
                inspect_workspace_backup(broken, database, root / "templates")
            database.close()

    def test_failed_template_install_rolls_back_database_and_templates(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source_db = Database(root / "source.db")
            source_db.initialize()
            source_templates = root / "source-templates"
            source_template = source_templates / "source.docx"
            self._write_template(source_template, "Source")
            self._add_workspace_data(
                source_db,
                source_template,
                preset_name="Source preset",
                title="Source report",
            )
            archive = root / "backup.zip"
            create_workspace_backup(source_db, source_templates, archive)

            target_db = Database(root / "target.db")
            target_db.initialize()
            target_templates = root / "target-templates"
            target_template = target_templates / "target.docx"
            original_template_bytes = self._write_template(target_template, "Target")
            self._add_workspace_data(
                target_db,
                target_template,
                preset_name="Target preset",
                title="Target report",
            )
            preview = inspect_workspace_backup(archive, target_db, target_templates)

            def fail_install(_source: Path, destination: Path) -> dict[str, Path]:
                target_template.write_bytes(b"partial-write")
                raise OSError(f"simulated failure in {destination}")

            with patch("core.workspace_backup._install_templates", side_effect=fail_install):
                with self.assertRaisesRegex(WorkspaceBackupError, "rolled back"):
                    restore_workspace_backup(
                        archive,
                        target_db,
                        target_templates,
                        confirmation_token=preview["confirmationToken"],
                    )

            self.assertEqual(target_db.list_presets()[0]["name"], "Target preset")
            self.assertEqual(target_template.read_bytes(), original_template_bytes)
            source_db.close()
            target_db.close()

    def test_restore_migrates_an_older_supported_database_schema(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source_db = Database(root / "source.db")
            source_db.initialize()
            source_db._execute_commit("DROP TABLE detection_rule_versions")
            source_db._execute_commit("DELETE FROM schema_migrations WHERE version >= 7")
            archive = root / "old-schema.zip"
            create_workspace_backup(source_db, root / "source-templates", archive)

            target_db = Database(root / "target.db")
            target_db.initialize()
            preview = inspect_workspace_backup(archive, target_db, root / "target-templates")
            self.assertEqual(preview["database"]["schemaVersion"], 6)
            self.assertTrue(any("migrated" in item for item in preview["warnings"]))

            restored = restore_workspace_backup(
                archive,
                target_db,
                root / "target-templates",
                confirmation_token=preview["confirmationToken"],
            )
            self.assertEqual(restored["database"]["sourceSchemaVersion"], 6)
            self.assertEqual(restored["database"]["schemaVersion"], 9)
            self.assertEqual(target_db.schema_version, 9)
            source_db.close()
            target_db.close()


if __name__ == "__main__":
    unittest.main()
