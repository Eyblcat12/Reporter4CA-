from __future__ import annotations

import hashlib
import json
import os
import re
import sys
import tempfile
import unicodedata
import unittest
import zipfile
from copy import deepcopy
from pathlib import Path, PurePosixPath
from xml.etree import ElementTree

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "apps" / "backend"
GOLDEN_ROOT = ROOT / "tests" / "golden" / "docx-v1"
TEMPLATE_ROOT = BACKEND / "templates"
sys.path.insert(0, str(BACKEND))

from core.report_generator import ReportType, generate_report
from tests.docx_golden_report import structural_diff, write_diff_reports


REPORT_TYPES = [item.value for item in ReportType]
REPORT_TEMPLATE_PATHS = {
    "full": TEMPLATE_ROOT / "report_template.docx",
    "server_only": TEMPLATE_ROOT / "server_only" / "report_server_only_default.docx",
    "client_only": TEMPLATE_ROOT / "client_only" / "report_client_only_default.docx",
    "summary": TEMPLATE_ROOT / "summary" / "report_summary_default.docx",
    "technical": TEMPLATE_ROOT / "technical" / "report_technical_default.docx",
    # The runtime currently has no category-specific IR template and deliberately
    # falls back to the full default template for this report type.
    "incident_response": TEMPLATE_ROOT / "report_template.docx",
}
TOKEN_PATTERN = re.compile(r"\{\{[^{}]+\}\}")
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
RELATIONSHIP_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
OFFICE_RELATIONSHIP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
XML_NAMESPACES = {"w": WORD_NS, "pr": RELATIONSHIP_NS}
GOLDEN_REPORT_DIR = Path(
    os.getenv("GOLDEN_DOCX_REPORT_DIR", str(ROOT / "artifacts" / "golden-docx"))
)


def fixture_data() -> dict:
    return {
        "servers": [{
            "hostname": "SRV-GOLDEN-01", "ip": "10.10.0.10", "os": "Windows Server 2022",
            "result": "Malware detected", "notes": "SHA256 evidence linked to T1055",
        }],
        "clients": [{
            "hostname": "WS-GOLDEN-01", "ip": "10.10.0.20", "os": "Windows 11",
            "result": "No finding", "notes": "Validated by endpoint telemetry",
        }],
        "metadata": {
            "incidentName": "Golden incident", "severity": "high",
            "timeline": [{"time": "2026-07-20T10:00:00Z", "event": "Detection", "evidence": "EDR-001"}],
        },
    }


def _normalize_text(value: object) -> str:
    normalized = unicodedata.normalize(
        "NFC",
        str(value or "").replace("\r\n", "\n").replace("\r", "\n"),
    )
    return "\n".join(line.rstrip() for line in normalized.split("\n")).strip()


def _canonical_digest(value: object) -> str:
    payload = json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


_VOLATILE_FORMAT_ATTRIBUTES = {
    "paraid",
    "textid",
}
_REVISION_PROPERTY_ELEMENTS = {
    "pprchange",
    "rprchange",
    "tblprchange",
    "tcprchange",
    "trprchange",
}


def _qualified_name(value: object) -> str:
    qualified = etree.QName(value)
    return f"{{{qualified.namespace or ''}}}{qualified.localname}"


def _format_property_projection(element) -> dict[str, object] | None:
    """Return canonical OOXML properties without text or revision metadata."""

    if element is None:
        return None

    attributes = []
    for name, value in element.attrib.items():
        qualified = etree.QName(name)
        local_name = qualified.localname.casefold()
        if (
            local_name.startswith("rsid")
            or local_name in _VOLATILE_FORMAT_ATTRIBUTES
        ):
            continue
        attributes.append((_qualified_name(name), value))

    children = []
    for child in element:
        if etree.QName(child).localname.casefold() in _REVISION_PROPERTY_ELEMENTS:
            continue
        children.append(_format_property_projection(child))

    return {
        "tag": _qualified_name(element.tag),
        "attributes": sorted(attributes),
        "children": children,
    }


def _table_format_projection(table) -> dict[str, object]:
    """Project table formatting while excluding volatile report text."""

    rows = []
    for row in table._tbl.findall(qn("w:tr")):
        cells = []
        for cell in row.findall(qn("w:tc")):
            paragraphs = []
            for paragraph in cell.findall(qn("w:p")):
                run_properties = []
                for run in paragraph.iter(qn("w:r")):
                    properties = _format_property_projection(
                        run.find(qn("w:rPr"))
                    )
                    # Splitting text into adjacent, identically styled runs is
                    # not a visual format change and must not churn goldens.
                    if not run_properties or run_properties[-1] != properties:
                        run_properties.append(properties)
                paragraphs.append({
                    "pPr": _format_property_projection(
                        paragraph.find(qn("w:pPr"))
                    ),
                    "rPr": run_properties,
                })
            cells.append({
                "tcPr": _format_property_projection(cell.find(qn("w:tcPr"))),
                "paragraphs": paragraphs,
            })
        rows.append({
            "trPr": _format_property_projection(row.find(qn("w:trPr"))),
            "cells": cells,
        })

    return {
        "tblPr": _format_property_projection(table._tbl.find(qn("w:tblPr"))),
        "tblGrid": _format_property_projection(
            table._tbl.find(qn("w:tblGrid"))
        ),
        "rows": rows,
    }


def _table_snapshot(table) -> dict:
    cells = [
        [_normalize_text(cell.text) for cell in row.cells]
        for row in table.rows
    ]
    return {
        "rows": len(cells),
        "columns": max((len(row) for row in cells), default=0),
        "header": cells[0] if cells else [],
        "nonEmptyCells": sum(bool(value) for row in cells for value in row),
        "contentDigest": _canonical_digest(cells),
        "rowDigests": [_canonical_digest(row) for row in cells],
        "formatDigest": _canonical_digest(_table_format_projection(table)),
    }


def _append_ooxml(
    parent,
    tag: str,
    attributes: dict[str, str] | None = None,
):
    element = OxmlElement(tag)
    for name, value in (attributes or {}).items():
        element.set(qn(name), value)
    parent.append(element)
    return element


def _element_attributes(parent, tag: str, names: tuple[str, ...]) -> dict[str, str]:
    element = parent.find(qn(f"w:{tag}"))
    if element is None:
        return {}
    return {
        name: value
        for name in names
        if (value := element.get(qn(f"w:{name}"))) is not None
    }


def _section_properties(section) -> dict:
    section_properties = section._sectPr
    normalized = deepcopy(section_properties)
    for element in normalized.iter():
        for attribute_name in list(element.attrib):
            qualified = etree.QName(attribute_name)
            if (
                qualified.namespace == OFFICE_RELATIONSHIP_NS
                or qualified.localname.casefold().startswith("rsid")
            ):
                del element.attrib[attribute_name]

    header_types = sorted(
        item.get(qn("w:type"), "default")
        for item in section_properties.findall(qn("w:headerReference"))
    )
    footer_types = sorted(
        item.get(qn("w:type"), "default")
        for item in section_properties.findall(qn("w:footerReference"))
    )
    return {
        "type": _element_attributes(section_properties, "type", ("val",)),
        "pageSize": _element_attributes(
            section_properties,
            "pgSz",
            ("w", "h", "orient", "code"),
        ),
        "pageMargins": _element_attributes(
            section_properties,
            "pgMar",
            ("top", "right", "bottom", "left", "header", "footer", "gutter"),
        ),
        "columns": _element_attributes(
            section_properties,
            "cols",
            ("num", "space", "sep", "equalWidth"),
        ),
        "pageNumbering": _element_attributes(
            section_properties,
            "pgNumType",
            ("fmt", "start", "chapStyle", "chapSep"),
        ),
        "titlePage": section_properties.find(qn("w:titlePg")) is not None,
        "headerTypes": header_types,
        "footerTypes": footer_types,
        "propertiesDigest": hashlib.sha256(
            etree.tostring(normalized, method="c14n")
        ).hexdigest(),
    }


def _xml_part_content(root: ElementTree.Element) -> dict:
    paragraphs = [
        _normalize_text(
            "".join(
                node.text or ""
                for node in paragraph.findall(".//w:t", XML_NAMESPACES)
            )
        )
        for paragraph in root.findall(".//w:p", XML_NAMESPACES)
    ]
    tables: list[list[list[str]]] = []
    for table in root.findall(".//w:tbl", XML_NAMESPACES):
        rows: list[list[str]] = []
        for row in table.findall("./w:tr", XML_NAMESPACES):
            rows.append([
                _normalize_text(
                    "".join(
                        node.text or ""
                        for node in cell.findall(".//w:t", XML_NAMESPACES)
                    )
                )
                for cell in row.findall("./w:tc", XML_NAMESPACES)
            ])
        tables.append(rows)
    instructions = [
        _normalize_text(node.text)
        for node in root.findall(".//w:instrText", XML_NAMESPACES)
        if _normalize_text(node.text)
    ]
    return {
        "paragraphs": paragraphs,
        "paragraphDigest": _canonical_digest(paragraphs),
        "tableCount": len(tables),
        "tableCellDigest": _canonical_digest(tables),
        "fieldInstructions": instructions,
    }


def _header_footer_snapshot(archive: zipfile.ZipFile) -> list[dict[str, object]]:
    parts: list[dict[str, object]] = []
    names = sorted(
        name
        for name in archive.namelist()
        if re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
    )
    for name in names:
        root = ElementTree.fromstring(archive.read(name))
        parts.append({
            "part": name,
            "kind": (
                "header"
                if PurePosixPath(name).name.startswith("header")
                else "footer"
            ),
            **_xml_part_content(root),
        })
    return parts


def _relationship_source(part_name: str) -> str:
    path = PurePosixPath(part_name)
    if part_name == "_rels/.rels":
        return "/"
    if path.parent.name != "_rels" or not path.name.endswith(".rels"):
        return part_name
    return str(path.parent.parent / path.name.removesuffix(".rels"))


def _relationship_snapshot(archive: zipfile.ZipFile) -> list[dict[str, str]]:
    relationships: list[dict[str, str]] = []
    for name in sorted(item for item in archive.namelist() if item.endswith(".rels")):
        root = ElementTree.fromstring(archive.read(name))
        for relationship in root.findall("pr:Relationship", XML_NAMESPACES):
            relationships.append({
                "source": _relationship_source(name),
                "type": relationship.get("Type", ""),
                "target": relationship.get("Target", "").replace("\\", "/"),
                "targetMode": relationship.get("TargetMode", "Internal"),
            })
    return sorted(
        relationships,
        key=lambda item: (
            item["source"],
            item["type"],
            item["target"],
            item["targetMode"],
        ),
    )


def _field_snapshot(archive: zipfile.ZipFile) -> dict[str, object]:
    document_root = ElementTree.fromstring(archive.read("word/document.xml"))
    instructions = [
        _normalize_text(node.text)
        for node in document_root.findall(".//w:instrText", XML_NAMESPACES)
        if _normalize_text(node.text)
    ]
    instructions.extend(
        _normalize_text(node.get(qn("w:instr")))
        for node in document_root.findall(".//w:fldSimple", XML_NAMESPACES)
        if _normalize_text(node.get(qn("w:instr")))
    )
    toc_instructions = sorted(
        instruction
        for instruction in instructions
        if re.match(r"^TOC(?:\s|$)", instruction, flags=re.IGNORECASE)
    )
    dirty_fields = sum(
        str(node.get(qn("w:dirty"), "")).casefold() in {"1", "on", "true", "yes"}
        for node in document_root.findall(".//w:fldChar", XML_NAMESPACES)
    )

    update_fields = False
    if "word/settings.xml" in archive.namelist():
        settings_root = ElementTree.fromstring(archive.read("word/settings.xml"))
        update = settings_root.find(".//w:updateFields", XML_NAMESPACES)
        if update is not None:
            update_fields = str(update.get(qn("w:val"), "true")).casefold() not in {
                "0",
                "false",
                "off",
                "no",
            }
    return {
        "tocCount": len(toc_instructions),
        "tocInstructions": toc_instructions,
        "updateFields": update_fields,
        "dirtyFieldCount": dirty_fields,
    }


def _semantic_counts(manifest: dict[str, object] | None) -> dict[str, object]:
    source = manifest or {}
    return {
        "reportType": str(source.get("reportType", "")),
        "assetCount": int(source.get("assetCount", 0)),
        "findingCount": int(source.get("findingCount", 0)),
        "evidenceCount": int(source.get("evidenceCount", 0)),
        "assetTypeCounts": dict(
            sorted(dict(source.get("assetTypeCounts", {})).items())
        ),
        "assessmentCounts": dict(
            sorted(dict(source.get("assessmentCounts", {})).items())
        ),
        "ruleCounts": dict(sorted(dict(source.get("ruleCounts", {})).items())),
    }


def document_snapshot(document, manifest: dict[str, object] | None = None) -> dict:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temporary:
        output_path = Path(temporary.name)
    try:
        document.save(output_path)
        reopened = Document(output_path)
        paragraphs = [
            _normalize_text(paragraph.text)
            for paragraph in reopened.paragraphs
            if _normalize_text(paragraph.text)
        ]
        headings = [
            {"style": paragraph.style.name, "text": _normalize_text(paragraph.text)}
            for paragraph in reopened.paragraphs
            if (
                _normalize_text(paragraph.text)
                and paragraph.style
                and paragraph.style.name.startswith("Heading")
            )
        ]
        tables = [_table_snapshot(table) for table in reopened.tables]
        numbered = sum(
            1 for paragraph in reopened.paragraphs
            if paragraph._p.pPr is not None
            and paragraph._p.pPr.numPr is not None
        )
        sections = [_section_properties(section) for section in reopened.sections]

        with zipfile.ZipFile(output_path) as archive:
            media = []
            for name in archive.namelist():
                if name.startswith("word/media/") and not name.endswith("/"):
                    content = archive.read(name)
                    media.append({
                        "name": PurePosixPath(name).name,
                        "size": len(content),
                        "sha256": hashlib.sha256(content).hexdigest(),
                    })
            headers_footers = _header_footer_snapshot(archive)
            relationships = _relationship_snapshot(archive)
            fields = _field_snapshot(archive)
    finally:
        output_path.unlink(missing_ok=True)

    return {
        "snapshotSchemaVersion": 3,
        "paragraphCount": len(paragraphs),
        "paragraphDigest": hashlib.sha256(
            "\n".join(paragraphs).encode("utf-8")
        ).hexdigest(),
        "headings": headings,
        "tables": tables,
        "numberedParagraphs": numbered,
        "remainingTokens": sorted(set(TOKEN_PATTERN.findall("\n".join(paragraphs)))),
        "sectionCount": len(sections),
        "sections": sections,
        "headersFooters": headers_footers,
        "relationships": relationships,
        "media": sorted(media, key=lambda item: item["name"]),
        "tocFields": fields,
        "semanticCounts": _semantic_counts(manifest),
    }


def readable_diff(expected: object, actual: object, path: str = "root") -> list[str]:
    markers = {"added": "+", "removed": "-", "changed": "~"}
    return [
        f"{markers[item['change']]} {item['path']}: "
        f"expected {item.get('expected')!r}; actual {item.get('actual')!r}"
        for item in structural_diff(expected, actual, path)
    ]


class DocxGoldenTests(unittest.TestCase):
    def test_snapshot_is_stable_and_separates_content_from_format(self) -> None:
        document = Document()
        document.sections[0].header.paragraphs[0].text = "Reporter Pro"
        table = document.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Hostname"
        table.rows[0].cells[1].text = "Result"
        table.rows[1].cells[0].text = "SRV-01"
        table.rows[1].cells[1].text = "Clean"
        manifest = {
            "reportType": "full",
            "assetCount": 1,
            "findingCount": 0,
            "evidenceCount": 0,
            "assetTypeCounts": {"server": 1},
            "assessmentCounts": {"clean": 1},
            "ruleCounts": {},
        }

        first = document_snapshot(document, manifest)
        second = document_snapshot(document, manifest)
        self.assertEqual(first, second)

        table.rows[1].cells[1].text = "Anomaly"
        changed = document_snapshot(document, manifest)
        self.assertNotEqual(
            first["tables"][0]["contentDigest"],
            changed["tables"][0]["contentDigest"],
        )
        self.assertEqual(
            first["tables"][0]["formatDigest"],
            changed["tables"][0]["formatDigest"],
        )

    def test_table_format_digest_and_diff_cover_ooxml_property_layers(self) -> None:
        def make_document():
            document = Document()
            table = document.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Hostname"
            table.rows[0].cells[1].text = "Result"
            table.rows[1].cells[0].text = "SRV-01"
            table.rows[1].cells[1].text = "Clean"
            return document, table

        def add_border(table) -> None:
            borders = _append_ooxml(table._tbl.tblPr, "w:tblBorders")
            _append_ooxml(
                borders,
                "w:top",
                {"w:val": "single", "w:sz": "8", "w:color": "445566"},
            )

        def change_grid_width(table) -> None:
            grid_column = table._tbl.tblGrid.find(qn("w:gridCol"))
            self.assertIsNotNone(grid_column)
            grid_column.set(qn("w:w"), "4321")

        def add_row_property(table) -> None:
            _append_ooxml(table.rows[1]._tr.get_or_add_trPr(), "w:cantSplit")

        def change_cell_width(table) -> None:
            tc_pr = table.rows[1].cells[0]._tc.get_or_add_tcPr()
            width = tc_pr.find(qn("w:tcW"))
            self.assertIsNotNone(width)
            width.set(qn("w:type"), "dxa")
            width.set(qn("w:w"), "3333")

        def add_numbering(table) -> None:
            p_pr = table.rows[1].cells[0].paragraphs[0]._p.get_or_add_pPr()
            num_pr = _append_ooxml(p_pr, "w:numPr")
            _append_ooxml(num_pr, "w:ilvl", {"w:val": "0"})
            _append_ooxml(num_pr, "w:numId", {"w:val": "7"})

        def change_run_format(table) -> None:
            table.rows[1].cells[1].paragraphs[0].runs[0].bold = True

        def add_merge(table) -> None:
            table.rows[1].cells[0].merge(table.rows[1].cells[1])

        def add_shading(table) -> None:
            tc_pr = table.rows[1].cells[0]._tc.get_or_add_tcPr()
            _append_ooxml(tc_pr, "w:shd", {"w:fill": "FFF2CC"})

        mutations = {
            "tblPr/border": add_border,
            "tblGrid/width": change_grid_width,
            "trPr": add_row_property,
            "tcPr/width": change_cell_width,
            "pPr/numbering": add_numbering,
            "rPr": change_run_format,
            "merge": add_merge,
            "shading": add_shading,
        }

        for name, mutate in mutations.items():
            with self.subTest(property_layer=name):
                document, table = make_document()
                baseline = document_snapshot(document)
                mutate(table)
                changed = document_snapshot(document)

                self.assertNotEqual(
                    baseline["tables"][0]["formatDigest"],
                    changed["tables"][0]["formatDigest"],
                )
                differences = structural_diff(baseline, changed)
                format_differences = [
                    item
                    for item in differences
                    if item["path"] == "root.tables[0].formatDigest"
                ]
                self.assertEqual(len(format_differences), 1)
                self.assertEqual(format_differences[0]["category"], "Tables")

    def test_template_matrix_covers_every_report_type(self) -> None:
        self.assertEqual(set(REPORT_TEMPLATE_PATHS), set(REPORT_TYPES))
        for report_type, template_path in REPORT_TEMPLATE_PATHS.items():
            with self.subTest(report_type=report_type):
                self.assertTrue(
                    template_path.is_file(),
                    f"Missing default template for {report_type}: {template_path}",
                )

    def test_all_report_types_match_structural_golden_files(self) -> None:
        update = os.getenv("UPDATE_GOLDEN_DOCX") == "1"
        regressions: list[dict[str, object]] = []
        if update:
            GOLDEN_ROOT.mkdir(parents=True, exist_ok=True)

        for report_type in REPORT_TYPES:
            with self.subTest(report_type=report_type):
                document = generate_report(
                    fixture_data(), title="Golden Report", organization="Reporter Team",
                    assessment_date="2026-07-20",
                    template_path=REPORT_TEMPLATE_PATHS[report_type],
                    report_type=report_type,
                )
                actual = document_snapshot(
                    document,
                    getattr(document, "_reporter_manifest", {}),
                )
                golden_path = GOLDEN_ROOT / f"{report_type}.json"
                if update:
                    golden_path.write_text(
                        json.dumps(actual, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
                    )
                self.assertTrue(golden_path.exists(), f"Missing golden file: {golden_path}")
                expected = json.loads(golden_path.read_text(encoding="utf-8"))
                structured = structural_diff(expected, actual)
                differences = readable_diff(expected, actual)
                if differences:
                    regressions.append({"reportType": report_type, "differences": structured})

        if regressions:
            json_path, html_path = write_diff_reports(regressions, GOLDEN_REPORT_DIR)
            self.fail(
                "DOCX structural regression detected:\n"
                + "\n".join(
                    f"- {item['reportType']}: {len(item['differences'])} differences"
                    for item in regressions
                )
                + f"\nJSON report: {json_path}\nHTML report: {html_path}"
                + "\nUse UPDATE_GOLDEN_DOCX=1 only after reviewing this diff."
            )

    def test_structured_diff_report_is_granular_and_readable(self) -> None:
        expected = {
            "headings": [{"style": "Heading 1", "text": "Overview"}],
            "tables": [{"rows": 2, "columns": 3}],
            "remainingTokens": [],
            "media": [{"name": "chart.png", "size": 100}],
            "sections": [{"pageSize": {"w": "11906"}}],
            "headersFooters": [
                {"part": "word/header1.xml", "paragraphs": ["Reporter"]}
            ],
            "relationships": [{"type": "image", "target": "media/logo.png"}],
            "tocFields": {"tocCount": 1, "updateFields": True},
            "semanticCounts": {"findingCount": 1, "evidenceCount": 1},
        }
        actual = {
            "headings": [{"style": "Heading 1", "text": "Summary"}],
            "tables": [{"rows": 3, "columns": 3}],
            "remainingTokens": ["{{UNRESOLVED}}"],
            "media": [],
            "sections": [{"pageSize": {"w": "12240"}}],
            "headersFooters": [
                {"part": "word/header1.xml", "paragraphs": ["Changed"]}
            ],
            "relationships": [{"type": "image", "target": "media/new-logo.png"}],
            "tocFields": {"tocCount": 0, "updateFields": False},
            "semanticCounts": {"findingCount": 2, "evidenceCount": 3},
        }
        differences = structural_diff(expected, actual)
        self.assertEqual(
            {item["category"] for item in differences},
            {
                "Findings/Evidence",
                "Header/Footer",
                "Headings",
                "Media",
                "Relationships",
                "Sections",
                "Tables",
                "TOC/Fields",
                "Tokens",
            },
        )
        self.assertIn("root.headings[0].text", {item["path"] for item in differences})
        self.assertIn("root.remainingTokens[0]", {item["path"] for item in differences})

        with tempfile.TemporaryDirectory() as directory:
            json_path, html_path = write_diff_reports(
                [{"reportType": "technical", "differences": differences}], Path(directory)
            )
            report = json.loads(json_path.read_text(encoding="utf-8"))
            rendered = html_path.read_text(encoding="utf-8")
            self.assertEqual(report["status"], "failed")
            self.assertEqual(report["summary"]["byCategory"]["Headings"], 1)
            self.assertIn("Golden DOCX structural diff", rendered)
            self.assertIn("root.headings[0].text", rendered)


if __name__ == "__main__":
    unittest.main()
