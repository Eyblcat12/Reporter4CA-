from __future__ import annotations

import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "apps" / "backend"
sys.path.insert(0, str(BACKEND))

from core.docx_field_updater import refresh_docx_fields  # noqa: E402
from core.report_generator import _enable_field_updates  # noqa: E402


class DocxFieldUpdaterTests(unittest.TestCase):
    def test_enable_field_updates_marks_fields_dirty(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        field_begin = OxmlElement("w:fldChar")
        field_begin.set(qn("w:fldCharType"), "begin")
        paragraph.add_run()._r.append(field_begin)

        _enable_field_updates(document)

        update_fields = document.settings.element.find(qn("w:updateFields"))
        self.assertIsNotNone(update_fields)
        self.assertEqual(update_fields.get(qn("w:val")), "true")
        self.assertEqual(field_begin.get(qn("w:dirty")), "true")

    def test_enable_field_updates_overrides_disabled_setting(self) -> None:
        document = Document()
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "false")
        document.settings.element.append(update_fields)

        _enable_field_updates(document)

        matches = document.settings.element.findall(qn("w:updateFields"))
        self.assertEqual(len(matches), 1)
        self.assertEqual(matches[0].get(qn("w:val")), "true")

    def test_non_windows_keeps_original_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "report.docx"
            document = Document()
            document.add_paragraph("Original")
            document.save(path)
            original = path.read_bytes()

            with patch("core.docx_field_updater.platform.system", return_value="Linux"):
                result = refresh_docx_fields(path)

            self.assertFalse(result.updated)
            self.assertEqual(result.engine, "deferred")
            self.assertEqual(path.read_bytes(), original)

    def test_failed_word_update_keeps_original_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "report.docx"
            document = Document()
            document.add_paragraph("Original")
            document.save(path)
            original = path.read_bytes()

            failed = subprocess.CompletedProcess(
                [], returncode=1, stdout="", stderr="Word unavailable"
            )
            with (
                patch("core.docx_field_updater.platform.system", return_value="Windows"),
                patch("core.docx_field_updater.shutil.which", return_value="powershell.exe"),
                patch("core.docx_field_updater._run_word_update", return_value=failed),
            ):
                result = refresh_docx_fields(path)

            self.assertFalse(result.updated)
            self.assertEqual(path.read_bytes(), original)
            self.assertFalse(any(path.parent.glob("*.field-refresh-*.docx")))

    def test_successful_word_update_replaces_file_atomically(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "report.docx"
            document = Document()
            paragraph = document.add_paragraph("Original")
            field_begin = OxmlElement("w:fldChar")
            field_begin.set(qn("w:fldCharType"), "begin")
            paragraph.add_run()._r.append(field_begin)
            _enable_field_updates(document)
            document.save(path)

            def update_copy(_powershell, working_copy, *, timeout_seconds):
                updated = Document(working_copy)
                updated.add_paragraph("Updated by Word")
                updated.save(working_copy)
                return subprocess.CompletedProcess([], returncode=0, stdout="", stderr="")

            with (
                patch("core.docx_field_updater.platform.system", return_value="Windows"),
                patch("core.docx_field_updater.shutil.which", return_value="powershell.exe"),
                patch("core.docx_field_updater._run_word_update", side_effect=update_copy),
            ):
                result = refresh_docx_fields(path)

            self.assertTrue(result.updated)
            self.assertEqual(result.engine, "microsoft-word")
            self.assertTrue(zipfile.is_zipfile(path))
            self.assertIn("Updated by Word", "\n".join(p.text for p in Document(path).paragraphs))
            with zipfile.ZipFile(path) as archive:
                settings = archive.read("word/settings.xml").decode("utf-8")
                document_xml = archive.read("word/document.xml").decode("utf-8")
            self.assertIn('w:updateFields w:val="false"', settings)
            self.assertNotIn('w:updateFields w:val="true"', settings)
            self.assertNotIn('w:dirty="true"', document_xml)

    def test_failed_word_update_preserves_update_on_open_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "report.docx"
            document = Document()
            paragraph = document.add_paragraph("Original")
            field_begin = OxmlElement("w:fldChar")
            field_begin.set(qn("w:fldCharType"), "begin")
            paragraph.add_run()._r.append(field_begin)
            _enable_field_updates(document)
            document.save(path)

            failed = subprocess.CompletedProcess(
                [], returncode=1, stdout="", stderr="Word unavailable"
            )
            with (
                patch("core.docx_field_updater.platform.system", return_value="Windows"),
                patch("core.docx_field_updater.shutil.which", return_value="cscript.exe"),
                patch("core.docx_field_updater._run_word_update", return_value=failed),
            ):
                result = refresh_docx_fields(path)

            self.assertFalse(result.updated)
            with zipfile.ZipFile(path) as archive:
                settings = archive.read("word/settings.xml").decode("utf-8")
                document_xml = archive.read("word/document.xml").decode("utf-8")
            self.assertIn('w:updateFields w:val="true"', settings)
            self.assertIn('w:dirty="true"', document_xml)


if __name__ == "__main__":
    unittest.main()
