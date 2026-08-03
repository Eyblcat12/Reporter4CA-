from __future__ import annotations

import sys
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "apps" / "backend"
sys.path.insert(0, str(BACKEND))

from core.template_blueprint import (  # noqa: E402
    UnsafeTableBlueprintError,
    classify_table_for_fast_path,
    compile_table_blueprint,
)


def simple_table(*, rows: int = 8, columns: int = 4):
    document = Document()
    table = document.add_table(rows=rows, cols=columns)
    for column, cell in enumerate(table.rows[0].cells):
        cell.text = f"Header {column + 1}"
        cell.paragraphs[0].runs[0].bold = True
    for row_index, row in enumerate(table.rows[1:], start=1):
        for column, cell in enumerate(row.cells):
            cell.text = f"R{row_index}C{column + 1}"
    return document, table


def add_cell_element(table, tag: str, *, relationship: bool = False) -> None:
    element = OxmlElement(tag)
    if relationship:
        element.set(qn("r:id"), "rId-test")
    table.rows[1].cells[0]._tc.append(element)


class TemplateBlueprintTests(unittest.TestCase):
    def test_compacts_large_simple_table_to_required_parts(self) -> None:
        _, table = simple_table(rows=101)

        blueprint = compile_table_blueprint(table)
        compact = blueprint.to_table_element()
        child_tags = [child.tag for child in compact]

        self.assertEqual(blueprint.column_count, 4)
        self.assertEqual(blueprint.data_row_variant_count, 1)
        self.assertFalse(blueprint.requires_integration_fallback)
        self.assertEqual(
            child_tags,
            [qn("w:tblPr"), qn("w:tblGrid"), qn("w:tr"), qn("w:tr")],
        )
        self.assertLess(blueprint.compact_xml_bytes, blueprint.source_xml_bytes / 5)
        self.assertLess(blueprint.compact_node_count, blueprint.source_node_count / 10)
        self.assertRegex(blueprint.fingerprint, r"^[0-9a-f]{64}$")

    def test_distinct_row_formatting_keeps_only_required_variants(self) -> None:
        _, table = simple_table(rows=11)
        for row_index, row in enumerate(table.rows[1:], start=1):
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EEEEEE" if row_index % 2 else "FFFFFF")
            row.cells[0]._tc.get_or_add_tcPr().append(shading)

        blueprint = compile_table_blueprint(table)

        self.assertEqual(blueprint.data_row_variant_count, 2)
        self.assertTrue(blueprint.requires_integration_fallback)
        self.assertEqual(len(set(blueprint.data_row_variant_fingerprints)), 2)
        self.assertLess(blueprint.compact_node_count, blueprint.source_node_count)

    def test_fingerprint_ignores_sample_data_but_tracks_format(self) -> None:
        _, first = simple_table(rows=4)
        _, second = simple_table(rows=4)
        second.rows[1].cells[1].text = "Different prototype text"

        first_blueprint = compile_table_blueprint(first)
        second_blueprint = compile_table_blueprint(second)
        self.assertEqual(first_blueprint.fingerprint, second_blueprint.fingerprint)

        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "FFF2CC")
        second.rows[1].cells[0]._tc.get_or_add_tcPr().append(shading)
        formatted_blueprint = compile_table_blueprint(second)
        self.assertNotEqual(first_blueprint.fingerprint, formatted_blueprint.fingerprint)

    def test_identically_styled_multiple_runs_remain_safe(self) -> None:
        _, table = simple_table(rows=3)
        paragraph = table.rows[1].cells[0].paragraphs[0]
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        first = paragraph.add_run("same")
        second = paragraph.add_run(" style")
        first.bold = True
        second.bold = True

        classification = classify_table_for_fast_path(table)

        self.assertTrue(classification.safe)
        self.assertNotIn("mixed_run_styles", classification.reasons)

    def test_unsafe_structures_are_sent_to_legacy_fallback(self) -> None:
        cases = {}

        _, merged = simple_table(rows=3)
        merged.cell(1, 0).merge(merged.cell(1, 1))
        cases["merged_cell"] = merged

        _, vertical = simple_table(rows=3)
        vertical.cell(1, 0).merge(vertical.cell(2, 0))
        cases["vertical_merge"] = vertical

        _, hyperlink = simple_table(rows=3)
        add_cell_element(hyperlink, "w:hyperlink")
        cases["hyperlink"] = hyperlink

        _, field = simple_table(rows=3)
        field_element = OxmlElement("w:fldSimple")
        field_element.set(qn("w:instr"), "PAGE")
        field.rows[1].cells[0]._tc.append(field_element)
        cases["field_code"] = field

        _, content_control = simple_table(rows=3)
        add_cell_element(content_control, "w:sdt")
        cases["content_control"] = content_control

        _, mixed_runs = simple_table(rows=3)
        paragraph = mixed_runs.rows[1].cells[0].paragraphs[0]
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        paragraph.add_run("bold").bold = True
        paragraph.add_run("italic").italic = True
        cases["mixed_run_styles"] = mixed_runs

        _, relationship = simple_table(rows=3)
        add_cell_element(relationship, "w:drawing", relationship=True)
        cases["cell_relationship"] = relationship

        _, multiple_paragraphs = simple_table(rows=3)
        multiple_paragraphs.rows[1].cells[0].add_paragraph("Second paragraph")
        cases["multiple_paragraphs"] = multiple_paragraphs

        _, unknown = simple_table(rows=3)
        add_cell_element(unknown, "w:customXml")
        cases["unsupported_cell_content"] = unknown

        for expected_reason, table in cases.items():
            with self.subTest(reason=expected_reason):
                classification = classify_table_for_fast_path(table)
                self.assertTrue(classification.requires_fallback)
                self.assertEqual(classification.path, "legacy")
                self.assertIn(expected_reason, classification.reasons)
                with self.assertRaises(UnsafeTableBlueprintError):
                    compile_table_blueprint(table)

    def test_variant_limit_falls_back_instead_of_retaining_large_prototype(self) -> None:
        _, table = simple_table(rows=5)
        for index, row in enumerate(table.rows[1:], start=1):
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), f"{index:06X}")
            row.cells[0]._tc.get_or_add_tcPr().append(shading)

        with self.assertRaises(UnsafeTableBlueprintError) as raised:
            compile_table_blueprint(table, max_data_variants=2)

        self.assertEqual(
            raised.exception.classification.reasons,
            ("too_many_data_row_variants",),
        )

    def test_every_row_must_match_header_and_grid_column_count(self) -> None:
        cases = {}

        _, short_data_row = simple_table(rows=3, columns=4)
        data_row = short_data_row.rows[1]._tr
        data_row.remove(data_row.findall(qn("w:tc"))[-1])
        cases["short_data_row"] = short_data_row

        _, short_header = simple_table(rows=3, columns=4)
        header_row = short_header.rows[0]._tr
        header_row.remove(header_row.findall(qn("w:tc"))[-1])
        cases["header_grid_mismatch"] = short_header

        for case, table in cases.items():
            with self.subTest(case=case):
                classification = classify_table_for_fast_path(table)
                self.assertTrue(classification.requires_fallback)
                self.assertIn("column_count_mismatch", classification.reasons)
                with self.assertRaises(UnsafeTableBlueprintError):
                    compile_table_blueprint(table)

    def test_row_property_exception_and_unknown_row_content_fall_back(self) -> None:
        cases = {}

        _, property_exception = simple_table(rows=3)
        property_exception.rows[1]._tr.insert(0, OxmlElement("w:tblPrEx"))
        cases["row_property_exception"] = property_exception

        _, unknown_row = simple_table(rows=3)
        unknown_row.rows[1]._tr.append(OxmlElement("w:customXml"))
        cases["unsupported_row_content"] = unknown_row

        for expected_reason, table in cases.items():
            with self.subTest(reason=expected_reason):
                classification = classify_table_for_fast_path(table)
                self.assertTrue(classification.requires_fallback)
                self.assertIn(expected_reason, classification.reasons)
                with self.assertRaises(UnsafeTableBlueprintError):
                    compile_table_blueprint(table)

    def test_unknown_direct_table_content_falls_back(self) -> None:
        _, table = simple_table(rows=3)
        table._tbl.append(OxmlElement("w:customXml"))

        classification = classify_table_for_fast_path(table)

        self.assertTrue(classification.requires_fallback)
        self.assertIn("unsupported_table_content", classification.reasons)
        with self.assertRaises(UnsafeTableBlueprintError):
            compile_table_blueprint(table)

    def test_malformed_table_grid_falls_back(self) -> None:
        cases = {}

        _, missing_grid = simple_table(rows=3)
        grid = missing_grid._tbl.find(qn("w:tblGrid"))
        missing_grid._tbl.remove(grid)
        cases["missing_grid"] = missing_grid

        _, empty_grid = simple_table(rows=3)
        grid = empty_grid._tbl.find(qn("w:tblGrid"))
        for column in list(grid):
            grid.remove(column)
        cases["empty_grid"] = empty_grid

        _, unsupported_grid_child = simple_table(rows=3)
        grid = unsupported_grid_child._tbl.find(qn("w:tblGrid"))
        grid.append(OxmlElement("w:customXml"))
        cases["unsupported_grid_child"] = unsupported_grid_child

        for case, table in cases.items():
            with self.subTest(case=case):
                classification = classify_table_for_fast_path(table)
                self.assertTrue(classification.requires_fallback)
                self.assertIn("malformed_table_grid", classification.reasons)
                with self.assertRaises(UnsafeTableBlueprintError):
                    compile_table_blueprint(table)


if __name__ == "__main__":
    unittest.main()
