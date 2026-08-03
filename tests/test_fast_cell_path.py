from __future__ import annotations

import os
import sys
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "apps" / "backend"
sys.path.insert(0, str(BACKEND))

from core import report_generator as generator  # noqa: E402
from core.config import fast_cell_enabled  # noqa: E402
from core.performance_metrics import PerformanceMetrics  # noqa: E402


def _simple_cell():
    document = Document()
    cell = document.add_table(rows=1, cols=1).cell(0, 0)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("old")
    run.bold = True
    return cell, paragraph, run


def _fixture_data(count: int = 4) -> dict:
    return {
        "servers": [
            {
                "hostname": f"SRV-{index:03d}",
                "ip": f"10.0.0.{index}",
                "os": "Windows Server 2022",
                "result": "Không phát hiện",
                "notes": "",
                "findings": [],
            }
            for index in range(1, count + 1)
        ],
        "clients": [],
        "metadata": {},
    }


class FastCellPathTests(unittest.TestCase):
    def test_flag_defaults_on_and_supports_explicit_rollback(self) -> None:
        with patch.dict(os.environ, {}, clear=True):
            self.assertTrue(fast_cell_enabled())
        for value in ("1", "true", "YES", "on"):
            with self.subTest(value=value), patch.dict(
                os.environ, {"AUTO_REPORT_FAST_CELL": value}
            ):
                self.assertTrue(fast_cell_enabled())
        for value in ("0", "false", "no", "off", ""):
            with self.subTest(value=value), patch.dict(
                os.environ, {"AUTO_REPORT_FAST_CELL": value}
            ):
                self.assertFalse(fast_cell_enabled())

    def test_simple_cell_reuses_run_and_preserves_canonical_final_format(self) -> None:
        cell, paragraph, run = _simple_cell()
        generator._format_cell(cell, bold=False, centered=True)
        run_element = run._r
        p_pr_before = paragraph._p.pPr.xml

        with patch.dict(os.environ, {"AUTO_REPORT_FAST_CELL": "1"}):
            generator._set_cell_text(
                cell,
                "new value",
                bold=False,
                centered=False,
            )

        self.assertIs(paragraph.runs[0]._r, run_element)
        self.assertEqual(paragraph.text, "new value")
        self.assertEqual(paragraph._p.pPr.xml, p_pr_before)
        generator._format_cell(cell, bold=False, centered=True)
        self.assertIn('w:rFonts w:ascii="Times New Roman"', cell._tc.xml)
        self.assertNotIn('w:rFonts w:cs="Times New Roman"', cell._tc.xml)

    def test_multiple_runs_and_merged_cells_use_safe_fallback(self) -> None:
        cell, paragraph, first_run = _simple_cell()
        generator._format_cell(cell, bold=False, centered=True)
        paragraph.add_run("second").italic = True
        old_run_element = first_run._r
        with patch.dict(os.environ, {"AUTO_REPORT_FAST_CELL": "1"}):
            generator._set_cell_text(cell, "fallback", bold=False, centered=False)
        self.assertEqual(paragraph.text, "fallback")
        self.assertEqual(len(paragraph.runs), 1)
        self.assertIsNot(paragraph.runs[0]._r, old_run_element)

        document = Document()
        table = document.add_table(rows=1, cols=2)
        merged = table.cell(0, 0).merge(table.cell(0, 1))
        merged_paragraph = merged.paragraphs[0]
        merged_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        merged_run = merged_paragraph.add_run("merged")
        merged_run.bold = True
        generator._format_cell(merged, bold=False, centered=True)
        old_merged_run = merged_run._r
        with patch.dict(os.environ, {"AUTO_REPORT_FAST_CELL": "1"}):
            generator._set_cell_text(merged, "safe", bold=False, centered=False)
        self.assertIsNot(merged.paragraphs[0].runs[0]._r, old_merged_run)

    def test_newlines_use_safe_fallback(self) -> None:
        cell, paragraph, run = _simple_cell()
        old_run_element = run._r
        with patch.dict(os.environ, {"AUTO_REPORT_FAST_CELL": "1"}):
            generator._set_cell_text(cell, "line 1\nline 2", bold=False, centered=False)
        self.assertIsNot(paragraph.runs[0]._r, old_run_element)
        self.assertEqual(paragraph.text, "line 1\nline 2")

    def test_generation_checks_cancel_between_bounded_row_batches(self) -> None:
        calls = 0
        progress: list[tuple[int, int]] = []

        class CancelledForTest(RuntimeError):
            pass

        def check_cancelled() -> None:
            nonlocal calls
            calls += 1
            if calls >= 3:
                raise CancelledForTest("cancelled")

        with patch.dict(
            os.environ,
            {
                "AUTO_REPORT_FAST_CELL": "1",
                "AUTO_REPORT_PREPARED_TEMPLATE": "1",
            },
        ):
            with self.assertRaises(CancelledForTest):
                generator.generate_report(
                    _fixture_data(8),
                    title="Cancellation test",
                    organization="Reporter Pro",
                    report_type="server_only",
                    check_cancelled=check_cancelled,
                    on_build_progress=lambda value, rows: progress.append((value, rows)),
                )

        self.assertGreaterEqual(calls, 3)
        self.assertTrue(progress)
        self.assertTrue(all(45 <= value <= 66 for value, _rows in progress))

    def test_checkpoint_metrics_are_aggregate_only(self) -> None:
        metrics = PerformanceMetrics()
        tracker = generator._BuildWorkTracker(
            estimated_rows=64,
            check_cancelled=None,
            on_progress=None,
            metrics=metrics,
            batch_size=32,
        )
        tracker.advance(32)
        payload = metrics.public()
        self.assertEqual(payload["aggregates"][0]["name"], "buildCheckpoint")
        self.assertEqual(payload["aggregates"][0]["category"], "rows")
        self.assertNotIn("hostname", str(payload).casefold())


if __name__ == "__main__":
    unittest.main()
